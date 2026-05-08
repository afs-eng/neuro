import os, sys, django
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings")
sys.path.insert(0, "/home/andre/neuro")
django.setup()

from docx import Document
from apps.reports.services.report_export_service import ReportExportService
from types import SimpleNamespace

class ReportStub:
    class Patient:
        full_name = "Leticia"
        birth_date = "2014-05-01"
        age = 11
        sex = "F"
    patient = Patient()
    sections = SimpleNamespace(filter=lambda **kw: [], all=lambda: [])

report = ReportStub()

sr = {"summary": {"qi_verbal": {"qi": 115}, "qi_execucao": {"qi": 123}, "qit_4": {"qi": 122}}}
cp = {"age": {"years": 11}, "composites": {"qi_verbal": {"qi": 115}, "qi_execucao": {"qi": 123}, "qit_4": {"qi": 122}, "subtests": {"vc": {}, "sm": {}, "cb": {}, "rm": {}}}
stale_text = "Leticia Bolonha Lucati obtuvo QIV = 118."
context = {
    "patient": {"full_name": "Leticia", "sex": "F"},
    "validated_tests": [
        {"instrument_code": "wasi", "applied_on": "2025-01-15", "structured_results": sr,
         "classified_payload": sr, "computed_payload": cp, "clinical_interpretation": stale_text}
    ]
}
sections = {
    "procedimentos": "Testes.",
    "eficiencia_intelectual": stale_text,
    "memoria_aprendizagem": "Mem.",
    "atencao": "At.",
    "linguagem": "Ling.",
    "funcoes_executivas": "FE",
    "gnosias_praxias": "GP",
    "conclusao": "Conc.",
}

document = Document(str(ReportExportService.WASI_TEMPLATE_PATH))
ReportExportService._ensure_model_table_styles(document)
ReportExportService._apply_base_styles(document)
ReportExportService._replace_simple_sections(document, report, sections, context)
ReportExportService._rebuild_qualitative_section(document, sections, context)

body = document._body._element
CHART_NS = ReportExportService.CHART_NS
charts = body.findall(".//c:chart", CHART_NS)
print(f"Charts after rebuild: {len(charts)}")

for i, p in enumerate(document.paragraphs):
    if 28 <= i <= 38:
        has_chart = ReportExportService._paragraph_contains_chart(p)
        print(f"  [{i:3}] has_chart={has_chart} text=\"{p.text.strip()[:60]}\"")

ReportExportService._sanitize_generated_document(document, report, context)

charts_after = body.findall(".//c:chart", CHART_NS)
print(f"Charts after sanitize: {len(charts_after)}")

for i, p in enumerate(document.paragraphs):
    if 28 <= i <= 38:
        has_chart = ReportExportService._paragraph_contains_chart(p)
        print(f"  [{i:3}] has_chart={has_chart} text=\"{p.text.strip()[:60]}\"")

for i, p in enumerate(document.paragraphs):
    if "Jo" in (p.text or "") and len(p.text.strip()) > 5:
        print(f"Para [{i}]: {p.text[:100]}")