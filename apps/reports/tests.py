from django.test import SimpleTestCase
from xml.etree import ElementTree as ET
from copy import deepcopy
from io import BytesIO
from zipfile import ZipFile
from datetime import date
from types import SimpleNamespace
from unittest.mock import patch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from lxml import etree as LET

from apps.reports.builders.tests_builder import _build_wais3_tables, build_validated_tests_snapshot
from apps.reports.builders.references_builder import build_references
from apps.reports.services.report_export_service import ReportExportService
from apps.reports.services.report_generation_service import ReportGenerationService
from apps.reports.services.wisc4_standardization import WISC4StandardizationService


class ReferencesBuilderTests(SimpleTestCase):
    def test_build_references_includes_only_validated_tests(self):
        references = build_references(
            [
                {"instrument_code": "wisc4"},
                {"instrument_code": "fdt"},
            ]
        )

        self.assertEqual(len(references), 3)
        self.assertIn("WISC-IV", references[1])
        self.assertIn("Teste dos Cinco Dígitos", references[2])

    @patch("apps.reports.builders.tests_builder.get_test_module", return_value=None)
    @patch("apps.reports.builders.tests_builder.get_validated_test_applications_by_evaluation")
    def test_build_validated_tests_snapshot_skips_empty_validated_stub(self, mocked_selector, _mocked_module):
        evaluation = SimpleNamespace(id=7, patient=SimpleNamespace(full_name="Paciente Exemplo"))
        mocked_selector.return_value = [
            SimpleNamespace(
                id=1,
                evaluation=evaluation,
                evaluation_id=evaluation.id,
                instrument=SimpleNamespace(code="epq_j", name="EPQ-J", category="Personalidade"),
                applied_on=None,
                raw_payload={},
                computed_payload={},
                classified_payload={},
                reviewed_payload={},
                interpretation_text="",
                is_validated=True,
            ),
            SimpleNamespace(
                id=2,
                evaluation=evaluation,
                evaluation_id=evaluation.id,
                instrument=SimpleNamespace(code="bpa2", name="BPA-2", category="Atenção"),
                applied_on=date(2026, 5, 5),
                raw_payload={},
                computed_payload={},
                classified_payload={"subtestes": []},
                reviewed_payload={},
                interpretation_text="",
                is_validated=True,
            ),
        ]

        snapshot = build_validated_tests_snapshot(evaluation)

        self.assertEqual([item["instrument_code"] for item in snapshot], ["bpa2"])

    def test_build_references_deduplicates_equivalent_entries(self):
        references = build_references(
            [
                {"instrument_code": "ebadep_ij"},
                {"instrument_code": "ebaped_ij"},
            ]
        )

        self.assertEqual(len(references), 2)
        self.assertIn("EBADEP-IJ", references[1])


class WISC4StandardizationTests(SimpleTestCase):
    def setUp(self):
        self.context = {
            "patient": {"full_name": "Maria Clara", "sex": "F"},
            "validated_tests": [
                {
                    "instrument_code": "wisc4",
                    "structured_results": {
                        "qit_data": {
                            "escore_composto": 61,
                            "classificacao": "Extremamente Baixo",
                        },
                        "gai_data": {
                            "escore_composto": 70,
                            "classificacao": "Limítrofe",
                        },
                        "cpi_data": {
                            "escore_composto": 58,
                            "classificacao": "Extremamente Baixo",
                        },
                        "indices": [
                            {
                                "indice": "icv",
                                "escore_composto": 84,
                                "classificacao": "Média Inferior",
                            },
                            {
                                "indice": "iop",
                                "escore_composto": 61,
                                "classificacao": "Extremamente Baixo",
                            },
                            {
                                "indice": "imt",
                                "escore_composto": 52,
                                "classificacao": "Extremamente Baixo",
                            },
                            {
                                "indice": "ivp",
                                "escore_composto": 68,
                                "classificacao": "Extremamente Baixo",
                            },
                        ],
                        "subtestes": [
                            {"codigo": "SM", "classificacao": "Média"},
                            {"codigo": "CN", "classificacao": "Média Inferior"},
                            {"codigo": "CO", "classificacao": "Média"},
                            {"codigo": "RM", "classificacao": "Limítrofe"},
                            {"codigo": "VC", "classificacao": "Limítrofe"},
                            {"codigo": "CB", "classificacao": "Extremamente Baixo"},
                            {"codigo": "SNL", "classificacao": "Extremamente Baixo"},
                            {"codigo": "DG", "classificacao": "Limítrofe"},
                        ],
                    },
                }
            ],
        }

    def test_build_global_section_uses_standardized_wisc_template(self):
        text = WISC4StandardizationService.build(
            "capacidade_cognitiva_global", self.context
        )

        self.assertIn("Capacidade Cognitiva Global: a avaliação neuropsicológica de Maria", text)
        self.assertIn("Quociente de Inteligência Total (QIT = 61)", text)
        self.assertIn("Compreensão Verbal (ICV) — 84 — Média Inferior", text)
        self.assertNotIn("Índice de Habilidade Geral (GAI = 70)", text)
        self.assertNotIn("O Índice de Compreensão Verbal", text)

    def test_build_global_section_includes_fixed_cognitive_age_phrase_when_available(self):
        self.context["validated_tests"][0]["structured_results"]["idade_cognitiva"] = "12 anos e 6 meses"

        text = WISC4StandardizationService.build(
            "capacidade_cognitiva_global", self.context
        )

        self.assertIn(
            "com idade cognitiva estimada de 12 anos e 6 meses",
            text,
        )

    def test_build_domain_section_uses_standardized_wisc_subscale_template(self):
        text = WISC4StandardizationService.build("funcoes_executivas", self.context)

        self.assertIn("A avaliação das funções executivas de Maria", text)
        self.assertIn("No subteste Semelhanças", text)
        self.assertIn("No subteste Raciocínio Matricial", text)
        self.assertIn("perfil executivo heterogêneo", text)


class ReportGenerationServiceConclusionTests(SimpleTestCase):
    def test_wisc4_adolescent_conclusion_follows_fixed_model_order(self):
        class Section:
            def __init__(self, key, generated):
                self.key = key
                self.content_edited = ""
                self.content_generated = generated

        class Sections(list):
            def all(self):
                return self

            def order_by(self, _field):
                return self

            def filter(self, **kwargs):
                key = kwargs.get("key")
                return Sections([item for item in self if item.key == key])

            def first(self):
                return self[0] if self else None

        class Report:
            sections = Sections(
                [
                    Section("capacidade_cognitiva_global", "Perfil heterogêneo com fragilidades cognitivas."),
                    Section("funcoes_executivas", "Fragilidades executivas e atencionais com oscilação clínica."),
                    Section("memoria_aprendizagem", "Perfil heterogêneo na memória e aprendizagem."),
                    Section("linguagem", "Recursos de linguagem relativamente preservados."),
                    Section("gnosias_praxias", "Fragilidades visuoconstrutivas pontuais."),
                    Section("aspectos_emocionais_comportamentais", "Indicadores emocionais e comportamentais relevantes."),
                    Section("srs2", "Indicadores sociais que exigem correlação clínica."),
                ]
            )

        context = {
            "patient": {"full_name": "Maria Clara", "sex": "F", "schooling": "fundamental", "birth_date": "2010-01-10"},
            "evaluation": {"clinical_hypothesis": "Transtorno do Déficit de Atenção e Hiperatividade (TDAH), apresentação combinada"},
            "validated_tests": [
                {
                    "instrument_code": "wisc4",
                    "structured_results": {
                        "qit_data": {"escore_composto": 61, "classificacao": "Extremamente Baixo"},
                        "indices": [
                            {"indice": "icv", "escore_composto": 84, "classificacao": "Média Inferior"},
                            {"indice": "iop", "escore_composto": 61, "classificacao": "Extremamente Baixo"},
                            {"indice": "imt", "escore_composto": 52, "classificacao": "Extremamente Baixo"},
                            {"indice": "ivp", "escore_composto": 68, "classificacao": "Extremamente Baixo"},
                        ],
                    },
                },
                {"instrument_code": "srs2"},
            ],
        }
        context["evaluation"]["start_date"] = "2024-04-20"

        text = ReportGenerationService._build_conclusion_text(Report(), context)

        self.assertIn("Maria Clara apresenta funcionamento neuropsicológico", text)
        self.assertIn("Na Escala Wechsler de Inteligência para Crianças – Quarta Edição (WISC-IV)", text)
        self.assertIn("No domínio atencional e executivo", text)
        self.assertIn("A avaliação da memória e aprendizagem", text)
        self.assertIn("Nos domínios de linguagem, gnosias e praxias", text)
        self.assertIn("Nos aspectos emocionais e comportamentais", text)
        self.assertIn("Na avaliação do funcionamento social", text)
        self.assertIn("Diante da integração dos resultados das testagens", text)
        self.assertIn("Hipótese Diagnóstica:", text)
        self.assertIn("Ressalta-se que o ser humano possui natureza dinâmica", text)
        self.assertNotIn("verifica-se", text.casefold())


class WISC4ExportTableTests(SimpleTestCase):
    def setUp(self):
        self.context = {
            "patient": {"birth_date": "2016-01-10"},
            "evaluation": {"start_date": "2024-04-20"},
            "validated_tests": [
                {
                    "instrument_code": "wisc4",
                    "applied_on": "2024-04-20",
                }
            ],
        }
        self.test = {
            "classified_payload": {
                "subtestes": [
                    {
                        "codigo": "SM",
                        "subteste": "Semelhanças",
                        "escore_bruto": 23,
                        "classificacao": "Média",
                    },
                    {
                        "codigo": "CN",
                        "subteste": "Conceitos Figurativos",
                        "escore_bruto": 11,
                        "classificacao": "Média Inferior",
                    },
                    {
                        "codigo": "CO",
                        "subteste": "Compreensão",
                        "escore_bruto": 24,
                        "classificacao": "Média",
                    },
                    {
                        "codigo": "RM",
                        "subteste": "Raciocínio Matricial",
                        "escore_bruto": 12,
                        "classificacao": "Limítrofe",
                    },
                ]
            }
        }

    def test_wisc_reference_scores_use_age_norm_table_ranges(self):
        table = ReportExportService._wisc_ncp_table(self.context)

        self.assertEqual(
            ReportExportService._wisc_reference_scores(table, "SM"),
            ("32-44", "6-16", "1"),
        )

    def test_wisc_rows_build_expected_columns_for_executive_domain(self):
        rows = ReportExportService._wisc_rows(
            self.test, self.context, "funcoes_executivas"
        )

        self.assertEqual(
            rows[0],
            [
                "Testes Utilizados",
                "Escore Máximo",
                "Escore Médio",
                "Escore Mínimo",
                "Escore Bruto",
                "Classificação",
            ],
        )
        self.assertEqual(rows[1], ["Semelhanças", "32-44", "6-16", "1", "23", "Média"])


class WAIS3ExportTableTests(SimpleTestCase):
    def _evaluation(self):
        class Patient:
            birth_date = date(2000, 1, 1)

        class Evaluation:
            patient = Patient()
            start_date = date(2020, 6, 1)

        return Evaluation()

    def test_build_wais3_tables_uses_normative_ranges_and_spontaneous_speech(self):
        payload = {
            "subtestes": {
                "semelhancas": {"nome": "Semelhanças", "pontos_brutos": 18, "classificacao": "Média"},
                "vocabulario": {"nome": "Vocabulário", "pontos_brutos": 28, "classificacao": "Média"},
                "compreensao": {"nome": "Compreensão", "pontos_brutos": 16, "classificacao": "Média"},
            }
        }

        tables = _build_wais3_tables(payload, self._evaluation(), date(2020, 6, 1))

        self.assertEqual(tables["linguagem"][0], {
            "label": "Semelhanças",
            "maxScore": "38",
            "avgScore": "17-28",
            "minScore": "9-10",
            "obtainedScore": "18",
            "classification": "Média",
        })
        self.assertEqual(tables["linguagem"][1], {
            "label": "Vocabulário",
            "maxScore": "66",
            "avgScore": "23-42",
            "minScore": "11-14",
            "obtainedScore": "28",
            "classification": "Média",
        })
        self.assertEqual(tables["linguagem"][2], {
            "label": "Compreensão",
            "maxScore": "33",
            "avgScore": "13-26",
            "minScore": "5-6",
            "obtainedScore": "16",
            "classification": "Média",
        })
        self.assertEqual(tables["linguagem"][3], {
            "label": "Fala Espontânea",
            "note": "Fala espontânea dentro do esperado para a sua idade",
        })

    def test_build_wais3_tables_recomputes_classification_from_normative_table(self):
        payload = {
            "subtestes": {
                "vocabulario": {
                    "nome": "Vocabulário",
                    "pontos_brutos": 39,
                    "classificacao": "Média Superior",
                },
            }
        }

        tables = _build_wais3_tables(payload, self._evaluation(), date(2020, 6, 1))

        self.assertEqual(tables["linguagem"][0], {
            "label": "Vocabulário",
            "maxScore": "66",
            "avgScore": "23-42",
            "minScore": "11-14",
            "obtainedScore": "39",
            "classification": "Média",
        })

    def test_build_wais3_tables_uses_normative_classification_across_domains(self):
        payload = {
            "subtestes": {
                "semelhancas": {"nome": "Semelhanças", "pontos_brutos": 28, "classificacao": "Média Superior"},
                "vocabulario": {"nome": "Vocabulário", "pontos_brutos": 42, "classificacao": "Média Superior"},
                "compreensao": {"nome": "Compreensão", "pontos_brutos": 26, "classificacao": "Média Superior"},
                "raciocinio_matricial": {"nome": "Raciocínio Matricial", "pontos_brutos": 21, "classificacao": "Média Superior"},
                "cubos": {"nome": "Cubos", "pontos_brutos": 44, "classificacao": "Média Superior"},
                "sequencia_numeros_letras": {"nome": "Sequência de Números e Letras", "pontos_brutos": 11, "classificacao": "Média Superior"},
                "digitos": {"nome": "Dígitos", "pontos_brutos": 17, "classificacao": "Média Superior"},
            }
        }

        tables = _build_wais3_tables(payload, self._evaluation(), date(2020, 6, 1))

        expected = {
            "linguagem": {
                "Semelhanças": "Média",
                "Vocabulário": "Média",
                "Compreensão": "Média",
            },
            "gnosias_praxias": {
                "Raciocínio Matricial": "Média",
                "Cubos": "Média",
            },
            "funcoes_executivas": {
                "Semelhanças": "Média",
                "Compreensão": "Média",
                "Raciocínio Matricial": "Média",
            },
            "memoria_aprendizagem": {
                "Sequência de Números e Letras": "Média",
                "Dígitos": "Média",
                "RAVLT": "Leitura do Gráfico",
            },
        }

        for domain, rows in tables.items():
            for row in rows:
                if row.get("note"):
                    continue
                self.assertEqual(row["classification"], expected[domain][row["label"]])

    def test_build_wais3_tables_keeps_non_average_normative_classification(self):
        payload = {
            "subtestes": {
                "vocabulario": {"nome": "Vocabulário", "pontos_brutos": 11, "classificacao": "Média Superior"},
                "cubos": {"nome": "Cubos", "pontos_brutos": 45, "classificacao": "Média"},
            }
        }

        tables = _build_wais3_tables(payload, self._evaluation(), date(2020, 6, 1))

        self.assertEqual(tables["linguagem"][0]["classification"], "Limítrofe")
        self.assertEqual(tables["gnosias_praxias"][0]["classification"], "Média Superior")

    def test_wais3_domain_rows_use_dynamic_template_structure(self):
        test = {
            "wais3_tables": {
                "linguagem": [
                    {
                        "label": "Semelhanças",
                        "maxScore": "38",
                        "avgScore": "17-28",
                        "minScore": "9-10",
                        "obtainedScore": "18",
                        "classification": "Média",
                    },
                    {
                        "label": "Fala Espontânea",
                        "note": "Fala espontânea dentro do esperado para a sua idade",
                    },
                ]
            }
        }

        rows = ReportExportService._wais3_domain_rows(test, "linguagem")

        self.assertEqual(rows[0], ["Testes Utilizados", "Escore Máximo", "Escore Médio", "Escore Mínimo", "Escore Bruto", "Classificação"])
        self.assertEqual(rows[1], ["Semelhanças", "38", "17-28", "9-10", "18", "Média"])
        self.assertEqual(rows[2], ["Fala Espontânea", "Fala espontânea dentro do esperado para a sua idade", "", "", "", ""])

    def test_build_wais3_tables_adds_ravlt_placeholder_to_memory_domain(self):
        payload = {
            "subtestes": {
                "sequencia_numeros_letras": {"nome": "Sequência de Números e Letras", "pontos_brutos": 11, "classificacao": "Média"},
                "digitos": {"nome": "Dígitos", "pontos_brutos": 17, "classificacao": "Média"},
            }
        }

        tables = _build_wais3_tables(payload, self._evaluation(), date(2020, 6, 1))

        self.assertEqual(tables["memoria_aprendizagem"][-1], {
            "label": "RAVLT",
            "maxScore": "-",
            "avgScore": "-",
            "minScore": "-",
            "obtainedScore": "-",
            "classification": "Leitura do Gráfico",
        })


class ReportExportChartSanitizationTests(SimpleTestCase):
    def test_primary_report_test_code_prefers_wisc4_then_wais3_then_wasi(self):
        self.assertEqual(
            ReportExportService._primary_report_test_code(
                {"validated_tests": [{"instrument_code": "wais3"}, {"instrument_code": "bfp"}]}
            ),
            "wais3",
        )
        self.assertEqual(
            ReportExportService._primary_report_test_code(
                {"validated_tests": [{"instrument_code": "wasi"}, {"instrument_code": "wais3"}]}
            ),
            "wais3",
        )
        self.assertEqual(
            ReportExportService._primary_report_test_code(
                {"validated_tests": [{"instrument_code": "wisc4"}, {"instrument_code": "wais3"}]}
            ),
            "wisc4",
        )

    def test_select_template_path_uses_primary_report_test(self):
        report = SimpleNamespace(patient=SimpleNamespace(age=18))
        self.assertEqual(
            ReportExportService._select_template_path(
                report,
                {"validated_tests": [{"instrument_code": "wais3"}, {"instrument_code": "bfp"}]},
            ),
            ReportExportService.WAIS3_TEMPLATE_PATH,
        )

    def test_etdah_table_title_distinguishes_ad_from_pais(self):
        self.assertEqual(ReportExportService._table_title_text("etdah_ad"), "E-TDAH-AD")
        self.assertEqual(ReportExportService._table_title_text("etdah_pais"), "E-TDAH-PAIS")

    def test_etdah_ad_rows_follow_official_factor_order(self):
        rows = ReportExportService._etdah_rows(
            {
                "classified_payload": {
                    "results": {
                        "D": {"name": "Fator 1 - Desatenção (D)", "raw_score": 64, "mean": 42.3, "percentile_text": "86", "classification": "Superior"},
                        "H": {"name": "Fator 5 - Hiperatividade (H)", "raw_score": 20, "mean": 14.35, "percentile_text": "85", "classification": "Superior"},
                        "I": {"name": "Fator 2 - Impulsividade (I)", "raw_score": 57, "mean": 44.3, "percentile_text": "78,8", "classification": "Média Superior"},
                        "AE": {"name": "Fator 3 - Aspectos Emocionais (AE)", "raw_score": 6, "mean": 5.77, "percentile_text": "60", "classification": "Média"},
                        "AAMA": {"name": "Fator 4 - Autorregulação da Atenção, Motivação e Ação (AAMA)", "raw_score": 25, "mean": 21.1, "percentile_text": "75", "classification": "Média Superior"},
                    }
                }
            }
        )

        self.assertEqual(
            [row[0] for row in rows[1:]],
            [
                "Fator 1 - Desatenção (D)",
                "Fator 2 - Impulsividade (I)",
                "Fator 3 - Aspectos Emocionais (AE)",
                "Fator 4 - Autorregulação da Atenção, Motivação e Ação (AAMA)",
                "Fator 5 - Hiperatividade (H)",
            ],
        )

    def test_insert_etdah_ad_interpretation_formats_factor_heading_and_body(self):
        document = Document()
        start = document.add_paragraph("ANÁLISE QUALITATIVA")
        interpretation = (
            "Interpretação e Observações Clínicas: **Texto introdutório.**\n\n"
            "No Fator 1 — Desatenção, o resultado classificado como **Média Inferior** indica funcionamento dentro dos limites esperados.\n\n"
            "Dessa forma, **não** há hipótese diagnóstica isolada."
        )

        ReportExportService._insert_etdah_ad_interpretation_block_after(
            start,
            ReportExportService._normalize_interpretation_text(interpretation),
        )

        paragraphs = [p for p in document.paragraphs if p.text.strip()]

        self.assertEqual(paragraphs[1].text.strip(), "Interpretação e Observações Clínicas: Texto introdutório.")
        self.assertEqual(paragraphs[2].text.strip(), "Fator 1 – Desatenção")
        self.assertTrue(any(run.bold for run in paragraphs[2].runs if run.text.strip()))
        self.assertEqual(paragraphs[3].text.strip(), "o resultado classificado como Média Inferior indica funcionamento dentro dos limites esperados.")
        self.assertEqual(paragraphs[3].alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
        self.assertAlmostEqual(paragraphs[3].paragraph_format.first_line_indent.cm, 1.5, places=2)
        self.assertEqual(paragraphs[4].text.strip(), "Dessa forma, não há hipótese diagnóstica isolada.")

    def test_wais3_template_is_organized_under_laudos_templates_assets(self):
        self.assertEqual(
            ReportExportService.WAIS3_TEMPLATE_PATH,
            ReportExportService.LAUDOS_TEMPLATE_DIR / "Modelo-WAIS3.docx",
        )
        self.assertTrue(ReportExportService.WAIS3_TEMPLATE_PATH.exists())

    def test_wisc4_template_is_organized_under_laudos_templates_assets(self):
        self.assertEqual(
            ReportExportService.WISC4_TEMPLATE_PATH,
            ReportExportService.LAUDOS_TEMPLATE_DIR / "Modelo-WISC4.docx",
        )
        self.assertTrue(ReportExportService.WISC4_TEMPLATE_PATH.exists())

    def test_wasi_template_is_organized_under_laudos_templates_assets(self):
        self.assertEqual(
            ReportExportService.WASI_TEMPLATE_PATH,
            ReportExportService.LAUDOS_TEMPLATE_DIR / "Modelo-WASI.docx",
        )
        self.assertTrue(ReportExportService.WASI_TEMPLATE_PATH.exists())

    def test_extract_template_chart_blocks_detects_wasi_native_charts(self):
        document = Document(str(ReportExportService.WASI_TEMPLATE_PATH))

        self.assertGreaterEqual(
            len(ReportExportService._extract_template_chart_blocks(document)),
            1,
        )

    def test_replace_simple_sections_rebuilds_wasi_procedures_block_from_context(self):
        document = Document(str(ReportExportService.WASI_TEMPLATE_PATH))

        ReportExportService._replace_simple_sections(
            document,
            self._report_stub(),
            sections={},
            context={
                "patient": {"full_name": "Paciente", "sex": "M", "birth_date": "2008-03-20"},
                "validated_tests": [
                    {"instrument_code": "wasi"},
                    {"instrument_code": "bpa2"},
                ],
            },
        )

        texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        procedures_texts = texts[texts.index("PROCEDIMENTOS"):texts.index("ANÁLISE")]

        self.assertIn(
            "• Escala Wechsler Abreviada de Inteligência (WASI): Utilizada para estimar o funcionamento intelectual global, verbal e de execução, bem como fornecer indicadores breves sobre a capacidade cognitiva geral.",
            procedures_texts,
        )
        self.assertIn(
            "• Bateria Psicológica para Avaliação da Atenção – Segunda Edição (BPA-2): Avalia a capacidade geral de atenção, incluindo atenção concentrada, alternada, dividida e atenção global.",
            procedures_texts,
        )
        self.assertFalse(any("WAIS-III" in text for text in procedures_texts))
        self.assertFalse(any("IPHEXA" in text for text in procedures_texts))

    def test_procedures_intro_text_ignores_embedded_bullet_list(self):
        text = ReportExportService._wais3_procedures_intro_text(
            self._report_stub(),
            {
                "procedimentos": "Para esta avaliação foram realizadas entrevista clínica inicial.\n• Anamnese: Texto antigo\n• BPA-2: Texto antigo"
            },
        )

        self.assertEqual(text, "Para esta avaliação foram realizadas entrevista clínica inicial.")

    def test_insert_bullet_paragraph_formats_label_in_bold(self):
        document = Document()
        anchor = document.add_paragraph("PROCEDIMENTOS")

        paragraph = ReportExportService._insert_bullet_paragraph_after(
            anchor,
            "Bateria Psicológica para Avaliação da Atenção – Segunda Edição (BPA-2): Avalia a atenção.",
        )

        self.assertEqual(paragraph.text, "• Bateria Psicológica para Avaliação da Atenção – Segunda Edição (BPA-2): Avalia a atenção.")
        label_run = next(run for run in paragraph.runs if run.text == "Bateria Psicológica para Avaliação da Atenção – Segunda Edição (BPA-2)")
        description_run = next(run for run in paragraph.runs if run.text == ": Avalia a atenção.")
        self.assertTrue(label_run.bold)
        self.assertFalse(description_run.bold)

    def test_rebuild_qualitative_section_for_wasi_omits_missing_sections(self):
        document = Document(str(ReportExportService.WASI_TEMPLATE_PATH))

        ReportExportService._replace_simple_sections(
            document,
            self._report_stub(),
            sections={
                "eficiencia_intelectual": "Interpretação WASI.",
                "memoria_aprendizagem": "Texto memória.",
            },
            context={
                "patient": {"full_name": "Paciente", "sex": "M", "birth_date": "2008-03-20"},
                "validated_tests": [
                    {
                        "instrument_code": "wasi",
                        "structured_results": {
                            "composites": {
                                "qi_verbal": {"qi": 118, "classification": "Media Superior"},
                                "qi_execucao": {"qi": 120, "classification": "Superior"},
                                "qit_4": {"qi": 122, "classification": "Superior"},
                            }
                        },
                    },
                    {
                        "instrument_code": "ravlt",
                        "classified_payload": {
                            "chart": {
                                "title": "RAVLT",
                                "labels": ["A1", "A2"],
                                "series": [{"label": "Paciente", "values": [4, 5], "color": "#70AD47"}],
                                "y_axis": {"min": 0, "max": 10},
                            }
                        },
                    },
                ],
            },
        )
        ReportExportService._rebuild_qualitative_section(
            document,
            sections={
                "eficiencia_intelectual": "Interpretação WASI.",
                "memoria_aprendizagem": "Texto memória.",
            },
            context={
                "patient": {"full_name": "Paciente", "sex": "M", "birth_date": "2008-03-20"},
                "validated_tests": [
                    {
                        "instrument_code": "wasi",
                        "structured_results": {
                            "composites": {
                                "qi_verbal": {"qi": 118, "classification": "Media Superior"},
                                "qi_execucao": {"qi": 120, "classification": "Superior"},
                                "qit_4": {"qi": 122, "classification": "Superior"},
                            }
                        },
                    },
                    {
                        "instrument_code": "ravlt",
                        "classified_payload": {
                            "chart": {
                                "title": "RAVLT",
                                "labels": ["A1", "A2"],
                                "series": [{"label": "Paciente", "values": [4, 5], "color": "#70AD47"}],
                                "y_axis": {"min": 0, "max": 10},
                            }
                        },
                    },
                ],
            },
        )

        texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

        self.assertIn("5.1. Memória e Aprendizagem", texts)
        self.assertNotIn("5.1. Atenção", texts)
        self.assertNotIn("5.2. Funções Executivas", texts)
        self.assertNotIn("5.3. Aspectos Emocionais, Comportamentais e Escalas Complementares", texts)

    def test_rebuild_qualitative_section_for_wasi_does_not_include_epq_j_without_validated_test(self):
        document = Document(str(ReportExportService.WASI_TEMPLATE_PATH))
        context = {
            "patient": {"full_name": "Paciente", "sex": "M", "birth_date": "2008-03-20"},
            "validated_tests": [
                {
                    "instrument_code": "wasi",
                    "structured_results": {
                        "composites": {
                            "qi_verbal": {"qi": 118, "classification": "Media Superior"},
                            "qi_execucao": {"qi": 120, "classification": "Superior"},
                            "qit_4": {"qi": 122, "classification": "Superior"},
                        }
                    },
                },
                {
                    "instrument_code": "bpa2",
                    "bpa_chart_data": {
                        "labels": ["AC", "AD", "AA", "AG"],
                        "series": [{"name": "Percentil", "values": [75, 70, 50, 70]}],
                    },
                    "classified_payload": {
                        "results": {
                            "ac": {"percentile": 75, "classification": "Media Superior"},
                            "ad": {"percentile": 70, "classification": "Media Superior"},
                            "aa": {"percentile": 50, "classification": "Media"},
                            "ag": {"percentile": 70, "classification": "Media Superior"},
                        }
                    },
                },
                {
                    "instrument_code": "fdt",
                    "classified_payload": {
                        "results": {
                            "leitura": {"tempo": 31.91, "erros": 0, "percentil_num": 5},
                            "contagem": {"tempo": 32.46, "erros": 0, "percentil_num": 5},
                            "escolha": {"tempo": 62.98, "erros": 1, "percentil_num": 5},
                            "alternancia": {"tempo": 71.85, "erros": 4, "percentil_num": 5},
                        }
                    },
                },
                {"instrument_code": "etdah_ad"},
                {
                    "instrument_code": "ravlt",
                    "classified_payload": {
                        "chart": {
                            "title": "RAVLT",
                            "labels": ["A1", "A2"],
                            "series": [{"label": "Paciente", "values": [4, 5], "color": "#70AD47"}],
                            "y_axis": {"min": 0, "max": 10},
                        }
                    },
                },
            ],
        }

        sections = {
            "eficiencia_intelectual": "Interpretação WASI.",
            "atencao": "Texto atenção.",
            "funcoes_executivas": "Texto funções executivas.",
            "etdah_ad": "Texto ETDAH-AD.",
            "memoria_aprendizagem": "Texto memória.",
        }

        ReportExportService._replace_simple_sections(
            document,
            self._report_stub(),
            sections=sections,
            context=context,
        )
        ReportExportService._rebuild_qualitative_section(
            document,
            sections=sections,
            context=context,
        )

        texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

        self.assertIn("5.1. Atenção", texts)
        self.assertIn("5.2. Funções Executivas", texts)
        self.assertIn("5.3. Aspectos Emocionais, Comportamentais e Escalas Complementares", texts)
        self.assertNotIn("EPQ-J", texts)
        self.assertFalse(any("EPQ-J" in text for text in texts))

    def test_rebuild_qualitative_section_for_wasi_preserves_native_template_charts(self):
        document = Document(str(ReportExportService.WASI_TEMPLATE_PATH))

        context = {
            "patient": {"full_name": "Paciente", "sex": "M", "birth_date": "2008-03-20"},
            "validated_tests": [
                {
                    "instrument_code": "wasi",
                    "structured_results": {
                        "composites": {
                            "qi_verbal": {"qi": 118, "classification": "Media Superior"},
                            "qi_execucao": {"qi": 120, "classification": "Superior"},
                            "qit_4": {"qi": 122, "classification": "Superior"},
                        }
                    },
                },
                {
                    "instrument_code": "bpa2",
                    "bpa_chart_data": {
                        "labels": ["AC", "AD", "AA", "AG"],
                        "series": [{"name": "Percentil", "values": [75, 70, 50, 70]}],
                    },
                    "classified_payload": {
                        "results": {
                            "ac": {"percentile": 75, "classification": "Media Superior"},
                            "ad": {"percentile": 70, "classification": "Media Superior"},
                            "aa": {"percentile": 50, "classification": "Media"},
                            "ag": {"percentile": 70, "classification": "Media Superior"},
                        }
                    },
                },
                {
                    "instrument_code": "ravlt",
                    "classified_payload": {
                        "chart": {
                            "title": "RAVLT",
                            "labels": ["A1", "A2"],
                            "series": [{"label": "Paciente", "values": [4, 5], "color": "#70AD47"}],
                            "y_axis": {"min": 0, "max": 10},
                        }
                    },
                },
            ],
        }

        ReportExportService._replace_simple_sections(
            document,
            self._report_stub(),
            sections={"eficiencia_intelectual": "Interpretação WASI.", "memoria_aprendizagem": "Texto memória."},
            context=context,
        )
        ReportExportService._rebuild_qualitative_section(
            document,
            sections={"eficiencia_intelectual": "Interpretação WASI.", "memoria_aprendizagem": "Texto memória."},
            context=context,
        )

        self.assertGreaterEqual(len(ReportExportService._extract_template_chart_blocks(document)), 3)

    def test_wasi_subscale_rows_use_expected_table_structure(self):
        test = {
            "structured_results": {
                "subtests": {
                    "vc": {"name": "Vocabulário", "raw_score": 61, "t_score": 56, "classification": "Média Superior"},
                    "sm": {"name": "Semelhanças", "raw_score": 61, "t_score": 57, "classification": "Média Superior"},
                    "cb": {"name": "Cubos", "raw_score": 53, "t_score": 50, "classification": "Média"},
                    "rm": {"name": "Raciocínio Matricial", "raw_score": 71, "t_score": 64, "classification": "Superior"},
                }
            }
        }

        verbal_rows = ReportExportService._wasi_subscale_rows(test, "verbal")
        execution_rows = ReportExportService._wasi_subscale_rows(test, "execucao")

        self.assertEqual(verbal_rows[0], ["Testes Utilizados", "Escore Máximo", "Escore Médio", "Escore Mínimo", "Escore Obtido", "Classificação"])
        self.assertEqual(verbal_rows[1], ["Vocabulário", "80", "40 - 60", "20", "56", "Média Superior"])
        self.assertEqual(verbal_rows[2], ["Semelhanças", "80", "40 - 60", "20", "57", "Média Superior"])
        self.assertEqual(execution_rows[1], ["Cubos", "80", "40 - 60", "20", "50", "Média"])
        self.assertEqual(execution_rows[2], ["Raciocínio Matricial", "80", "40 - 60", "20", "64", "Superior"])

    def test_rebuild_qualitative_section_for_wasi_includes_subscale_sections(self):
        document = Document(str(ReportExportService.WASI_TEMPLATE_PATH))
        context = {
            "patient": {"full_name": "João Vitor", "sex": "M", "birth_date": "2008-03-20"},
            "validated_tests": [
                {
                    "instrument_code": "wasi",
                    "structured_results": {
                        "composites": {
                            "qi_verbal": {"qi": 118, "classification": "Média Superior"},
                            "qi_execucao": {"qi": 120, "classification": "Superior"},
                            "qit_4": {"qi": 122, "classification": "Superior"},
                        },
                        "subtests": {
                            "vc": {"name": "Vocabulário", "raw_score": 61, "t_score": 56, "classification": "Média Superior"},
                            "sm": {"name": "Semelhanças", "raw_score": 61, "t_score": 57, "classification": "Média Superior"},
                            "cb": {"name": "Cubos", "raw_score": 53, "t_score": 50, "classification": "Média"},
                            "rm": {"name": "Raciocínio Matricial", "raw_score": 71, "t_score": 64, "classification": "Superior"},
                        },
                    },
                },
            ],
        }

        ReportExportService._replace_simple_sections(document, self._report_stub(), sections={}, context=context)
        ReportExportService._rebuild_qualitative_section(document, sections={}, context=context)

        texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        self.assertIn("5.2. Subescalas WASI", texts)
        self.assertIn("5.2.1. Escala Verbal", texts)
        self.assertIn("5.2.2. Escala de Execução", texts)
        self.assertTrue(any(text.startswith("Interpretação: A avaliação da escala verbal") for text in texts))
        self.assertTrue(any(text.startswith("Interpretação: A avaliação da escala de execução") for text in texts))

    def test_rebuild_qualitative_section_for_wasi_skips_bfp_png_chart_fallback(self):
        document = Document(str(ReportExportService.WASI_TEMPLATE_PATH))
        context = {
            "patient": {"full_name": "Paciente", "sex": "M", "birth_date": "2008-03-20"},
            "validated_tests": [
                {
                    "instrument_code": "wasi",
                    "structured_results": {
                        "composites": {
                            "qi_verbal": {"qi": 118, "classification": "Média Superior"},
                            "qi_execucao": {"qi": 120, "classification": "Superior"},
                            "qit_4": {"qi": 122, "classification": "Superior"},
                        }
                    },
                },
                {
                    "instrument_code": "bfp",
                    "computed_payload": {
                        "factor_order": ["NN", "EE", "SS", "RR", "AA"],
                        "factors": {
                            "NN": {"name": "Neuroticismo", "raw_score": 4.2, "percentile": 88, "classification": "Superior"},
                            "EE": {"name": "Extroversão", "raw_score": 3.3, "percentile": 22, "classification": "Média Inferior"},
                            "SS": {"name": "Socialização", "raw_score": 4.8, "percentile": 50, "classification": "Média"},
                            "RR": {"name": "Realização", "raw_score": 3.0, "percentile": 18, "classification": "Média Inferior"},
                            "AA": {"name": "Abertura", "raw_score": 4.1, "percentile": 60, "classification": "Média"},
                        },
                    },
                    "clinical_interpretation": "Texto interpretativo do BFP.",
                },
            ],
        }

        ReportExportService._replace_simple_sections(document, self._report_stub(), sections={}, context=context)
        ReportExportService._rebuild_qualitative_section(
            document,
            sections={"aspectos_emocionais_comportamentais": "Texto integrado do BFP."},
            context=context,
        )

        body_children = list(document._body._element)
        bfp_index = next(
            index
            for index, child in enumerate(body_children)
            if "BFP – Bateria Fatorial de Personalidade" in "".join(child.itertext())
        )
        conclusion_index = next(
            index
            for index, child in enumerate(body_children)
            if "Conclusão" in "".join(child.itertext())
        )

        drawings_between = [
            child
            for child in body_children[bfp_index:conclusion_index]
            if child.findall('.//w:drawing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        ]

        self.assertEqual(drawings_between, [])

    def test_wasi_payload_merges_snapshot_structured_and_computed_data(self):
        test = {
            "structured_results": {
                "results": {"status": "ok"},
            },
            "computed_payload": {
                "composites": {
                    "qi_verbal": {"qi": 118, "classification": "Média Superior"},
                    "qi_execucao": {"qi": 120, "classification": "Superior"},
                    "qit_4": {"qi": 122, "classification": "Superior"},
                },
                "subtests": {
                    "vc": {"name": "Vocabulário", "raw_score": 61, "t_score": 56, "classification": "Média Superior"},
                    "sm": {"name": "Semelhanças", "raw_score": 61, "t_score": 57, "classification": "Média Superior"},
                    "cb": {"name": "Cubos", "raw_score": 53, "t_score": 50, "classification": "Média"},
                    "rm": {"name": "Raciocínio Matricial", "raw_score": 71, "t_score": 64, "classification": "Superior"},
                },
            },
        }

        intro = ReportExportService._wasi_intro_text(test, {"patient": {"full_name": "João Vitor"}})
        chart_labels, chart_values = ReportExportService._wasi_chart_payload(test)
        verbal_rows = ReportExportService._wasi_subscale_rows(test, "verbal")

        self.assertIn("QIT = 122", intro)
        self.assertEqual(chart_labels, ["QIE", "QIV", "QI TOTAL"])
        self.assertEqual(chart_values, [120.0, 118.0, 122.0])
        self.assertEqual(verbal_rows[1], ["Vocabulário", "80", "40 - 60", "20", "56", "Média Superior"])

    def test_strip_markdown_heading_prefix_for_wasi_sections(self):
        self.assertEqual(
            ReportExportService._strip_markdown_heading_prefix("**Linguagem**\nTexto", "Linguagem"),
            "Texto",
        )
        self.assertEqual(
            ReportExportService._strip_markdown_heading_prefix("**Gnosias e Praxias**\nTexto", "Gnosias e Praxias"),
            "Texto",
        )

    def test_adult_reports_restore_header_footer_from_papel_timbrado(self):
        class Patient:
            age = 23

        class Report:
            patient = Patient()

        template_path = ReportExportService.WAIS3_TEMPLATE_PATH

        self.assertEqual(
            ReportExportService._header_footer_template_path(
                template_path,
                Report(),
                {"patient": {"birth_date": "2002-09-12"}},
            ),
            ReportExportService.DEFAULT_TEMPLATE_PATH,
        )

    def test_wais3_reports_restore_header_footer_from_own_template(self):
        class Patient:
            age = 23

        class Report:
            patient = Patient()

        template_path = ReportExportService.WAIS3_TEMPLATE_PATH

        self.assertEqual(
            ReportExportService._header_footer_template_path(
                template_path,
                Report(),
                {"patient": {"birth_date": "2002-09-12"}, "validated_tests": [{"instrument_code": "wais3"}]},
            ),
            template_path,
        )

    def _report_stub(self):
        class Report:
            evaluation = None
            interested_party = ""
            purpose = ""

        return Report()

    def test_apply_base_styles_preserves_template_header_footer_layout(self):
        document = Document(str(ReportExportService.WAIS3_TEMPLATE_PATH))
        section = document.sections[0]

        original = {
            "top": section.top_margin,
            "bottom": section.bottom_margin,
            "left": section.left_margin,
            "right": section.right_margin,
            "header": section.header_distance,
            "footer": section.footer_distance,
        }

        ReportExportService._apply_base_styles(document)

        self.assertEqual(section.top_margin, original["top"])
        self.assertEqual(section.bottom_margin, original["bottom"])
        self.assertEqual(section.left_margin, original["left"])
        self.assertEqual(section.right_margin, original["right"])
        self.assertEqual(section.header_distance, original["header"])
        self.assertEqual(section.footer_distance, original["footer"])

    def test_restore_template_header_footer_replaces_exported_header_with_template(self):
        template_bytes = ReportExportService.WAIS3_TEMPLATE_PATH.read_bytes()
        mutated = template_bytes.replace(b"word/media/image4.jpeg", b"word/media/image1.jpeg")

        restored = ReportExportService._restore_template_header_footer(
            mutated,
            ReportExportService.WAIS3_TEMPLATE_PATH,
        )

        with ZipFile(BytesIO(restored), "r") as archive:
            header_xml = archive.read("word/header1.xml")
            header_rels = archive.read("word/_rels/header1.xml.rels")

        with ZipFile(BytesIO(template_bytes), "r") as template_archive:
            expected_header_xml = template_archive.read("word/header1.xml")
            expected_header_rels = template_archive.read("word/_rels/header1.xml.rels")

        self.assertEqual(header_xml, expected_header_xml)
        self.assertEqual(header_rels, expected_header_rels)

    def test_restore_template_header_footer_restores_document_section_references(self):
        template_bytes = ReportExportService.WAIS3_TEMPLATE_PATH.read_bytes()

        with ZipFile(BytesIO(template_bytes), "r") as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")

        mutated_document_xml = document_xml.replace(
            '<w:sectPr w:rsidR="00F43DEA" w:rsidRPr="00A64CAA" w:rsidSect="00C77ADB"><w:headerReference w:type="default" r:id="rId18"/><w:footerReference w:type="default" r:id="rId19"/><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1701" w:right="1134" w:bottom="1304" w:left="1418" w:header="397" w:footer="283" w:gutter="0"/><w:cols w:space="720"/><w:formProt w:val="0"/><w:docGrid w:linePitch="326" w:charSpace="8192"/></w:sectPr>',
            '',
        )

        buffer = BytesIO()
        with ZipFile(BytesIO(template_bytes), "r") as source_zip, ZipFile(buffer, "w") as target_zip:
            for item in source_zip.infolist():
                data = source_zip.read(item.filename)
                if item.filename == "word/document.xml":
                    data = mutated_document_xml.encode("utf-8")
                target_zip.writestr(item, data)

        restored = ReportExportService._restore_template_header_footer(
            buffer.getvalue(),
            ReportExportService.WAIS3_TEMPLATE_PATH,
        )

        with ZipFile(BytesIO(restored), "r") as archive:
            restored_document_xml = archive.read("word/document.xml").decode("utf-8")

        self.assertIn("sectPr", restored_document_xml)
        self.assertIn("headerReference", restored_document_xml)
        self.assertIn("footerReference", restored_document_xml)

    def test_rebuild_qualitative_section_preserves_template_chart_for_wais3(self):
        template = Document(str(ReportExportService.WAIS3_TEMPLATE_PATH))
        template_chart_blocks = ReportExportService._extract_template_chart_blocks(template)

        document = Document()
        start = document.add_paragraph("ANÁLISE QUALITATIVA")
        start._p.addnext(deepcopy(template_chart_blocks[0]))
        end = document.add_paragraph("Conclusão")
        start._p.getnext().addnext(end._p)

        ReportExportService._rebuild_qualitative_section(
            document,
            sections={},
            context={
                "patient": {"sex": "F"},
                "validated_tests": [
                    {
                        "instrument_code": "wais3",
                        "structured_results": {
                            "indices": {
                                "qi_total": {
                                    "pontuacao_composta": 95,
                                    "classificacao": "Média",
                                }
                            }
                        },
                    }
                ],
            },
        )

        self.assertTrue(ReportExportService._extract_template_chart_blocks(document))

    def test_rebuild_qualitative_section_uses_wais3_template_chart_when_document_has_no_chart_block(self):
        document = Document()
        start = document.add_paragraph("ANÁLISE QUALITATIVA")
        end = document.add_paragraph("Conclusão")
        start._p.addnext(end._p)

        ReportExportService._rebuild_qualitative_section(
            document,
            sections={},
            context={
                "patient": {"sex": "F"},
                "validated_tests": [
                    {
                        "instrument_code": "wais3",
                        "structured_results": {
                            "indices": {
                                "qi_total": {
                                    "pontuacao_composta": 95,
                                    "classificacao": "Média",
                                }
                            }
                        },
                    }
                ],
            },
        )

        self.assertTrue(ReportExportService._extract_template_chart_blocks(document))

    def test_rebuild_qualitative_section_uses_wisc_template_chart_for_scared_in_wais3_report(self):
        document = Document()
        start = document.add_paragraph("ANÁLISE QUALITATIVA")
        end = document.add_paragraph("Conclusão")
        start._p.addnext(end._p)

        ReportExportService._rebuild_qualitative_section(
            document,
            sections={},
            context={
                "patient": {"sex": "F"},
                "validated_tests": [
                    {
                        "instrument_code": "wais3",
                        "structured_results": {
                            "indices": {
                                "qi_total": {"pontuacao_composta": 95, "classificacao": "Média"}
                            }
                        },
                    },
                    {
                        "instrument_code": "scared",
                        "respondent": "self",
                        "structured_results": {
                            "results": {
                                "panic_somatic": {"raw_score": 1, "classification": "Sem indicativo"},
                                "generalized_anxiety": {"raw_score": 2, "classification": "Sem indicativo"},
                                "separation_anxiety": {"raw_score": 3, "classification": "Sem indicativo"},
                                "social_phobia": {"raw_score": 4, "classification": "Sem indicativo"},
                                "school_avoidance": {"raw_score": 5, "classification": "Sem indicativo"},
                                "total": {"raw_score": 15, "classification": "Sem indicativo"},
                            }
                        },
                    },
                ],
            },
        )

        self.assertGreaterEqual(len(ReportExportService._extract_template_chart_blocks(document)), 2)

    def test_rebuild_qualitative_section_uses_wais3_model_titles_and_order(self):
        template = Document(str(ReportExportService.WAIS3_TEMPLATE_PATH))
        template_chart_blocks = ReportExportService._extract_template_chart_blocks(template)

        document = Document()
        start = document.add_paragraph("ANÁLISE QUALITATIVA")
        start._p.addnext(deepcopy(template_chart_blocks[0]))
        end = document.add_paragraph("Conclusão")
        start._p.getnext().addnext(end._p)

        ReportExportService._rebuild_qualitative_section(
            document,
            sections={
                "eficiencia_intelectual": "Interpretação global do WAIS3.",
                "linguagem": "Interpretação da linguagem.",
            },
            context={
                "patient": {"sex": "M"},
                "validated_tests": [
                    {
                        "instrument_code": "wais3",
                        "structured_results": {
                            "indices": {
                                "qi_total": {
                                    "pontuacao_composta": 95,
                                    "classificacao": "Média",
                                },
                                "qi_verbal": {
                                    "pontuacao_composta": 93,
                                    "classificacao": "Média",
                                },
                                "qi_execucao": {
                                    "pontuacao_composta": 91,
                                    "classificacao": "Média",
                                },
                                "compreensao_verbal": {
                                    "pontuacao_composta": 94,
                                    "classificacao": "Média",
                                },
                                "organizacao_perceptual": {
                                    "pontuacao_composta": 92,
                                    "classificacao": "Média",
                                },
                                "memoria_operacional": {
                                    "pontuacao_composta": 96,
                                    "classificacao": "Média",
                                },
                                "velocidade_processamento": {
                                    "pontuacao_composta": 90,
                                    "classificacao": "Média",
                                },
                            }
                        },
                        "wais3_tables": {
                            "linguagem": [
                                {
                                    "label": "Semelhanças",
                                    "maxScore": "38",
                                    "avgScore": "17-28",
                                    "minScore": "9-10",
                                    "obtainedScore": "18",
                                    "classification": "Média",
                                }
                            ]
                        },
                    }
                ],
            },
        )

        texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

        self.assertIn("Desempenho do paciente no WAIS III", texts)
        self.assertIn("Subescalas WAIS III", texts)
        self.assertIn("Gráfico 1 WAIS III - INDICES DE QIS", texts)
        self.assertNotIn("5.1. Desempenho do paciente no WAIS-III", texts)

        desempenho_index = texts.index("Desempenho do paciente no WAIS III")
        interpretacao_index = next(
            i for i, text in enumerate(texts) if "Interpretação global do WAIS3." in text
        )
        subescalas_index = texts.index("Subescalas WAIS III")

        self.assertLess(desempenho_index, interpretacao_index)
        self.assertLess(interpretacao_index, subescalas_index)

    def test_rebuild_qualitative_section_includes_bfp_only_when_applied(self):
        document = Document(str(ReportExportService.WAIS3_TEMPLATE_PATH))

        ReportExportService._rebuild_qualitative_section(
            document,
            sections={},
            context={
                "patient": {"sex": "M"},
                "validated_tests": [
                    {
                        "instrument_code": "wais3",
                        "structured_results": {
                            "indices": {
                                "qi_total": {"pontuacao_composta": 95, "classificacao": "Média"},
                            }
                        },
                    },
                    {
                        "instrument_code": "bfp",
                        "computed_payload": {
                            "factor_order": ["NN", "EE", "SS", "RR", "AA"],
                            "factors": {
                                "NN": {"name": "Neuroticismo", "raw_score": 4.2, "percentile": 88, "classification": "Superior"},
                                "EE": {"name": "Extroversão", "raw_score": 3.3, "percentile": 22, "classification": "Média Inferior"},
                                "SS": {"name": "Socialização", "raw_score": 4.8, "percentile": 50, "classification": "Média"},
                                "RR": {"name": "Realização", "raw_score": 3.0, "percentile": 18, "classification": "Média Inferior"},
                                "AA": {"name": "Abertura", "raw_score": 4.1, "percentile": 60, "classification": "Média"},
                            },
                        },
                        "clinical_interpretation": "Texto interpretativo do BFP.",
                    },
                ],
            },
        )

        texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        self.assertIn("BFP – Bateria Fatorial de Personalidade", texts)
        self.assertTrue(any(text.endswith("BFP Resultados dos fatores") for text in texts))
        self.assertTrue(any("Texto interpretativo do BFP." in text for text in texts))

        document_without_bfp = Document(str(ReportExportService.WAIS3_TEMPLATE_PATH))
        ReportExportService._rebuild_qualitative_section(
            document_without_bfp,
            sections={},
            context={
                "patient": {"sex": "M"},
                "validated_tests": [
                    {
                        "instrument_code": "wais3",
                        "structured_results": {
                            "indices": {
                                "qi_total": {"pontuacao_composta": 95, "classificacao": "Média"},
                            }
                        },
                    }
                ],
            },
        )

        texts_without_bfp = [p.text.strip() for p in document_without_bfp.paragraphs if p.text.strip()]
        self.assertNotIn("BFP – Bateria Fatorial de Personalidade", texts_without_bfp)

    def test_rebuild_qualitative_section_places_bfp_between_ebadep_and_srs2_in_wais3_model(self):
        document = Document(str(ReportExportService.WAIS3_TEMPLATE_PATH))

        ReportExportService._rebuild_qualitative_section(
            document,
            sections={
                "aspectos_emocionais_comportamentais": "Interpretacao integrada."
            },
            context={
                "patient": {"sex": "M"},
                "validated_tests": [
                    {
                        "instrument_code": "wais3",
                        "structured_results": {
                            "indices": {
                                "qi_total": {"pontuacao_composta": 95, "classificacao": "Média"},
                            }
                        },
                    },
                    {
                        "instrument_code": "ebadep_a",
                        "classified_payload": {"resultado": {"percentil": 43, "classificacao": "Mínimo"}},
                    },
                    {
                        "instrument_code": "bfp",
                        "computed_payload": {
                            "factor_order": ["NN", "EE", "SS", "RR", "AA"],
                            "factors": {
                                "NN": {"name": "Neuroticismo", "raw_score": 4.2, "percentile": 88, "classification": "Superior"},
                                "EE": {"name": "Extroversão", "raw_score": 3.3, "percentile": 22, "classification": "Média Inferior"},
                                "SS": {"name": "Socialização", "raw_score": 4.8, "percentile": 50, "classification": "Média"},
                                "RR": {"name": "Realização", "raw_score": 3.0, "percentile": 18, "classification": "Média Inferior"},
                                "AA": {"name": "Abertura", "raw_score": 4.1, "percentile": 60, "classification": "Média"},
                            },
                        },
                        "clinical_interpretation": "Texto interpretativo do BFP.",
                    },
                    {
                        "instrument_code": "srs2",
                        "classified_payload": {"resultados": []},
                    },
                ],
            },
        )

        texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        ebadep_index = texts.index("EBADEP-A")
        bfp_index = texts.index("BFP – Bateria Fatorial de Personalidade")
        srs2_index = texts.index("SRS-2 Escala de Responsividade Social")

        self.assertLess(ebadep_index, bfp_index)
        self.assertLess(bfp_index, srs2_index)

    def test_rebuild_qualitative_section_uses_wais3_skill_labels_for_late_chapters(self):
        template = Document(str(ReportExportService.WAIS3_TEMPLATE_PATH))
        template_chart_blocks = ReportExportService._extract_template_chart_blocks(template)

        document = Document()
        start = document.add_paragraph("ANÁLISE QUALITATIVA")
        start._p.addnext(deepcopy(template_chart_blocks[0]))
        end = document.add_paragraph("Conclusão")
        start._p.getnext().addnext(end._p)

        ReportExportService._rebuild_qualitative_section(
            document,
            sections={
                "etdah_ad": "Interpretacao do ETDAH-AD.",
                "srs2": "Interpretacao do SRS-2.",
                "aspectos_emocionais_comportamentais": "Interpretacao da EBADEP-A.",
            },
            context={
                "patient": {"sex": "F"},
                "validated_tests": [
                    {"instrument_code": "wais3", "structured_results": {"indices": {"qi_total": {"pontuacao_composta": 95, "classificacao": "Média"}}}},
                    {"instrument_code": "etdah_ad"},
                    {"instrument_code": "ebadep_a"},
                    {"instrument_code": "srs2"},
                ],
            },
        )

        texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

        self.assertIn("Função Executiva", texts)
        self.assertTrue(any(text.startswith("A E-TDAH-AD e um instrumento de autorrelato destinado a adolescentes e adultos") for text in texts))
        self.assertTrue(any(text.startswith("A EBADEP-A avalia a presenca e a intensidade de sintomas depressivos em adultos") for text in texts))
        self.assertTrue(any(text.startswith("A SRS-2 e uma escala destinada a investigacao de dificuldades") for text in texts))

    def test_rebuild_qualitative_section_for_plain_document_hides_epq_and_includes_bai_bfp(self):
        document = Document()
        start = document.add_paragraph("ANÁLISE QUALITATIVA")
        end = document.add_paragraph("Conclusão")
        start._p.addnext(end._p)

        ReportExportService._rebuild_qualitative_section(
            document,
            sections={"aspectos_emocionais_comportamentais": "Texto integrado."},
            context={
                "patient": {"sex": "F"},
                "validated_tests": [
                    {
                        "instrument_code": "bai",
                        "computed_payload": {"total_raw_score": 24, "t_score": 58, "percentile": 78.4},
                        "classified_payload": {
                            "faixa_normativa": "Moderado",
                            "interpretacao_faixa": "nível moderado de ansiedade",
                            "classificacao": {"label": "Moderado"},
                        },
                    },
                    {
                        "instrument_code": "ebadep_a",
                        "classified_payload": {
                            "result": {
                                "escore_total": 69,
                                "percentil": 81,
                                "classificacao": "Sintomatologia Depressiva Leve",
                            }
                        },
                    },
                    {
                        "instrument_code": "bfp",
                        "computed_payload": {
                            "factor_order": ["NN"],
                            "factors": {
                                "NN": {
                                    "name": "Neuroticismo",
                                    "raw_score": 4.2,
                                    "percentile": 88,
                                    "classification": "Superior",
                                }
                            },
                        },
                        "clinical_interpretation": "Texto interpretativo do BFP.",
                    },
                ],
            },
        )

        texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

        self.assertNotIn("EPQ-J", texts)
        self.assertIn("BAI", texts)
        self.assertIn("EBADEP-A", texts)
        self.assertIn("BFP – Bateria Fatorial de Personalidade", texts)
        self.assertTrue(any(text.endswith("BAI - Resultado da sintomatologia ansiosa") for text in texts))
        self.assertTrue(any(text.endswith("BFP Resultados dos fatores") for text in texts))

        drawings = document._body._element.findall(
            './/w:drawing',
            {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        )
        self.assertGreaterEqual(len(drawings), 2)

    def test_rebuild_qualitative_section_does_not_duplicate_wais3_global_intro(self):
        document = Document()
        start = document.add_paragraph("ANÁLISE QUALITATIVA")
        end = document.add_paragraph("Conclusão")
        start._p.addnext(end._p)

        ReportExportService._rebuild_qualitative_section(
            document,
            sections={},
            context={
                "patient": {"sex": "M"},
                "validated_tests": [
                    {
                        "instrument_code": "wais3",
                        "structured_results": {
                            "indices": {
                                "qi_total": {
                                    "pontuacao_composta": 98,
                                    "classificacao": "Média",
                                },
                                "qi_verbal": {
                                    "pontuacao_composta": 99,
                                    "classificacao": "Média",
                                },
                                "qi_execucao": {
                                    "pontuacao_composta": 99,
                                    "classificacao": "Média",
                                },
                                "compreensao_verbal": {
                                    "pontuacao_composta": 102,
                                    "classificacao": "Média",
                                },
                                "organizacao_perceptual": {
                                    "pontuacao_composta": 96,
                                    "classificacao": "Média",
                                },
                                "memoria_operacional": {
                                    "pontuacao_composta": 90,
                                    "classificacao": "Média",
                                },
                                "velocidade_processamento": {
                                    "pontuacao_composta": 105,
                                    "classificacao": "Média",
                                },
                            }
                        },
                    }
                ],
            },
        )

        global_lines = [
            p.text.strip()
            for p in document.paragraphs
            if p.text.strip().startswith("Capacidade Cognitiva Global:")
        ]

        self.assertEqual(len(global_lines), 1)
        self.assertEqual(global_lines[0].count("Capacidade Cognitiva Global:"), 1)

    def test_rebuild_qualitative_section_uses_wisc4_model_headings(self):
        document = Document()
        start = document.add_paragraph("ANÁLISE QUALITATIVA")
        end = document.add_paragraph("Conclusão")
        start._p.addnext(end._p)

        ReportExportService._rebuild_qualitative_section(
            document,
            sections={},
            context={
                "patient": {"full_name": "Maria Clara", "sex": "F", "birth_date": "2016-01-10"},
                "validated_tests": [
                    {
                        "instrument_code": "wisc4",
                        "structured_results": {
                            "qit_data": {"escore_composto": 61, "classificacao": "Extremamente Baixo"},
                            "indices": [
                                {"indice": "icv", "escore_composto": 84, "classificacao": "Média Inferior"},
                                {"indice": "iop", "escore_composto": 61, "classificacao": "Extremamente Baixo"},
                                {"indice": "imt", "escore_composto": 52, "classificacao": "Extremamente Baixo"},
                                {"indice": "ivp", "escore_composto": 68, "classificacao": "Extremamente Baixo"},
                            ],
                            "subtestes": [
                                {"codigo": "SM", "subteste": "Semelhanças", "escore_bruto": 23, "classificacao": "Média"},
                                {"codigo": "CN", "subteste": "Conceitos Figurativos", "escore_bruto": 11, "classificacao": "Média Inferior"},
                                {"codigo": "CO", "subteste": "Compreensão", "escore_bruto": 24, "classificacao": "Média"},
                                {"codigo": "RM", "subteste": "Raciocínio Matricial", "escore_bruto": 12, "classificacao": "Limítrofe"},
                                {"codigo": "VC", "subteste": "Vocabulário", "escore_bruto": 10, "classificacao": "Limítrofe"},
                                {"codigo": "CB", "subteste": "Cubos", "escore_bruto": 8, "classificacao": "Extremamente Baixo"},
                                {"codigo": "SNL", "subteste": "Sequência de Números e Letras", "escore_bruto": 5, "classificacao": "Extremamente Baixo"},
                                {"codigo": "DG", "subteste": "Dígitos", "escore_bruto": 7, "classificacao": "Limítrofe"},
                            ],
                        },
                        "classified_payload": {
                            "subtestes": [
                                {"codigo": "SM", "subteste": "Semelhanças", "escore_bruto": 23, "classificacao": "Média"},
                                {"codigo": "CN", "subteste": "Conceitos Figurativos", "escore_bruto": 11, "classificacao": "Média Inferior"},
                                {"codigo": "CO", "subteste": "Compreensão", "escore_bruto": 24, "classificacao": "Média"},
                                {"codigo": "RM", "subteste": "Raciocínio Matricial", "escore_bruto": 12, "classificacao": "Limítrofe"},
                                {"codigo": "VC", "subteste": "Vocabulário", "escore_bruto": 10, "classificacao": "Limítrofe"},
                                {"codigo": "CB", "subteste": "Cubos", "escore_bruto": 8, "classificacao": "Extremamente Baixo"},
                                {"codigo": "SNL", "subteste": "Sequência de Números e Letras", "escore_bruto": 5, "classificacao": "Extremamente Baixo"},
                                {"codigo": "DG", "subteste": "Dígitos", "escore_bruto": 7, "classificacao": "Limítrofe"},
                            ]
                        },
                        "applied_on": "2024-04-20",
                    }
                ],
                "evaluation": {"start_date": "2024-04-20"},
            },
        )

        texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

        self.assertIn("Capacidade Cognitiva Global", texts)
        self.assertIn("Desempenho da paciente no WISC-IV", texts)
        self.assertIn("Subescalas WISC-IV", texts)
        self.assertIn("Função Executiva", texts)
        self.assertTrue(any(text.startswith("Interpretação e Observações Clínicas:") for text in texts))

    def test_build_adolescent_document_uses_model_wisc4_titles_when_optional_tests_are_missing(self):
        document = ReportExportService._build_adolescent_document(
            self._report_stub(),
            sections={"conclusao": "Conclusão clínica.", "sugestoes_conduta": "- Sugestão 1"},
            context={
                "patient": {"full_name": "Maria Clara", "sex": "F", "birth_date": "2016-01-10"},
                "evaluation": {"start_date": "2024-04-20"},
                "validated_tests": [
                    {
                        "instrument_code": "wisc4",
                        "structured_results": {
                            "qit_data": {"escore_composto": 61, "classificacao": "Extremamente Baixo"},
                            "indices": [
                                {"indice": "icv", "escore_composto": 84, "classificacao": "Média Inferior"},
                                {"indice": "iop", "escore_composto": 61, "classificacao": "Extremamente Baixo"},
                                {"indice": "imt", "escore_composto": 52, "classificacao": "Extremamente Baixo"},
                                {"indice": "ivp", "escore_composto": 68, "classificacao": "Extremamente Baixo"},
                            ],
                            "subtestes": [
                                {"codigo": "SM", "subteste": "Semelhanças", "escore_bruto": 23, "classificacao": "Média"},
                                {"codigo": "CN", "subteste": "Conceitos Figurativos", "escore_bruto": 11, "classificacao": "Média Inferior"},
                                {"codigo": "CO", "subteste": "Compreensão", "escore_bruto": 24, "classificacao": "Média"},
                                {"codigo": "RM", "subteste": "Raciocínio Matricial", "escore_bruto": 12, "classificacao": "Limítrofe"},
                                {"codigo": "VC", "subteste": "Vocabulário", "escore_bruto": 10, "classificacao": "Limítrofe"},
                                {"codigo": "CB", "subteste": "Cubos", "escore_bruto": 8, "classificacao": "Extremamente Baixo"},
                                {"codigo": "SNL", "subteste": "Sequência de Números e Letras", "escore_bruto": 5, "classificacao": "Extremamente Baixo"},
                                {"codigo": "DG", "subteste": "Dígitos", "escore_bruto": 7, "classificacao": "Limítrofe"},
                            ],
                        },
                        "classified_payload": {
                            "subtestes": [
                                {"codigo": "SM", "subteste": "Semelhanças", "escore_bruto": 23, "classificacao": "Média"}
                            ]
                        },
                        "applied_on": "2024-04-20",
                    }
                ],
            },
        )

        texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

        self.assertIn("5. ANÁLISE QUALITATIVA", texts)
        self.assertIn("Capacidade Cognitiva Global", texts)
        self.assertIn("Desempenho da paciente no WISC-IV", texts)
        self.assertIn("Subescalas WISC-IV", texts)
        self.assertIn("Conclusão", texts)
        self.assertIn("Sugestões de Conduta (Encaminhamentos)", texts)
        self.assertIn("Referencia Bibliográfica", texts)
        self.assertNotIn("6. BPA-2 – BATERIA PSICOLÓGICA PARA AVALIAÇÃO DA ATENÇÃO", texts)
        self.assertNotIn("6. CONCLUSÃO", texts)
        self.assertNotIn("7. SUGESTÕES DE CONDUTA (ENCAMINHAMENTOS)", texts)
        self.assertNotIn("8. REFERÊNCIA BIBLIOGRÁFICA", texts)

    def test_replace_simple_sections_uses_wais3_skill_structure_for_demand_and_procedures(self):
        document = Document()
        document.add_paragraph("IDENTIFICAÇÃO")
        document.add_paragraph("DESCRIÇÃO DA DEMANDA")
        document.add_paragraph("PROCEDIMENTOS")
        document.add_paragraph("ANÁLISE")

        ReportExportService._replace_simple_sections(
            document,
            self._report_stub(),
            sections={
                "identificacao": "Identificação pronta.",
                "descricao_demanda": "O paciente Marcos foi encaminhado para avaliação neuropsicológica por apresentar queixas relacionadas a desatenção e esquecimentos frequentes.",
                "procedimentos": "Para esta avaliação foram realizadas: uma sessão de anamnese, 05 sessões de testagem com o paciente e uma sessão de devolutiva.",
                "historia_pessoal": "História pessoal resumida.",
            },
            context={
                "patient": {"full_name": "Marcos Henrique Carvalho Santos"},
                "validated_tests": [
                    {"instrument_code": "wais3"},
                    {"instrument_code": "bpa2"},
                ],
            },
        )

        texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

        self.assertIn("PROCEDIMENTOS", texts)

    def test_replace_simple_sections_uses_default_wais3_demand_text_when_empty(self):
        document = Document()
        document.add_paragraph("IDENTIFICAÇÃO")
        document.add_paragraph("DESCRIÇÃO DA DEMANDA")
        document.add_paragraph("PROCEDIMENTOS")
        document.add_paragraph("ANÁLISE")

        ReportExportService._replace_simple_sections(
            document,
            self._report_stub(),
            sections={},
            context={"patient": {"full_name": "Marcos Henrique Carvalho Santos"}, "validated_tests": []},
        )

        texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

        self.assertIn("Motivo do Encaminhamento", texts)
        self.assertIn(
            "Levando em consideração a demanda apresentada, objetivou-se avaliar as funções cognitivas, aspectos comportamentais, emocionais e sociais do paciente, a fim de compreender seu funcionamento global e subsidiar orientações para o manejo adequado em contextos familiar e escolar.",
            texts,
        )

    def test_replace_simple_sections_rebuilds_identification_chapter_with_fixed_labels(self):
        document = Document()
        document.add_paragraph("IDENTIFICAÇÃO")
        document.add_paragraph("DESCRIÇÃO DA DEMANDA")

        report = self._report_stub()
        report.interested_party = "Dra. Exemplo"
        report.purpose = "Auxílio diagnóstico"

        ReportExportService._replace_simple_sections(
            document,
            report,
            sections={},
            context={
                "patient": {
                    "full_name": "Marcos Henrique Carvalho Santos",
                    "sex": "M",
                    "birth_date": "2002-09-12",
                    "schooling": "superior incompleto",
                }
            },
        )

        texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

        self.assertEqual(texts[1], "1.1. Identificação do laudo:")
        self.assertEqual(texts[2], "Autora: Jacqueline Oliveira Caires (CRP 09/6017)")
        self.assertEqual(texts[3], "Interessado: Dra. Exemplo")
        self.assertEqual(texts[4], "Finalidade: Auxílio diagnóstico")
        self.assertEqual(texts[5], "1.2. Identificação do paciente:")
        self.assertEqual(texts[6], "Nome: Marcos Henrique Carvalho Santos")
        self.assertEqual(texts[7], "Sexo: Masculino")
        self.assertEqual(texts[8], "Data de nascimento: 12/09/2002")
        self.assertTrue(texts[9].startswith("Idade: "))
        self.assertIn("Filiação: Não informada", texts)
        self.assertIn("Escolaridade: superior incompleto", texts)

    def test_sanitize_chart_xml_inlines_cached_refs_and_removes_external_data(self):
        chart_xml = b'''<?xml version="1.0" encoding="UTF-8"?>
<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <c:externalData r:id="rId1"/>
  <c:chart>
    <c:plotArea>
      <c:barChart>
        <c:ser>
          <c:idx val="0"/>
          <c:order val="0"/>
          <c:tx>
            <c:strRef>
              <c:f>Plan1!$B$1</c:f>
              <c:strCache><c:ptCount val="1"/><c:pt idx="0"><c:v>Serie A</c:v></c:pt></c:strCache>
            </c:strRef>
          </c:tx>
          <c:cat>
            <c:strRef>
              <c:f>Plan1!$A$2:$A$3</c:f>
              <c:strCache><c:ptCount val="2"/><c:pt idx="0"><c:v>A</c:v></c:pt><c:pt idx="1"><c:v>B</c:v></c:pt></c:strCache>
            </c:strRef>
          </c:cat>
          <c:val>
            <c:numRef>
              <c:f>Plan1!$B$2:$B$3</c:f>
              <c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="2"/><c:pt idx="0"><c:v>10</c:v></c:pt><c:pt idx="1"><c:v>20</c:v></c:pt></c:numCache>
            </c:numRef>
          </c:val>
        </c:ser>
      </c:barChart>
    </c:plotArea>
  </c:chart>
</c:chartSpace>'''

        sanitized = ReportExportService._sanitize_chart_xml_bytes(chart_xml)
        root = ET.fromstring(sanitized)
        ns = ReportExportService.CHART_NS

        self.assertIsNone(root.find('.//c:externalData', ns))
        self.assertIsNone(root.find('.//c:strRef', ns))
        self.assertIsNone(root.find('.//c:numRef', ns))
        self.assertEqual(root.find('.//c:tx/c:v', ns).text, 'Serie A')
        self.assertEqual(
            [node.text for node in root.findall('.//c:cat/c:strLit/c:pt/c:v', ns)],
            ['A', 'B'],
        )
        self.assertEqual(
            [node.text for node in root.findall('.//c:val/c:numLit/c:pt/c:v', ns)],
            ['10', '20'],
        )

    def test_strip_external_chart_relationships_keeps_internal_links(self):
        rels_xml = b'''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.microsoft.com/office/2011/relationships/chartStyle" Target="style1.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" Target="https://example.com/test.xlsx" TargetMode="External"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chartUserShapes" Target="../drawings/drawing1.xml"/>
</Relationships>'''

        sanitized = ReportExportService._strip_external_chart_relationships(rels_xml)
        text = sanitized.decode('utf-8')

        self.assertIn('style1.xml', text)
        self.assertIn('../drawings/drawing1.xml', text)
        self.assertNotIn('example.com/test.xlsx', text)
        self.assertNotIn('oleObject', text)

    def test_sanitize_chart_xml_flattens_multi_level_categories(self):
        chart_xml = b'''<?xml version="1.0" encoding="UTF-8"?>
<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart">
  <c:chart>
    <c:plotArea>
      <c:barChart>
        <c:ser>
          <c:idx val="0"/>
          <c:order val="0"/>
          <c:cat>
            <c:multiLvlStrRef>
              <c:f>Plan1!$A$1:$B$3</c:f>
              <c:multiLvlStrCache>
                <c:ptCount val="3"/>
                <c:lvl>
                  <c:pt idx="0"><c:v>A</c:v></c:pt>
                  <c:pt idx="1"><c:v>B</c:v></c:pt>
                </c:lvl>
                <c:lvl>
                  <c:pt idx="2"><c:v>C</c:v></c:pt>
                </c:lvl>
              </c:multiLvlStrCache>
            </c:multiLvlStrRef>
          </c:cat>
          <c:val>
            <c:numRef>
              <c:f>Plan1!$C$1:$C$3</c:f>
              <c:numCache><c:ptCount val="3"/><c:pt idx="0"><c:v>1</c:v></c:pt><c:pt idx="1"><c:v>2</c:v></c:pt><c:pt idx="2"><c:v>3</c:v></c:pt></c:numCache>
            </c:numRef>
          </c:val>
        </c:ser>
      </c:barChart>
    </c:plotArea>
  </c:chart>
</c:chartSpace>'''

        sanitized = ReportExportService._sanitize_chart_xml_bytes(chart_xml)
        root = ET.fromstring(sanitized)
        ns = ReportExportService.CHART_NS

        self.assertIsNone(root.find('.//c:multiLvlStrRef', ns))
        self.assertEqual(
            [node.text for node in root.findall('.//c:cat/c:strLit/c:pt/c:v', ns)],
            ['A', 'B', 'C'],
        )

    def test_sanitize_chart_xml_preserves_mc_requires_namespace_prefix(self):
        chart_xml = b'''<?xml version="1.0" encoding="UTF-8"?>
<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:c16r2="http://schemas.microsoft.com/office/drawing/2015/06/chart">
  <mc:AlternateContent xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
    <mc:Choice Requires="c14" xmlns:c14="http://schemas.microsoft.com/office/drawing/2007/8/2/chart">
      <c14:style val="102"/>
    </mc:Choice>
    <mc:Fallback><c:style val="2"/></mc:Fallback>
  </mc:AlternateContent>
  <c:chart><c:plotArea/></c:chart>
</c:chartSpace>'''

        sanitized = ReportExportService._sanitize_chart_xml_bytes(chart_xml).decode('utf-8')

        self.assertIn('Requires="c14"', sanitized)
        self.assertIn('xmlns:c14="http://schemas.microsoft.com/office/drawing/2007/8/2/chart"', sanitized)
        self.assertIn('<c14:style', sanitized)

    def test_update_chart_series_before_sanitize_replaces_template_values(self):
        chart_xml = LET.fromstring(b'''<?xml version="1.0" encoding="UTF-8"?>
<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart">
  <c:chart>
    <c:plotArea>
      <c:barChart>
        <c:ser>
          <c:idx val="0"/>
          <c:order val="0"/>
          <c:cat>
            <c:strRef>
              <c:f>Plan1!$A$1:$A$2</c:f>
              <c:strCache>
                <c:ptCount val="2"/>
                <c:pt idx="0"><c:v>Old A</c:v></c:pt>
                <c:pt idx="1"><c:v>Old B</c:v></c:pt>
              </c:strCache>
            </c:strRef>
          </c:cat>
          <c:val>
            <c:numRef>
              <c:f>Plan1!$B$1:$B$2</c:f>
              <c:numCache>
                <c:formatCode>General</c:formatCode>
                <c:ptCount val="2"/>
                <c:pt idx="0"><c:v>84</c:v></c:pt>
                <c:pt idx="1"><c:v>61</c:v></c:pt>
              </c:numCache>
            </c:numRef>
          </c:val>
        </c:ser>
      </c:barChart>
    </c:plotArea>
  </c:chart>
</c:chartSpace>''')

        ReportExportService._update_chart_series(chart_xml, 0, ['ICV', 'IOP'], [113, 90])
        ReportExportService._sanitize_chart_root(chart_xml)

        ns = ReportExportService.CHART_NS
        self.assertEqual(
            [node.text for node in chart_xml.findall('.//c:cat//c:pt/c:v', ns)],
            ['ICV', 'IOP'],
        )
        self.assertEqual(
            [node.text for node in chart_xml.findall('.//c:val//c:pt/c:v', ns)],
            ['113', '90'],
        )

    def test_docx_package_is_valid_detects_missing_relationship_target(self):
        buffer = BytesIO()
        with ZipFile(buffer, 'w') as docx:
            docx.writestr(
                '[Content_Types].xml',
                '''<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/_rels/document.xml.rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
</Types>''',
            )
            docx.writestr(
                'word/document.xml',
                '''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><w:body>
<w:p><w:r><w:drawing><c:chart r:id="rId1"/></w:drawing></w:r></w:p>
</w:body></w:document>''',
            )
            docx.writestr(
                'word/_rels/document.xml.rels',
                '''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart1.xml"/>
</Relationships>''',
            )

        self.assertFalse(ReportExportService._docx_package_is_valid(buffer.getvalue()))

    def test_wais3_chart_payload_matches_model_labels(self):
        categories, values = ReportExportService._wais3_chart_payload(
            {
                "structured_results": {
                    "indices": {
                        "compreensao_verbal": {"pontuacao_composta": 104},
                        "organizacao_perceptual": {"pontuacao_composta": 96},
                        "memoria_operacional": {"pontuacao_composta": 100},
                        "velocidade_processamento": {"pontuacao_composta": 98},
                        "qi_verbal": {"pontuacao_composta": 104},
                        "qi_execucao": {"pontuacao_composta": 100},
                        "qi_total": {"pontuacao_composta": 102},
                        "gai": {"pontuacao_composta": 101},
                    }
                }
            }
        )

        self.assertEqual(categories, ["ICV", "IOP", "IMO", "IVP", "QIV", "QIE", "QIT"])
        self.assertEqual(values, [104.0, 96.0, 100.0, 98.0, 104.0, 100.0, 102.0])

    def test_populate_wais3_excel_charts_keeps_scared_and_srs2_in_separate_slots(self):
        def chart_xml():
            return b'''<?xml version="1.0" encoding="UTF-8"?>
<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart">
  <c:chart>
    <c:plotArea>
      <c:barChart>
        <c:ser>
          <c:idx val="0"/>
          <c:order val="0"/>
          <c:cat><c:strRef><c:f>Plan1!$A$1:$A$2</c:f><c:strCache><c:ptCount val="2"/><c:pt idx="0"><c:v>Old A</c:v></c:pt><c:pt idx="1"><c:v>Old B</c:v></c:pt></c:strCache></c:strRef></c:cat>
          <c:val><c:numRef><c:f>Plan1!$B$1:$B$2</c:f><c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="2"/><c:pt idx="0"><c:v>1</c:v></c:pt><c:pt idx="1"><c:v>2</c:v></c:pt></c:numCache></c:numRef></c:val>
        </c:ser>
        <c:ser>
          <c:idx val="1"/>
          <c:order val="1"/>
          <c:cat><c:strRef><c:f>Plan1!$A$1:$A$2</c:f><c:strCache><c:ptCount val="2"/><c:pt idx="0"><c:v>Old A</c:v></c:pt><c:pt idx="1"><c:v>Old B</c:v></c:pt></c:strCache></c:strRef></c:cat>
          <c:val><c:numRef><c:f>Plan1!$C$1:$C$2</c:f><c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="2"/><c:pt idx="0"><c:v>3</c:v></c:pt><c:pt idx="1"><c:v>4</c:v></c:pt></c:numCache></c:numRef></c:val>
        </c:ser>
        <c:ser>
          <c:idx val="2"/>
          <c:order val="2"/>
          <c:cat><c:strRef><c:f>Plan1!$A$1:$A$2</c:f><c:strCache><c:ptCount val="2"/><c:pt idx="0"><c:v>Old A</c:v></c:pt><c:pt idx="1"><c:v>Old B</c:v></c:pt></c:strCache></c:strRef></c:cat>
          <c:val><c:numRef><c:f>Plan1!$D$1:$D$2</c:f><c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="2"/><c:pt idx="0"><c:v>5</c:v></c:pt><c:pt idx="1"><c:v>6</c:v></c:pt></c:numCache></c:numRef></c:val>
        </c:ser>
        <c:ser>
          <c:idx val="3"/>
          <c:order val="3"/>
          <c:cat><c:strRef><c:f>Plan1!$A$1:$A$2</c:f><c:strCache><c:ptCount val="2"/><c:pt idx="0"><c:v>Old A</c:v></c:pt><c:pt idx="1"><c:v>Old B</c:v></c:pt></c:strCache></c:strRef></c:cat>
          <c:val><c:numRef><c:f>Plan1!$E$1:$E$2</c:f><c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="2"/><c:pt idx="0"><c:v>7</c:v></c:pt><c:pt idx="1"><c:v>8</c:v></c:pt></c:numCache></c:numRef></c:val>
        </c:ser>
      </c:barChart>
    </c:plotArea>
  </c:chart>
</c:chartSpace>'''

        buffer = BytesIO()
        with ZipFile(buffer, 'w') as docx:
            docx.writestr(
                '[Content_Types].xml',
                '''<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/charts/chart1.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
  <Override PartName="/word/charts/chart2.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
  <Override PartName="/word/charts/chart3.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
  <Override PartName="/word/charts/chart4.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
  <Override PartName="/word/charts/chart5.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
  <Override PartName="/word/charts/chart6.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
  <Override PartName="/word/charts/chart7.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
  <Override PartName="/word/charts/chart8.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
</Types>''',
            )
            docx.writestr(
                'word/document.xml',
                '''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><w:body>
<w:p><w:r><w:drawing><c:chart r:id="rId1"/></w:drawing></w:r></w:p>
<w:p><w:r><w:drawing><c:chart r:id="rId2"/></w:drawing></w:r></w:p>
<w:p><w:r><w:drawing><c:chart r:id="rId3"/></w:drawing></w:r></w:p>
<w:p><w:r><w:drawing><c:chart r:id="rId4"/></w:drawing></w:r></w:p>
<w:p><w:r><w:drawing><c:chart r:id="rId5"/></w:drawing></w:r></w:p>
<w:p><w:r><w:drawing><c:chart r:id="rId6"/></w:drawing></w:r></w:p>
<w:p><w:r><w:drawing><c:chart r:id="rId7"/></w:drawing></w:r></w:p>
<w:p><w:r><w:drawing><c:chart r:id="rId8"/></w:drawing></w:r></w:p>
</w:body></w:document>''',
            )
            docx.writestr(
                'word/_rels/document.xml.rels',
                '''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart1.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart2.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart3.xml"/>
  <Relationship Id="rId4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart4.xml"/>
  <Relationship Id="rId5" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart5.xml"/>
  <Relationship Id="rId6" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart6.xml"/>
  <Relationship Id="rId7" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart7.xml"/>
  <Relationship Id="rId8" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart8.xml"/>
</Relationships>''',
            )
            for idx in range(1, 9):
                docx.writestr(f'word/charts/chart{idx}.xml', chart_xml())
                docx.writestr(f'word/charts/_rels/chart{idx}.xml.rels', '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>')

        docx_bytes = ReportExportService._populate_wais3_excel_charts(
            buffer.getvalue(),
            {
                'validated_tests': [
                    {'instrument_code': 'wais3', 'structured_results': {'indices': {'compreensao_verbal': {'pontuacao_composta': 104}, 'organizacao_perceptual': {'pontuacao_composta': 96}, 'memoria_operacional': {'pontuacao_composta': 100}, 'velocidade_processamento': {'pontuacao_composta': 98}, 'qi_verbal': {'pontuacao_composta': 104}, 'qi_execucao': {'pontuacao_composta': 100}, 'qi_total': {'pontuacao_composta': 102}}}},
                    {'instrument_code': 'bpa2', 'classified_payload': {'resultados': [{'codigo': 'ac', 'pontuacao': 64, 'media': 91, 'minimo': 63, 'maximo': 120, 'percentil': 20}, {'codigo': 'ad', 'pontuacao': 80, 'media': 83, 'minimo': 49, 'maximo': 119, 'percentil': 40}, {'codigo': 'aa', 'pontuacao': 86, 'media': 95, 'minimo': 61, 'maximo': 120, 'percentil': 30}, {'codigo': 'ag', 'pontuacao': 230, 'media': 267, 'minimo': 217, 'maximo': 352, 'percentil': 25}]}} ,
                    {'instrument_code': 'ravlt', 'structured_results': {'ravlt_esperado': {'a1': 7, 'a2': 9, 'a3': 11, 'a4': 12, 'a5': 12, 'b1': 6, 'a6': 12, 'a7': 11, 'r': 13, 'alt': 18, 'ret': 1, 'ip': 0.89, 'ir': 0.92}, 'ravlt_minimo': {'a1': 6, 'a2': 8, 'a3': 10, 'a4': 10, 'a5': 11, 'b1': 5, 'a6': 9, 'a7': 9, 'r': 5, 'alt': 12, 'ret': 0.91, 'ip': 0.73, 'ir': 0.82}, 'ravlt_obtido': {'a1': 4, 'a2': 5, 'a3': 7, 'a4': 7, 'a5': 7, 'b1': 4, 'a6': 5, 'a7': 5, 'r': 11, 'alt': 10, 'ret': 1, 'ip': 1, 'ir': 0.71}}},
                    {'instrument_code': 'fdt', 'classified_payload': {'metric_results': [{'codigo': 'leitura', 'media': 20.4, 'valor': 31.91, 'percentil_num': 5}, {'codigo': 'contagem', 'media': 23.8, 'valor': 32.46, 'percentil_num': 5}, {'codigo': 'escolha', 'media': 34, 'valor': 62.98, 'percentil_num': 5}, {'codigo': 'alternancia', 'media': 44.8, 'valor': 71.85, 'percentil_num': 5}, {'codigo': 'inibicao', 'media': 13.6, 'valor': 31.07, 'percentil_num': 5}, {'codigo': 'flexibilidade', 'media': 24.4, 'valor': 39.94, 'percentil_num': 5}], 'erros': {'leitura': {'qtde_erros': 0}, 'contagem': {'qtde_erros': 0}, 'escolha': {'qtde_erros': 1}, 'alternancia': {'qtde_erros': 4}, 'inibicao': {'qtde_erros': 0}, 'flexibilidade': {'qtde_erros': 0}}}},
                    {'instrument_code': 'etdah_ad', 'classified_payload': {'results': {'D': {'percentil': '30,48'}, 'I': {'percentil': '40,14'}, 'AE': {'percentil': '78,25'}, 'AAMA': {'percentil': '1,99'}, 'H': {'percentil': '27,21'}}}},
                    {'instrument_code': 'scared', 'classified_payload': {'analise_geral': [{'percentil': 10}, {'percentil': 20}, {'percentil': 30}, {'percentil': 40}, {'percentil': 50}, {'percentil': 60}]}},
                    {'instrument_code': 'srs2', 'classified_payload': {'resultados': [{'tscore': 66}, {'tscore': 61}, {'tscore': 59}, {'tscore': 68}, {'tscore': 66}, {'tscore': 65}, {'tscore': 65}] }},
                ]
            },
        )

        with ZipFile(BytesIO(docx_bytes), 'r') as docx:
            ns = ReportExportService.CHART_NS
            chart7 = LET.fromstring(docx.read('word/charts/chart7.xml'))
            chart8 = LET.fromstring(docx.read('word/charts/chart8.xml'))

        self.assertEqual(
            [node.text for node in chart7.findall('.//c:cat//c:pt/c:v', ns)],
            ['Panic / S.S.', 'AG', 'AS', 'FS', 'EE', 'TOTAL'],
        )
        self.assertEqual(
            [node.text for node in chart8.findall('.//c:cat//c:pt/c:v', ns)],
            ['Perc.S', 'Cog.S', 'Com.S', 'Mot.S', 'PRR', 'CIS', 'TOTAL'],
        )

    def test_populate_wasi_excel_charts_updates_all_present_chart_slots(self):
        def chart_xml():
            return b'''<?xml version="1.0" encoding="UTF-8"?>
<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart">
  <c:chart>
    <c:plotArea>
      <c:barChart>
        <c:ser>
          <c:idx val="0"/>
          <c:order val="0"/>
          <c:cat><c:strRef><c:f>Plan1!$A$1:$A$2</c:f><c:strCache><c:ptCount val="2"/><c:pt idx="0"><c:v>Old A</c:v></c:pt><c:pt idx="1"><c:v>Old B</c:v></c:pt></c:strCache></c:strRef></c:cat>
          <c:val><c:numRef><c:f>Plan1!$B$1:$B$2</c:f><c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="2"/><c:pt idx="0"><c:v>1</c:v></c:pt><c:pt idx="1"><c:v>2</c:v></c:pt></c:numCache></c:numRef></c:val>
        </c:ser>
      </c:barChart>
    </c:plotArea>
  </c:chart>
</c:chartSpace>'''

        buffer = BytesIO()
        with ZipFile(buffer, 'w') as docx:
            docx.writestr(
                '[Content_Types].xml',
                '''<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/charts/chart1.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
  <Override PartName="/word/charts/chart2.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
  <Override PartName="/word/charts/chart3.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
</Types>''',
            )
            docx.writestr(
                'word/document.xml',
                '''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><w:body>
<w:p><w:r><w:drawing><c:chart r:id="rId1"/></w:drawing></w:r></w:p>
<w:p><w:r><w:drawing><c:chart r:id="rId2"/></w:drawing></w:r></w:p>
<w:p><w:r><w:drawing><c:chart r:id="rId3"/></w:drawing></w:r></w:p>
</w:body></w:document>''',
            )
            docx.writestr(
                'word/_rels/document.xml.rels',
                '''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart1.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart2.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart3.xml"/>
</Relationships>''',
            )
            for idx in range(1, 4):
                docx.writestr(f'word/charts/chart{idx}.xml', chart_xml())
                docx.writestr(f'word/charts/_rels/chart{idx}.xml.rels', '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>')

        docx_bytes = ReportExportService._populate_wasi_excel_charts(
            buffer.getvalue(),
            {
                'validated_tests': [
                    {
                        'instrument_code': 'wasi',
                        'computed_payload': {
                            'composites': {
                                'qi_verbal': {'qi': 89},
                                'qi_execucao': {'qi': 89},
                                'qit_4': {'qi': 86},
                            }
                        },
                    },
                    {
                        'instrument_code': 'bpa2',
                        'bpa_chart_data': {
                            'domains': [
                                {'label': 'ATENÇÃO CONCENTRADA', 'values': {'maximo': 120, 'medio': 91, 'minimo': 63, 'bruto': 64, 'percentil': 20}},
                                {'label': 'ATENÇÃO DIVIDIDA', 'values': {'maximo': 119, 'medio': 83, 'minimo': 49, 'bruto': 80, 'percentil': 40}},
                                {'label': 'ATENÇÃO ALTERNADA', 'values': {'maximo': 120, 'medio': 95, 'minimo': 61, 'bruto': 86, 'percentil': 30}},
                                {'label': 'ATENÇÃO GERAL', 'values': {'maximo': 352, 'medio': 267, 'minimo': 217, 'bruto': 230, 'percentil': 25}},
                            ]
                        },
                    },
                    {
                        'instrument_code': 'ravlt',
                        'structured_results': {
                            'ravlt_esperado': {'a1': 7, 'a2': 9, 'a3': 11, 'a4': 12, 'a5': 12, 'b1': 6, 'a6': 12, 'a7': 11, 'r': 13, 'alt': 18, 'ret': 1, 'ip': 0.89, 'ir': 0.92},
                            'ravlt_minimo': {'a1': 6, 'a2': 8, 'a3': 10, 'a4': 10, 'a5': 11, 'b1': 5, 'a6': 9, 'a7': 9, 'r': 5, 'alt': 12, 'ret': 0.91, 'ip': 0.73, 'ir': 0.82},
                            'ravlt_obtido': {'a1': 4, 'a2': 5, 'a3': 7, 'a4': 7, 'a5': 7, 'b1': 4, 'a6': 5, 'a7': 5, 'r': 11, 'alt': 10, 'ret': 1, 'ip': 1, 'ir': 0.71},
                        },
                    },
                ]
            },
        )

        with ZipFile(BytesIO(docx_bytes), 'r') as docx:
            ns = ReportExportService.CHART_NS
            chart1 = LET.fromstring(docx.read('word/charts/chart1.xml'))
            chart2 = LET.fromstring(docx.read('word/charts/chart2.xml'))
            chart3 = LET.fromstring(docx.read('word/charts/chart3.xml'))

        self.assertEqual(
            [node.text for node in chart1.findall('.//c:cat//c:pt/c:v', ns)],
            ['QIE', 'QIV', 'QI TOTAL'],
        )
        self.assertEqual(
            [node.text for node in chart2.findall('.//c:val//c:pt/c:v', ns)][:5],
            ['120.0', '91.0', '63.0', '64.0', '20.0'],
        )
        self.assertEqual(chart3.tag, '{http://schemas.openxmlformats.org/drawingml/2006/chart}chartSpace')

    def test_etdah_ad_table_widths_follow_etdah_model(self):
        self.assertEqual(
            ReportExportService._table_widths("etdah_ad"),
            ReportExportService._table_widths("etdah"),
        )
        self.assertEqual(
            ReportExportService._table_widths("etdah_pais"),
            ReportExportService._table_widths("etdah"),
        )

    def test_document_chart_targets_follow_document_order(self):
        buffer = BytesIO()
        with ZipFile(buffer, 'w') as docx:
            docx.writestr(
                'word/document.xml',
                '''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:body>
    <w:p><w:r><w:drawing><c:chart r:id="rId9"/></w:drawing></w:r></w:p>
    <w:p><w:r><w:drawing><c:chart r:id="rId3"/></w:drawing></w:r></w:p>
    <w:p><w:r><w:drawing><c:chart r:id="rId11"/></w:drawing></w:r></w:p>
  </w:body>
</w:document>''',
            )
            docx.writestr(
                'word/_rels/document.xml.rels',
                '''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart7.xml"/>
  <Relationship Id="rId9" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart2.xml"/>
  <Relationship Id="rId11" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart10.xml"/>
</Relationships>''',
            )

        self.assertEqual(
            ReportExportService._document_chart_targets(buffer.getvalue()),
            [
                'word/charts/chart2.xml',
                'word/charts/chart7.xml',
                'word/charts/chart10.xml',
            ],
        )

    def test_fdt_chart_payload_uses_table_order_not_saved_chart_order(self):
        test = {
            'classified_payload': {
                'metric_results': [
                    {'codigo': 'contagem', 'media': 68.91, 'valor': 30, 'percentil_num': 12},
                    {'codigo': 'leitura', 'media': 37.44, 'valor': 23.3, 'percentil_num': 18},
                    {'codigo': 'escolha', 'media': 127, 'valor': 47.8, 'percentil_num': 20},
                    {'codigo': 'alternancia', 'media': 235, 'valor': 56.9, 'percentil_num': 22},
                    {'codigo': 'inibicao', 'media': 89.56, 'valor': 23.8, 'percentil_num': 24},
                    {'codigo': 'flexibilidade', 'media': 197.56, 'valor': 33.6, 'percentil_num': 26},
                ],
                'erros': {
                    'contagem': {'qtde_erros': 1},
                    'leitura': {'qtde_erros': 0},
                    'escolha': {'qtde_erros': 10},
                    'alternancia': {'qtde_erros': 16},
                },
                'charts': {
                    'automaticos': {
                        'categories': ['saved'],
                        'series': [{'values': [999]}],
                    },
                },
            }
        }

        automatic_categories, automatic_series = ReportExportService._fdt_chart_payload(test, automatic=True)
        controlled_categories, controlled_series = ReportExportService._fdt_chart_payload(test, automatic=False)

        self.assertEqual(
            automatic_categories,
            ['Tempo Médio', 'Tempo Obtido', 'Erros', 'Desempenho %', 'Indicativo de Déficit', 'Indicativo de Dificuldade Discreta', 'Sem Indicativo de Déficit'],
        )
        self.assertEqual(automatic_series[0], [68.91, 30.0, 1.0, 5.0, 5.0, 25.0, 50.0])
        self.assertEqual(automatic_series[1], [37.44, 23.3, 0.0, 5.0, 5.0, 25.0, 50.0])
        self.assertEqual(controlled_series[0], [127.0, 47.8, 10.0, 5.0, 5.0, 25.0, 50.0])
        self.assertEqual(controlled_series[3], [197.56, 33.6, 0.0, 25.0, 5.0, 25.0, 50.0])

    def test_ravlt_chart_payload_accepts_pt_br_numeric_strings(self):
        service = ReportExportService
        original = service._ravlt_rows
        service._ravlt_rows = classmethod(lambda cls, test, context=None: [
            ['Desempenho', 'A1', 'RET', 'I.P.'],
            ['Esperado', '6', '1', '0,88'],
            ['Mínimo', '5', '0,89', '0,73'],
            ['Obtido', '6', '0,75', '0,86'],
        ])
        try:
            categories, series = service._ravlt_chart_payload({}, None)
        finally:
            service._ravlt_rows = original

        self.assertEqual(categories, ['A1', 'RET', 'I.P.'])
        self.assertEqual(series, [[6.0, 1.0, 0.88], [5.0, 0.89, 0.73], [6.0, 0.75, 0.86]])

    def test_prune_unused_chart_parts_removes_unreferenced_chart_files(self):
        buffer = BytesIO()
        with ZipFile(buffer, 'w') as docx:
            docx.writestr(
                '[Content_Types].xml',
                '''<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/charts/chart1.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
  <Override PartName="/word/charts/chart2.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
</Types>''',
            )
            docx.writestr(
                'word/document.xml',
                '''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><w:body><w:p><w:r><w:drawing><c:chart r:id="rId1"/></w:drawing></w:r></w:p></w:body></w:document>''',
            )
            docx.writestr(
                'word/_rels/document.xml.rels',
                '''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart1.xml"/>
</Relationships>''',
            )
            docx.writestr('word/charts/chart1.xml', '<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart"/>')
            docx.writestr('word/charts/chart2.xml', '<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart"/>')
            docx.writestr('word/charts/_rels/chart1.xml.rels', '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>')
            docx.writestr('word/charts/_rels/chart2.xml.rels', '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>')

        pruned = ReportExportService._prune_unused_chart_parts(buffer.getvalue())
        with ZipFile(BytesIO(pruned), 'r') as docx:
            names = set(docx.namelist())
            self.assertIn('word/charts/chart1.xml', names)
            self.assertNotIn('word/charts/chart2.xml', names)
            self.assertNotIn('word/charts/_rels/chart2.xml.rels', names)
            self.assertNotIn('/word/charts/chart2.xml', docx.read('[Content_Types].xml').decode('utf-8'))
            rels_text = docx.read('word/_rels/document.xml.rels').decode('utf-8')
            self.assertIn('charts/chart1.xml', rels_text)
            self.assertNotIn('charts/chart2.xml', rels_text)

    def test_prune_unused_chart_parts_keeps_chart_user_shapes_targets(self):
        buffer = BytesIO()
        with ZipFile(buffer, 'w') as docx:
            docx.writestr(
                '[Content_Types].xml',
                '''<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/charts/chart1.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
  <Override PartName="/word/drawings/drawing1.xml" ContentType="application/vnd.openxmlformats-officedocument.drawing+xml"/>
</Types>''',
            )
            docx.writestr(
                'word/document.xml',
                '''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><w:body><w:p><w:r><w:drawing><c:chart r:id="rId1"/></w:drawing></w:r></w:p></w:body></w:document>''',
            )
            docx.writestr(
                'word/_rels/document.xml.rels',
                '''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart1.xml"/>
</Relationships>''',
            )
            docx.writestr('word/charts/chart1.xml', '<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart"/>')
            docx.writestr(
                'word/charts/_rels/chart1.xml.rels',
                '''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chartUserShapes" Target="../drawings/drawing1.xml"/>
</Relationships>''',
            )
            docx.writestr('word/drawings/drawing1.xml', '<xdr:wsDr xmlns:xdr="http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing"/>')

        pruned = ReportExportService._prune_unused_chart_parts(buffer.getvalue())
        with ZipFile(BytesIO(pruned), 'r') as docx:
            names = set(docx.namelist())
            self.assertIn('word/drawings/drawing1.xml', names)

    def test_sanitize_generated_document_removes_invalid_text_and_tables_after_references(self):
        document = Document()
        document.add_paragraph('Texto  com  espaços  extras .')
        document.add_paragraph('Sem conteúdo disponível para esta seção.')
        document.add_paragraph('17. REFERÊNCIAS BIBLIOGRÁFICAS')
        document.add_paragraph('WECHSLER, D. Referência válida.')
        document.add_paragraph('{"raw": true}')
        document.add_table(rows=1, cols=1)

        ReportExportService._sanitize_generated_document(
            document,
            report=SimpleNamespace(patient=SimpleNamespace(full_name='Maria Clara Souza')),
            context={'patient': {'full_name': 'Maria Clara Souza'}},
        )

        texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

        self.assertIn('Texto com espaços extras.', texts)
        self.assertIn('17. REFERÊNCIAS BIBLIOGRÁFICAS', texts)
        self.assertIn('WECHSLER, D. Referência válida.', texts)
        self.assertNotIn('Sem conteúdo disponível para esta seção.', texts)
        self.assertNotIn('{"raw": true}', texts)
        self.assertEqual(len(document.tables), 0)

    def test_validate_patient_identity_blocks_foreign_name_before_references(self):
        document = Document()
        document.add_paragraph('Nome: Maria Clara Souza')
        document.add_paragraph('A conclusão descreve João Vitor Almeida com outro histórico clínico.')

        with self.assertRaisesMessage(ValueError, 'nomes divergentes de pacientes'):
            ReportExportService._validate_patient_identity(
                document,
                report=SimpleNamespace(patient=SimpleNamespace(full_name='Maria Clara Souza')),
                context={'patient': {'full_name': 'Maria Clara Souza'}},
            )

    def test_validate_unique_wasi_result_blocks_distinct_qit_values(self):
        document = Document()
        document.add_paragraph('QI Total (QIT = 68), classificado como limítrofe.')
        document.add_paragraph('Em outra seção, QIT = 122 foi informado indevidamente.')

        with self.assertRaisesMessage(ValueError, 'resultados divergentes de QIT/QI Total'):
            ReportExportService._validate_unique_wasi_result(document)

    def test_validate_unique_wasi_result_allows_repeated_same_qit_value(self):
        document = Document()
        document.add_paragraph('QI Total (QIT = 86), classificado como média inferior.')
        document.add_paragraph('Síntese clínica: QIT = 86, mantendo a mesma interpretação.')

        ReportExportService._validate_unique_wasi_result(document)

    def test_append_body_element_before_sectpr_keeps_section_properties_last(self):
        document = Document()
        document.add_paragraph('Antes')

        new_paragraph = parse_xml(
            '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:t>Chart placeholder</w:t></w:r></w:p>'
        )

        ReportExportService._append_body_element_before_sectpr(document, new_paragraph)

        body_children = list(document._body._element.iterchildren())
        self.assertEqual(body_children[-1].tag.split('}')[-1], 'sectPr')
        self.assertEqual(body_children[-2].tag.split('}')[-1], 'p')
