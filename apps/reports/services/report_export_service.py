from __future__ import annotations

from copy import deepcopy
from io import BytesIO
import logging
import posixpath
import re
from pathlib import Path
from pathlib import PurePosixPath
from zipfile import ZipFile

import matplotlib
from lxml import etree as ET

matplotlib.use("Agg")
logging.getLogger("matplotlib").setLevel(logging.WARNING)

import matplotlib.pyplot as plt
import numpy as np
from matplotlib import font_manager
from matplotlib.patches import Rectangle
from docx import Document
from docx.table import Table
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor
from docx.shared import Cm

from apps.reports.charts import gerar_grafico_bpa_bytes, gerar_grafico_wasi_bytes
from apps.reports.models import Report
from apps.reports.specs import TABLE_LAYOUT_SPECS, WASI_CHART_SPEC, WASI_LAYOUT_SPEC, WASI_REPORT_SPEC, WASI_TABLE_SPECS
from apps.reports.builders.references_builder import build_references
from apps.reports.builders.wais3_report_builder import WAIS3ReportBuilder
from apps.reports.services.report_context_service import ReportContextService
from apps.reports.services.ptbr_text_service import PtBrTextService
from apps.reports.services.wisc4_standardization import WISC4StandardizationService
from apps.reports.services.wais3_standardization import WAIS3StandardizationService
from apps.tests.bpa2.interpreters import build_clinical_summary, SUBTEST_OPENINGS
from apps.tests.etdah_pais.interpreters import (
    FACTOR_LABELS as ETDAH_PAIS_FACTOR_LABELS,
    FACTOR_ORDER as ETDAH_PAIS_FACTOR_ORDER,
    _analysis_text as build_etdah_pais_analysis_text,
    _factor_paragraph as build_etdah_pais_factor_paragraph,
)
from apps.tests.srs2.interpreters import interpret_srs2_results
from apps.tests.ravlt.norms import NORMS as RAVLT_NORMS
from apps.tests.ravlt.norms import get_age_band as get_ravlt_age_band
from apps.tests.wisc4.calculators import _calcular_idade, _carregar_tabela_ncp
from apps.tests.wisc4.interpreters import interpret_wisc4_profile


ET.register_namespace("a", "http://schemas.openxmlformats.org/drawingml/2006/main")
ET.register_namespace("c", "http://schemas.openxmlformats.org/drawingml/2006/chart")
ET.register_namespace("mc", "http://schemas.openxmlformats.org/markup-compatibility/2006")
ET.register_namespace("r", "http://schemas.openxmlformats.org/officeDocument/2006/relationships")
ET.register_namespace("c14", "http://schemas.microsoft.com/office/drawing/2007/8/2/chart")
ET.register_namespace("c16r2", "http://schemas.microsoft.com/office/drawing/2015/06/chart")
ET.register_namespace("cx", "http://schemas.microsoft.com/office/drawing/2014/chartex")
ET.register_namespace("cx1", "http://schemas.microsoft.com/office/drawing/2015/9/8/chartex")
ET.register_namespace("cx2", "http://schemas.microsoft.com/office/drawing/2015/10/21/chartex")
ET.register_namespace("cx3", "http://schemas.microsoft.com/office/drawing/2016/5/9/chartex")
ET.register_namespace("cx4", "http://schemas.microsoft.com/office/drawing/2016/5/10/chartex")
ET.register_namespace("cx5", "http://schemas.microsoft.com/office/drawing/2016/5/11/chartex")
ET.register_namespace("cx6", "http://schemas.microsoft.com/office/drawing/2016/5/12/chartex")
ET.register_namespace("cx7", "http://schemas.microsoft.com/office/drawing/2016/5/13/chartex")
ET.register_namespace("cx8", "http://schemas.microsoft.com/office/drawing/2016/5/14/chartex")


class ReportExportService:
    TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates_assets"
    LAUDOS_TEMPLATE_DIR = TEMPLATE_DIR / "laudo-modelo"
    DEFAULT_TEMPLATE_PATH = LAUDOS_TEMPLATE_DIR / "PAPEL-TIMBRADO-MODELO.docx"
    WAIS3_TEMPLATE_PATH = LAUDOS_TEMPLATE_DIR / "Modelo-WAIS3.docx"
    WISC4_TEMPLATE_PATH = LAUDOS_TEMPLATE_DIR / "Modelo-WISC4.docx"
    WASI_TEMPLATE_PATH = LAUDOS_TEMPLATE_DIR / "Modelo-WASI.docx"
    TABLE_STYLE_SOURCE_PATH = WISC4_TEMPLATE_PATH
    FONT_NAME = WASI_LAYOUT_SPEC["font_family"]
    BODY_SIZE = Pt(12)
    TABLE_SIZE = Pt(WASI_LAYOUT_SPEC["table_size_pt"])
    TABLE_HEADER_SIZE = Pt(WASI_LAYOUT_SPEC["table_header_size_pt"])
    TITLE_SIZE = Pt(12)
    CAPTION_SIZE = Pt(WASI_LAYOUT_SPEC["caption_size_pt"])
    TABLE_CELL_MARGIN_TOP = 30
    TABLE_CELL_MARGIN_START = 60
    TABLE_CELL_MARGIN_BOTTOM = 30
    TABLE_CELL_MARGIN_END = 60
    BODY_FIRST_LINE_INDENT = Cm(1.5)
    BODY_LINE_SPACING = 1.5
    CHART_NS = {
        "c": "http://schemas.openxmlformats.org/drawingml/2006/chart",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    IDENTIFICATION_LINE_SPACING = 1.15
    HEADING_SPACE_BEFORE = Pt(10)
    HEADING_SPACE_AFTER = Pt(4)
    SUBHEADING_SPACE_BEFORE = Pt(6)
    SUBHEADING_SPACE_AFTER = Pt(2)
    BODY_SPACE_AFTER = Pt(6)
    IMAGE_WIDTH = Inches(6.2)
    DEFAULT_TABLE_WIDTH = Inches(6.5)
    HEADER_FILL = "DCE6F1"
    BPA_HEADER_FILL = "A8D08D"
    BPA_BODY_FILL = "E2EFD9"
    WISC_HEADER_FILL = "A8D08D"
    WISC_NAME_FILL = "A9D18E"
    WISC_VALUE_FILL = "E2EFD9"
    RAVLT_HEADER_FILL = "A9D08E"
    FDT_HEADER_FILL = "A8D08D"
    FDT_BODY_FILL = "E2EFD9"
    TABLE_TITLE_FILL = "538135"
    LOCAL_TIMES_NEW_ROMAN = Path("/mnt/c/Windows/Fonts/times.ttf")
    LOCAL_LIBERATION_SERIF = Path.home() / ".local/share/fonts/liberation/LiberationSerif-Regular.ttf"
    MATPLOTLIB_FALLBACK_FONT = "DejaVu Serif"
    INTERPRETATION_LABELS = (
        "Interpretação clínica:",
        "Interpretação e Observações Clínicas:",
    )
    EMPTY_INTERPRETATION_MESSAGES = {
        "Nenhum instrumento específico deste domínio apresentou interpretação clínica consolidada.",
    }
    FIXED_AUTHOR = "Jacqueline Oliveira Caires (CRP 09/6017)"
    FIXED_INTERESTED_PARTY = "Familiares"
    FIXED_PURPOSE = "Averiguação das capacidades cognitivas para auxílio diagnóstico"
    PRIMARY_REPORT_TEST_CODES = ("wisc4", "wais3", "wasi")
    INVALID_DOCX_PATTERNS = (
        "Sem conteúdo disponível para esta seção.",
        "undefined",
        "null",
        "None",
        "Interp Interpretação",
        "Parte superior do formulário",
        "Parte inferior do formulário",
    )
    REFERENCE_SECTION_HEADINGS = (
        "17. REFERÊNCIAS BIBLIOGRÁFICAS",
        "REFERÊNCIAS BIBLIOGRÁFICAS",
        "REFERENCIA BIBLIOGRÁFICA",
        "REFERÊNCIAS",
    )
    SKIP_PATTERNS = (
        "Laudo", "Neuropsicológico", "Documento", "Familiares", "Finalidade",
        "Masculino", "Feminino", "Data", "Segunda", "Edição", "Dígitos",
        "Child", "Anxiety", "Related", "Emotional", "Disorders",
        "Aprendizagem", "Auditivo", "Versão", "Pais", "História", "Pessoal",
        "Não", "Sim", "Resultado", "Observação", "Índice", "Nota",
        "Encaminhamento", "Interessado", "Autora", "Filiação", "Escolaridade",
        "Profissão", "Gênero", "Paciente", "Capacidade Cognitiva", "Global",
        "Wechsler", "Abreviada", "Escala", "Psicológica", "Conselho",
        "Média", "Inferior", "Superior", "Média Superior", "Média Inferior",
        "Abaixo", "Acima", "Faixa", "Limítrofe", "Muito Superior",
        "Quociente", "Inteligência", "Escalas", "Primárias",
        "Atenção", "Concentrada", "Dividida", "Alternada", "Processos",
        "Automáticos", "Gráfico", "Cinco", "Grandes", "Fatores",
        "Interpretação", "Clínica", "Cognição", "Social", "Comunicação",
        "Análise", "Integrada", "Os", "Na", "Em", "Entre", "Diferenças",
    )
    logger = logging.getLogger(__name__)

    @classmethod
    def _wisc_section_text(cls, sections: dict, section_key: str, context: dict) -> str:
        text = (sections.get(section_key) or "").strip()
        if text:
            return text

        if WISC4StandardizationService.supports(section_key, context):
            return PtBrTextService.normalize(WISC4StandardizationService.build(section_key, context))

        if WAIS3StandardizationService.supports(section_key, context):
            return PtBrTextService.normalize(WAIS3StandardizationService.build(section_key, context))

        return ""
    INSTRUMENT_CATALOG = {
        "anamnese": {
            "nome": "Anamnese",
            "descricao": "Entrevista clínica inicial destinada ao levantamento da história do desenvolvimento, queixas atuais, antecedentes médicos, escolares, familiares, emocionais e comportamentais, com o objetivo de contextualizar os dados obtidos nos demais instrumentos.",
        },
        "wisc4": {
            "nome": "Escala Wechsler de Inteligência para Crianças – Quarta Edição (WISC-IV)",
            "descricao": "Utilizada para avaliar o funcionamento intelectual global, os índices fatoriais (Compreensão Verbal, Organização Perceptual, Memória Operacional e Velocidade de Processamento) e fornecer indicadores sobre o perfil cognitivo.",
        },
        "wais3": {
            "nome": "Escala Wechsler de Inteligência para Adultos – Terceira Edição (WAIS-III)",
            "descricao": "Instrumento destinado à avaliação da capacidade intelectual global em adultos, permitindo analisar habilidades verbais, não verbais, memória operacional e velocidade de processamento, além de oferecer indicadores sobre o perfil cognitivo geral.",
        },
        "wasi": {
            "nome": "Escala Wechsler Abreviada de Inteligência (WASI)",
            "descricao": "Utilizada para estimar o funcionamento intelectual global, verbal e de execução, bem como fornecer indicadores breves sobre a capacidade cognitiva geral.",
        },
        "bpa2": {
            "nome": "Bateria Psicológica para Avaliação da Atenção – Segunda Edição (BPA-2)",
            "descricao": "Avalia a capacidade geral de atenção, incluindo atenção concentrada, alternada, dividida e atenção global.",
        },
        "fdt": {
            "nome": "Teste dos Cinco Dígitos (FDT)",
            "descricao": "Investiga a velocidade de processamento, controle inibitório, alternância e flexibilidade cognitiva.",
        },
        "ravlt": {
            "nome": "Teste de Aprendizagem Auditivo-Verbal de Rey (RAVLT)",
            "descricao": "Avalia memória verbal episódica, aquisição, retenção, recuperação e resistência à interferência.",
        },
        "figuras_complexas_rey": {
            "nome": "Figuras Complexas de Rey",
            "descricao": "Instrumento utilizado para avaliar habilidades visuoconstrutivas, planejamento perceptomotor e memória visual, por meio das etapas de cópia e evocação.",
        },
        "etdah_ad": {
            "nome": "Escala para Transtorno de Déficit de Atenção e Hiperatividade – Autorrelato (E-TDAH-AD)",
            "descricao": "Questionário destinado à identificação de sintomas relacionados à desatenção, impulsividade, hiperatividade, aspectos emocionais e dificuldades funcionais associadas ao TDAH, com base no autorrelato.",
        },
        "etdah_pais": {
            "nome": "Escala para Transtorno de Déficit de Atenção e Hiperatividade – Versão Pais (E-TDAH-PAIS)",
            "descricao": "Questionário para identificação de sintomas de desatenção, impulsividade, hiperatividade, regulação emocional e comportamento adaptativo, a partir da percepção dos responsáveis.",
        },
        "scared": {
            "nome": "SCARED – Autorrelato (Screen for Child Anxiety Related Emotional Disorders)",
            "descricao": "Avalia sintomas ansiosos em crianças e adolescentes, incluindo ansiedade generalizada, ansiedade de separação, fobia social, pânico e sintomas somáticos, com base no autorrelato.",
        },
        "scared_mae": {
            "nome": "SCARED – Versão Mãe",
            "descricao": "Avalia sintomas ansiosos, incluindo ansiedade generalizada, de separação, fobia social e somatizações, a partir da percepção materna.",
        },
        "srs2": {
            "nome": "Escala de Responsividade Social – Segunda Edição (SRS-2)",
            "descricao": "Rastreia dificuldades em comunicação social, cognição social, padrões restritos e repetitivos e comportamentos associados ao espectro autista.",
        },
        "epq_j": {
            "nome": "Inventário de Personalidade de Eysenck para Jovens (EPQ-J)",
            "descricao": "Avalia traços de personalidade em crianças e adolescentes, contemplando dimensões como extroversão, neuroticismo, psicoticismo e sinceridade.",
        },
        "bfp": {
            "nome": "Bateria Fatorial de Personalidade (BFP)",
            "descricao": "Instrumento fundamentado no modelo dos Cinco Grandes Fatores, destinado à avaliação dos traços de personalidade: neuroticismo, extroversão, socialização, realização e abertura à experiência.",
        },
        "iphexa": {
            "nome": "Inventário de Personalidade HEXA-Flexível para Adultos (IPHEXA)",
            "descricao": "Avalia traços de personalidade com base em dimensões amplas do funcionamento da personalidade, contribuindo para a compreensão do estilo emocional, interpessoal e adaptativo do avaliado.",
        },
        "bai": {
            "nome": "Inventário de Ansiedade de Beck (BAI)",
            "descricao": "Instrumento de autorrelato utilizado para mensurar a intensidade de sintomas ansiosos, com ênfase em manifestações fisiológicas e subjetivas de ansiedade.",
        },
        "bdi": {
            "nome": "Inventário de Depressão de Beck (BDI)",
            "descricao": "Instrumento de autorrelato utilizado para avaliar a intensidade de sintomas depressivos, incluindo aspectos afetivos, cognitivos, motivacionais e somáticos.",
        },
        "ebadep_a": {
            "nome": "Escala Baptista de Depressão – Versão Adulto (EBADEP-A)",
            "descricao": "Avalia a presença e a intensidade de sintomas depressivos em adultos, contemplando aspectos emocionais, cognitivos, comportamentais e fisiológicos.",
        },
        "ebadep_ij": {
            "nome": "Escala Baptista de Depressão – Versão Infantojuvenil (EBADEP-IJ)",
            "descricao": "Instrumento destinado à identificação de sintomas depressivos em crianças e adolescentes, abrangendo indicadores emocionais, cognitivos e comportamentais.",
        },
        "ebaped_ij": {
            "nome": "Escala Baptista de Depressão – Versão Infantojuvenil (EBADEP-IJ)",
            "descricao": "Instrumento destinado à identificação de sintomas depressivos em crianças e adolescentes, abrangendo indicadores emocionais, cognitivos e comportamentais.",
        },
        "scq": {
            "nome": "Social Communication Questionnaire (SCQ)",
            "descricao": "Questionário de rastreio voltado à identificação de dificuldades na comunicação social e comportamentos associados ao Transtorno do Espectro Autista.",
        },
        "cars2_hf": {
            "nome": "Childhood Autism Rating Scale – High Functioning (CARS2-HF)",
            "descricao": "Escala destinada à avaliação de comportamentos associados ao Transtorno do Espectro Autista em indivíduos com maior nível de funcionamento, considerando interação social, comunicação, flexibilidade e padrões comportamentais.",
        },
        "mchat": {
            "nome": "Modified Checklist for Autism in Toddlers (M-CHAT)",
            "descricao": "Instrumento de rastreio precoce para sinais de risco para Transtorno do Espectro Autista em crianças pequenas, com foco em interação social, atenção compartilhada e comportamentos comunicativos.",
        },
        "son_r_2_5_7": {
            "nome": "SON-R 2½–7[a]",
            "descricao": "Teste não verbal de inteligência utilizado para avaliar o raciocínio geral, a percepção visual e a organização espacial, minimizando a influência da linguagem.",
        },
        "thcp": {
            "nome": "Teste de Habilidades Cognitivas Primárias (THCP)",
            "descricao": "Instrumento utilizado para investigar habilidades cognitivas básicas relacionadas ao desenvolvimento infantil, contribuindo para a análise do perfil cognitivo e do potencial de aprendizagem.",
        },
        "mmse_2": {
            "nome": "Miniexame do Estado Mental – Segunda Edição (MMSE-2)",
            "descricao": "Instrumento de rastreio cognitivo utilizado para avaliar orientação, atenção, memória, linguagem e habilidades construtivas, auxiliando na investigação de possíveis alterações cognitivas.",
        },
        "vineland": {
            "nome": "Vineland – Escala de Comportamento Adaptativo",
            "descricao": "Avalia o comportamento adaptativo em domínios como comunicação, socialização, habilidades de vida diária e, em algumas versões, habilidades motoras, contribuindo para a compreensão do funcionamento adaptativo global.",
        },
        "portage": {
            "nome": "Portage",
            "descricao": "Instrumento utilizado para avaliação do desenvolvimento infantil em diferentes áreas, como linguagem, cognição, socialização, autocuidados e desenvolvimento motor.",
        },
        "htp": {
            "nome": "HTP (House-Tree-Person)",
            "descricao": "Técnica projetiva gráfica utilizada para investigação de aspectos emocionais, vivenciais e da dinâmica da personalidade, por meio da interpretação dos desenhos da casa, árvore e pessoa.",
        },
    }

    @classmethod
    def _chart_font(cls):
        if cls.LOCAL_TIMES_NEW_ROMAN.exists():
            return font_manager.FontProperties(fname=str(cls.LOCAL_TIMES_NEW_ROMAN))
        if cls.LOCAL_LIBERATION_SERIF.exists():
            return font_manager.FontProperties(fname=str(cls.LOCAL_LIBERATION_SERIF))
        return font_manager.FontProperties(family=cls.MATPLOTLIB_FALLBACK_FONT)

    @classmethod
    def _apply_figure_border(cls, fig, color: str = "#D9D9D9", linewidth: float = 1.0):
        fig.add_artist(
            Rectangle(
                (0, 0),
                1,
                1,
                transform=fig.transFigure,
                fill=False,
                edgecolor=color,
                linewidth=linewidth,
            )
        )
    PROCEDURES_ORDER = [
        "anamnese",
        "wisc4",
        "wais3",
        "wasi",
        "son_r_2_5_7",
        "thcp",
        "mmse_2",
        "vineland",
        "portage",
        "bpa2",
        "fdt",
        "ravlt",
        "figuras_complexas_rey",
        "etdah_pais",
        "etdah_ad",
        "scared",
        "scared_mae",
        "srs2",
        "epq_j",
        "bfp",
        "iphexa",
        "bai",
        "bdi",
        "ebadep_a",
        "ebadep_ij",
        "ebaped_ij",
        "scq",
        "cars2_hf",
        "mchat",
        "htp",
    ]

    @staticmethod
    def generate_html(report: Report):
        ReportContextService.sync_report_context(report)
        content = str(report.final_text or report.edited_text or "")
        created_at = report.created_at.strftime("%d/%m/%Y") if report.created_at else ""
        patient_name = report.patient.full_name if report.patient else ""
        author_name = ReportExportService.FIXED_AUTHOR
        interested_party = ReportExportService.FIXED_INTERESTED_PARTY
        purpose = ReportExportService.FIXED_PURPOSE

        return f"""
        <html>
        <head>
            <style>
                body {{ font-family: 'Times New Roman', serif; line-height: 1.5; color: #222; font-size: 12pt; }}
                h1 {{ text-align: center; font-size: 16pt; }}
                .metadata {{ margin-bottom: 24px; }}
            </style>
        </head>
        <body>
            <h1>{report.title}</h1>
            <div class="metadata">
                <p><b>Paciente:</b> {patient_name}</p>
                <p><b>Profissional:</b> {author_name}</p>
                <p><b>Interessado:</b> {interested_party}</p>
                <p><b>Finalidade:</b> {purpose}</p>
                <p><b>Data:</b> {created_at}</p>
            </div>
            <div>{content.replace(chr(10), "<br>")}</div>
        </body>
        </html>
        """

    @classmethod
    def generate_docx_bytes(cls, report: Report) -> bytes:
        context = ReportContextService.sync_report_context(report)
        context = dict(context) if isinstance(context, dict) else {}
        sections = {
            section.key: cls._sanitize_section_text_for_patient(
                str(section.content_edited or section.content_generated or ""),
                context,
            )
            for section in report.sections.all()
        }

        template_path = cls._select_template_path(report, context)
        if cls._is_adolescent_context(context, report):
            document = cls._build_adolescent_document(report, context, sections)
        elif not template_path.exists():
            cls.logger.warning(
                "Template DOCX ausente em %s. Gerando fallback sem papel timbrado.",
                template_path,
            )
            document = cls._build_fallback_document(report, context)
        else:
            document = Document(str(template_path))
            if not cls._template_has_editable_body(document):
                cls.logger.warning(
                    "Template DOCX em %s não possui corpo editável. Gerando fallback com conteúdo do laudo.",
                    template_path,
                )
                document = cls._build_fallback_document(report, context)
            else:
                cls._ensure_model_table_styles(document)
                cls._apply_base_styles(document)
                cls._replace_simple_sections(document, report, sections, context)
                cls._rebuild_qualitative_section(document, sections, context)
                cls._populate_wasi_tables(document, context)

        cls._ensure_model_table_styles(document)
        cls._normalize_model_header_footer(document)

        body = document._body._element
        ns_c = "{http://schemas.openxmlformats.org/drawingml/2006/chart}"
        charts_in_doc = [e for e in body.iter() if e.tag == f"{ns_c}chart"]
        cls._sanitize_generated_document(document, report, context)
        if charts_in_doc:
            cls._remove_invalid_content_after_references(document)

        cls._validate_patient_identity(document, report, context)
        cls._validate_unique_wasi_result(document)

        output = BytesIO()
        document.save(output)
        docx_bytes = output.getvalue()
        header_footer_template_path = cls._header_footer_template_path(template_path, report, context)
        if header_footer_template_path.exists():
            docx_bytes = cls._restore_template_header_footer(docx_bytes, header_footer_template_path)

        pre_chart_docx_bytes = docx_bytes
        if cls._find_test(context, "wisc4"):
            docx_bytes = cls._populate_wisc4_excel_charts(docx_bytes, context)
        if cls._find_test(context, "wais3") and not cls._is_adolescent_context(context, report):
            docx_bytes = cls._populate_wais3_excel_charts(docx_bytes, context)
        if cls._find_test(context, "wasi"):
            docx_bytes = cls._populate_wasi_excel_charts(docx_bytes, context)
        if not cls._docx_package_is_valid(docx_bytes):
            cls.logger.warning(
                "DOCX gerado com estrutura invalida apos atualizacao de graficos; retornando versao anterior aos graficos."
            )
            return pre_chart_docx_bytes
        return docx_bytes

    @classmethod
    def _header_footer_template_path(cls, template_path: Path, report, context: dict) -> Path:
        if cls._is_adolescent_context(context, report):
            return cls.DEFAULT_TEMPLATE_PATH
        if cls._primary_report_test_code(context) == "wais3" and template_path.exists():
            return template_path
        return cls.DEFAULT_TEMPLATE_PATH

    @classmethod
    def _select_template_path(cls, report: Report, context: dict | None = None):
        available_codes = set(cls._available_test_codes(context or {}))

        if "wasi" in available_codes and cls.WASI_TEMPLATE_PATH.exists():
            return cls.WASI_TEMPLATE_PATH

        primary_test_code = cls._primary_report_test_code(context or {}) if context else None

        if primary_test_code == "wisc4" and cls.WISC4_TEMPLATE_PATH.exists():
            return cls.WISC4_TEMPLATE_PATH
        if primary_test_code == "wais3" and cls.WAIS3_TEMPLATE_PATH.exists():
            return cls.WAIS3_TEMPLATE_PATH

        return cls.DEFAULT_TEMPLATE_PATH

    @classmethod
    def _available_test_codes(cls, context: dict) -> list[str]:
        return [
            item.get("instrument_code")
            for item in context.get("validated_tests") or []
            if item.get("instrument_code")
        ]

    @classmethod
    def _primary_report_test_code(cls, context: dict) -> str | None:
        available_codes = set(cls._available_test_codes(context))
        for code in cls.PRIMARY_REPORT_TEST_CODES:
            if code in available_codes:
                return code
        return None

    @classmethod
    def _extract_template_chart_blocks(cls, document: Document) -> list:
        blocks = []
        for child in document._body._element.iterchildren():
            if cls._element_contains_chart(child):
                blocks.append(deepcopy(child))
        if blocks:
            return blocks
        ns_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        for paragraph in document._body._element.iter(ns_w + 'p'):
            if cls._element_contains_chart(paragraph):
                blocks.append(deepcopy(paragraph))
        return blocks

    @classmethod
    def _extract_chart_blocks_from_template_path(cls, template_path: Path) -> list:
        if not template_path.exists():
            return []
        return cls._extract_template_chart_blocks(Document(str(template_path)))

    @classmethod
    def _is_etdah_table_key(cls, table_key: str) -> bool:
        return table_key in {"etdah", "etdah_ad", "etdah_pais"}

    @classmethod
    def _wais3_report_builder(cls, context: dict) -> WAIS3ReportBuilder:
        available = set(cls._available_test_codes(context))
        is_adolescent = cls._is_adolescent_context(context, None)
        if is_adolescent:
            return WAIS3ReportBuilder.for_adolescent(available)
        return WAIS3ReportBuilder.for_adult(available)

    @classmethod
    def _element_contains_chart(cls, element) -> bool:
        try:
            return bool(
                element.xpath(
                    './/*[local-name()="chart" and namespace-uri()="http://schemas.openxmlformats.org/drawingml/2006/chart"]'
                )
            )
        except TypeError:
            xml = getattr(element, "xml", "")
            return (
                "charts/chart" in xml
                or 'uri="http://schemas.openxmlformats.org/drawingml/2006/chart"' in xml
                or re.search(r"<[A-Za-z0-9_]+:chart\\b", xml) is not None
            )

    @classmethod
    def _chart_series(cls, root, index: int):
        series = root.findall('.//c:ser', cls.CHART_NS)
        if index >= len(series):
            cls.logger.warning("Chart has %d series, requested index %d", len(series), index)
            return None
        return series[index]

    @classmethod
    def _set_chart_cache_points(cls, cache_element, values: list, tag_name: str):
        if cache_element is None:
            return
        pt_count = cache_element.find('c:ptCount', cls.CHART_NS)
        existing_points = cache_element.findall('c:pt', cls.CHART_NS)
        existing_indexes = [point.get('idx', str(idx)) for idx, point in enumerate(existing_points)]
        for child in list(cache_element.findall('c:pt', cls.CHART_NS)):
            cache_element.remove(child)
        if pt_count is None:
            pt_count = ET.SubElement(cache_element, f'{{{cls.CHART_NS["c"]}}}ptCount')
        pt_count.set('val', str(len(values)))
        indexes = existing_indexes if len(existing_indexes) == len(values) else [str(idx) for idx, _ in enumerate(values)]
        for idx, value in enumerate(values):
            point = ET.SubElement(cache_element, f'{{{cls.CHART_NS["c"]}}}pt')
            point.set('idx', indexes[idx])
            value_node = ET.SubElement(point, f'{{{cls.CHART_NS["c"]}}}{tag_name}')
            value_node.text = str(value)

    @classmethod
    def _replace_child(cls, parent, old_child, new_child):
        children = list(parent)
        index = children.index(old_child)
        parent.remove(old_child)
        parent.insert(index, new_child)

    @classmethod
    def _inline_series_title(cls, series):
        tx = series.find('c:tx', cls.CHART_NS)
        if tx is None:
            return
        str_ref = tx.find('c:strRef', cls.CHART_NS)
        if str_ref is None:
            return
        cache = str_ref.find('c:strCache', cls.CHART_NS)
        value = ""
        if cache is not None:
            point = cache.find('c:pt/c:v', cls.CHART_NS)
            value = point.text if point is not None else ""
        literal = ET.Element(f'{{{cls.CHART_NS["c"]}}}v')
        literal.text = value
        cls._replace_child(tx, str_ref, literal)

    @classmethod
    def _inline_string_categories(cls, category_node):
        str_ref = category_node.find('c:strRef', cls.CHART_NS)
        if str_ref is not None:
            cache = str_ref.find('c:strCache', cls.CHART_NS)
            literal = ET.Element(f'{{{cls.CHART_NS["c"]}}}strLit')
            if cache is not None:
                cls._copy_cache(cache, literal)
            cls._replace_child(category_node, str_ref, literal)
            return

        multi_ref = category_node.find('c:multiLvlStrRef', cls.CHART_NS)
        if multi_ref is not None:
            cache = multi_ref.find('c:multiLvlStrCache', cls.CHART_NS)
            literal = ET.Element(f'{{{cls.CHART_NS["c"]}}}strLit')
            if cache is not None:
                values_by_idx: dict[int, str] = {}
                max_idx = -1
                for level in cache.findall('c:lvl', cls.CHART_NS):
                    for point in level.findall('c:pt', cls.CHART_NS):
                        idx = int(point.get('idx', '0'))
                        max_idx = max(max_idx, idx)
                        text = point.findtext('c:v', default='', namespaces=cls.CHART_NS)
                        if idx not in values_by_idx and text.strip():
                            values_by_idx[idx] = text
                pt_count = ET.SubElement(literal, f'{{{cls.CHART_NS["c"]}}}ptCount')
                pt_count.set('val', str(max_idx + 1 if max_idx >= 0 else 0))
                for idx in range(max_idx + 1):
                    point = ET.SubElement(literal, f'{{{cls.CHART_NS["c"]}}}pt')
                    point.set('idx', str(idx))
                    value_node = ET.SubElement(point, f'{{{cls.CHART_NS["c"]}}}v')
                    value_node.text = values_by_idx.get(idx, '')
            cls._replace_child(category_node, multi_ref, literal)

    @classmethod
    def _inline_numeric_values(cls, value_node):
        num_ref = value_node.find('c:numRef', cls.CHART_NS)
        if num_ref is None:
            return
        cache = num_ref.find('c:numCache', cls.CHART_NS)
        literal = ET.Element(f'{{{cls.CHART_NS["c"]}}}numLit')
        if cache is not None:
            cls._copy_cache(cache, literal)
        cls._replace_child(value_node, num_ref, literal)

    @classmethod
    def _inline_cached_reference(cls, node):
        if node is None:
            return

        str_ref = node.find('c:strRef', cls.CHART_NS)
        if str_ref is not None:
            cache = str_ref.find('c:strCache', cls.CHART_NS)
            literal = ET.Element(f'{{{cls.CHART_NS["c"]}}}strLit')
            if cache is not None:
                cls._copy_cache(cache, literal)
            cls._replace_child(node, str_ref, literal)
            return

        multi_ref = node.find('c:multiLvlStrRef', cls.CHART_NS)
        if multi_ref is not None:
            cls._inline_string_categories(node)
            return

        cls._inline_numeric_values(node)

    @classmethod
    def _copy_cache(cls, source_cache, target_literal):
        format_code = source_cache.find('c:formatCode', cls.CHART_NS)
        if format_code is not None:
            cloned_format = ET.SubElement(target_literal, f'{{{cls.CHART_NS["c"]}}}formatCode')
            cloned_format.text = format_code.text
        pt_count = source_cache.find('c:ptCount', cls.CHART_NS)
        if pt_count is not None:
            cloned_count = ET.SubElement(target_literal, f'{{{cls.CHART_NS["c"]}}}ptCount')
            cloned_count.set('val', pt_count.get('val', '0'))
        for point in source_cache.findall('c:pt', cls.CHART_NS):
            target_literal.append(deepcopy(point))

    @classmethod
    def _strip_external_chart_relationships(cls, rels_bytes: bytes) -> bytes:
        root = ET.fromstring(rels_bytes)
        for rel in list(root):
            target_mode = rel.attrib.get('TargetMode')
            rel_type = rel.attrib.get('Type', '')
            if target_mode == 'External' or rel_type.endswith('/oleObject'):
                root.remove(rel)
        return ET.tostring(root, encoding='utf-8', xml_declaration=True)

    @classmethod
    def _update_direct_chart_values(cls, value_node, values: list[float]):
        c_ns = cls.CHART_NS["c"]
        literal_node = value_node.find(f"{{{c_ns}}}numLit")
        if literal_node is None:
            literal_node = ET.SubElement(value_node, f"{{{c_ns}}}numLit")
        existing_pts = {
            pt.get("idx", str(i)): pt
            for i, pt in enumerate(literal_node.findall(f"{{{c_ns}}}pt"))
        }
        pt_count = literal_node.find(f"{{{c_ns}}}ptCount")
        if pt_count is None:
            pt_count = ET.SubElement(literal_node, f"{{{c_ns}}}ptCount")
        pt_count.set("val", str(len(values)))
        for key, pt in list(existing_pts.items()):
            try:
                idx = int(key)
            except ValueError:
                idx = -1
            if idx >= len(values):
                literal_node.remove(pt)
                existing_pts.pop(key, None)
        for i, v in enumerate(values):
            pt = existing_pts.get(str(i))
            if pt is None:
                pt = ET.SubElement(literal_node, f"{{{c_ns}}}pt")
                pt.set("idx", str(i))
            for child in list(pt):
                if child.tag == f"{{{c_ns}}}v":
                    pt.remove(child)
            v_elem = ET.SubElement(pt, f"{{{c_ns}}}v")
            v_elem.text = str(v) if v is not None else "0"

    @classmethod
    def _update_chart_series(cls, root, index: int, categories: list[str], values: list[float]):
        series = cls._chart_series(root, index)
        if series is None:
            return
        category_cache = series.find('.//c:cat//c:strCache', cls.CHART_NS)
        value_cache = series.find('.//c:val//c:numCache', cls.CHART_NS)
        cls._set_chart_cache_points(category_cache, categories, 'v')
        cls._set_chart_cache_points(value_cache, values, 'v')
        cls._inline_series_title(series)
        category_node = series.find('c:cat', cls.CHART_NS)
        cls._inline_cached_reference(category_node)
        value_node = series.find('c:val', cls.CHART_NS)
        if value_node is not None:
            cls._inline_numeric_values(value_node)
            cls._update_direct_chart_values(value_node, values)
        cls._inline_cached_reference(value_node)

    @classmethod
    def _update_chart_title(cls, root, title: str | None):
        if not title:
            return
        for text_node in root.findall('.//c:title//a:t', cls.CHART_NS):
            text_node.text = title

    @classmethod
    def _detach_chart_external_data(cls, root):
        for parent in root.iter():
            for child in list(parent):
                if child.tag == f'{{{cls.CHART_NS["c"]}}}externalData':
                    parent.remove(child)

    @classmethod
    def _sanitize_chart_root(cls, root):
        cls._detach_chart_external_data(root)
        for series in root.findall('.//c:ser', cls.CHART_NS):
            cls._inline_series_title(series)
            for tag_name in ('cat', 'xVal', 'val', 'yVal', 'bubbleSize'):
                cls._inline_cached_reference(series.find(f'c:{tag_name}', cls.CHART_NS))

    @classmethod
    def _sanitize_chart_xml_bytes(cls, chart_bytes: bytes) -> bytes:
        root = ET.fromstring(chart_bytes)
        cls._sanitize_chart_root(root)
        return ET.tostring(root, encoding='utf-8', xml_declaration=True)

    @classmethod
    def _wisc4_chart_payload(cls, test: dict | None):
        payload = cls._wisc_payload(test)
        labels, values = [], []
        index_labels = {"icv": "ICV", "iop": "IOP", "imt": "IMO", "ivp": "IVP"}
        for code in ("icv", "iop", "imt", "ivp"):
            item = next((entry for entry in payload.get("indices") or [] if entry.get("indice") == code), None)
            if not item:
                continue
            labels.append(index_labels[code])
            values.append(cls._to_float(item.get("escore_composto")))
        qit = payload.get("qit_data") or {}
        if qit.get("escore_composto"):
            labels.append("QI Total")
            values.append(cls._to_float(qit.get("escore_composto")))
        for key, label in (("gai_data", "GAI"), ("cpi_data", "CPI")):
            item = payload.get(key) or {}
            if item.get("escore_composto"):
                labels.append(label)
                values.append(cls._to_float(item.get("escore_composto")))
        return labels, values

    @classmethod
    def _wasi_chart_payload(cls, test: dict | None):
        if not test:
            return [], []
        payload = cls._wasi_payload(test)
        labels, values = [], []
        composites = payload.get("composites") or {}
        index_labels = {
            "qi_execucao": "QIE",
            "qi_verbal": "QIV",
            "qit_4": "QI TOTAL",
        }
        for code, label in index_labels.items():
            composite = composites.get(code) or {}
            qi = composite.get("qi")
            if qi is not None:
                labels.append(label)
                values.append(float(qi))
        return labels, values

    @classmethod
    def _bpa_excel_series(cls, test: dict | None):
        chart_data = (test or {}).get("bpa_chart_data") or {}
        domains = {item.get("label"): item.get("values") or {} for item in chart_data.get("domains") or []}
        categories = ["Escore Máximo", "Escore Médio", "Escore Minímo", "Escore Bruto", "Percentil Obtido"]
        order = [
            "ATENÇÃO CONCENTRADA",
            "ATENÇÃO DIVIDIDA",
            "ATENÇÃO ALTERNADA",
            "ATENÇÃO GERAL",
        ]
        key_order = ["maximo", "medio", "minimo", "bruto", "percentil"]
        series = []
        for label in order:
            values = domains.get(label) or {}
            series.append([cls._to_float(values.get(key, 0)) for key in key_order])
        return categories, series

    @classmethod
    def _ravlt_chart_payload(cls, test: dict | None, context: dict | None = None):
        rows = cls._ravlt_rows(test, context)
        if not rows or len(rows) < 4 or not rows[0] or not isinstance(rows[0], list):
            return [], []
        categories = rows[0][1:]
        series_values = []
        for row in rows[1:4]:
            if not row or not isinstance(row, list):
                continue
            series_values.append([cls._to_float(value) for value in row[1:] or []])
        return categories, series_values

    @classmethod
    def _fdt_chart_payload(cls, test: dict | None, automatic: bool):
        payload = (test or {}).get("classified_payload") or {}
        categories = [
            "Tempo Médio",
            "Tempo Obtido",
            "Erros",
            "Desempenho %",
            "Indicativo de Déficit",
            "Indicativo de Dificuldade Discreta",
            "Sem Indicativo de Déficit",
        ]
        metric_map = {
            item.get("codigo"): item for item in payload.get("metric_results") or [] if item.get("codigo")
        }
        errors = payload.get("erros") or {}
        order = (
            ("contagem", "leitura")
            if automatic
            else ("escolha", "alternancia", "inibicao", "flexibilidade")
        )
        series_values = []
        for code in order:
            metric = metric_map.get(code)
            if not metric:
                continue
            performance = cls._fdt_table_result(metric.get("percentil_num")).get("desempenho") or 0
            error_count = (errors.get(code) or {}).get("qtde_erros")
            series_values.append(
                [
                    cls._to_float(metric.get("media")),
                    cls._to_float(metric.get("valor")),
                    cls._to_float(error_count if error_count is not None else 0),
                    cls._to_float(performance),
                    5.0,
                    25.0,
                    50.0,
                ]
            )
        return categories, series_values

    @classmethod
    def _etdah_pais_chart_values(cls, test: dict | None):
        results = ((test or {}).get("classified_payload") or {}).get("results") or {}
        return [
            cls._extract_percentile_value((results.get(key) or {}).get("percentile_text") or (results.get(key) or {}).get("percentil"))
            for key in ("fator_1", "fator_2", "fator_3", "fator_4", "escore_geral")
        ]

    @classmethod
    def _etdah_ad_chart_values(cls, test: dict | None):
        results = ((test or {}).get("classified_payload") or {}).get("results") or {}
        return [
            cls._extract_percentile_value((results.get(key) or {}).get("percentile_text") or (results.get(key) or {}).get("percentil"))
            for key in ("D", "I", "AE", "AAMA", "H")
        ]

    @classmethod
    def _scared_self_chart_values(cls, test: dict | None):
        rows = ((test or {}).get("classified_payload") or {}).get("analise_geral") or []
        return [float(item.get("percentil") or item.get("percentual") or 0) for item in rows]

    @classmethod
    def _scared_parent_chart_series(cls, test: dict | None):
        rows = ((test or {}).get("classified_payload") or {}).get("analise_geral") or []
        lower = []
        cutoff = []
        upper = []
        patient = []
        for item in rows:
            max_percent = float(item.get("nota_corte") or 0)
            lower.append(round(max_percent * (2 / 3), 2))
            cutoff.append(max_percent)
            upper.append(round(max_percent * (4 / 3), 2))
            patient.append(float(item.get("percentual") or 0))
        return [lower, cutoff, upper, patient]

    @classmethod
    def _srs2_chart_values(cls, test: dict | None):
        rows = ((test or {}).get("classified_payload") or {}).get("resultados") or []
        return [float(item.get("tscore") or 0) for item in rows]

    @classmethod
    def _document_chart_targets(cls, docx_bytes: bytes) -> list[str]:
        with ZipFile(BytesIO(docx_bytes), "r") as source_zip:
            document_root = ET.fromstring(source_zip.read("word/document.xml"))
            rels_root = ET.fromstring(source_zip.read("word/_rels/document.xml.rels"))

        relationship_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
        relationship_type = (
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart"
        )
        rel_map = {
            rel.get("Id"): f"word/{rel.get('Target')}"
            for rel in rels_root.findall(f"{{{relationship_ns}}}Relationship")
            if rel.get("Type") == relationship_type and rel.get("Id") and rel.get("Target")
        }

        chart_targets: list[str] = []
        for chart in document_root.findall('.//c:chart', cls.CHART_NS):
            rel_id = chart.get(f'{{{cls.CHART_NS["r"]}}}id')
            target = rel_map.get(rel_id)
            if target:
                chart_targets.append(target)
        return chart_targets

    @classmethod
    def _document_chart_reference_is_valid(
        cls,
        document_root,
        rels_root,
        package_names: set[str],
    ) -> bool:
        relationship_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
        relationship_type = (
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart"
        )
        rel_map = {
            rel.get("Id"): f"word/{rel.get('Target')}"
            for rel in rels_root.findall(f"{{{relationship_ns}}}Relationship")
            if rel.get("Type") == relationship_type and rel.get("Id") and rel.get("Target")
        }
        for chart in document_root.findall('.//c:chart', cls.CHART_NS):
            rel_id = chart.get(f'{{{cls.CHART_NS["r"]}}}id')
            if not rel_id:
                return False
            target = rel_map.get(rel_id)
            if not target or target not in package_names:
                return False
        return True

    @classmethod
    def _resolve_package_target(cls, source_part: str, target: str) -> str:
        source_path = PurePosixPath(source_part)
        if target.startswith('/'):
            return target.lstrip('/')
        resolved = posixpath.normpath((source_path.parent / target).as_posix())
        return resolved.lstrip('./')

    @classmethod
    def _docx_package_is_valid(cls, docx_bytes: bytes) -> bool:
        relationship_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
        content_type_ns = "http://schemas.openxmlformats.org/package/2006/content-types"

        try:
            with ZipFile(BytesIO(docx_bytes), 'r') as source_zip:
                names = set(source_zip.namelist())
                if source_zip.testzip() is not None:
                    return False

                parsed_parts: dict[str, ET.Element] = {}
                for name in names:
                    if not (name.endswith('.xml') or name.endswith('.rels')):
                        continue
                    parsed_parts[name] = ET.fromstring(source_zip.read(name))

                for rel_name, rels_root in parsed_parts.items():
                    if not rel_name.endswith('.rels'):
                        continue
                    if rel_name == '_rels/.rels':
                        source_part = ''
                    else:
                        source_part = rel_name.replace('_rels/', '').replace('.rels', '')
                    for rel in rels_root.findall(f'{{{relationship_ns}}}Relationship'):
                        target = rel.get('Target')
                        if not target or rel.get('TargetMode') == 'External':
                            continue
                        resolved = cls._resolve_package_target(source_part, target)
                        if resolved not in names:
                            return False

                document_root = parsed_parts.get('word/document.xml')
                document_rels_root = parsed_parts.get('word/_rels/document.xml.rels')
                if document_root is None or document_rels_root is None:
                    return False
                if not cls._document_chart_reference_is_valid(
                    document_root,
                    document_rels_root,
                    names,
                ):
                    return False

                content_types = parsed_parts.get('[Content_Types].xml')
                if content_types is None:
                    return False
                defaults = {
                    item.get('Extension')
                    for item in content_types.findall(f'{{{content_type_ns}}}Default')
                    if item.get('Extension')
                }
                overrides = {
                    (item.get('PartName') or '').lstrip('/')
                    for item in content_types.findall(f'{{{content_type_ns}}}Override')
                    if item.get('PartName')
                }
                for name in names:
                    if name == '[Content_Types].xml' or name.endswith('/'):
                        continue
                    extension = name.rsplit('.', 1)[-1] if '.' in name else ''
                    if extension not in defaults and name not in overrides:
                        return False
        except Exception:
            return False

        return True

    @classmethod
    def _prune_unused_chart_parts(cls, docx_bytes: bytes) -> bytes:
        used_chart_parts = set(cls._document_chart_targets(docx_bytes))
        if not used_chart_parts:
            return docx_bytes

        with ZipFile(BytesIO(docx_bytes), 'r') as source_zip:
            names = set(source_zip.namelist())
            used_support_parts: set[str] = set()
            document_rels_root = ET.fromstring(source_zip.read('word/_rels/document.xml.rels'))
            for chart_part in used_chart_parts:
                rels_part = f"word/charts/_rels/{PurePosixPath(chart_part).name}.rels"
                if rels_part not in names:
                    continue
                rels_root = ET.fromstring(source_zip.read(rels_part))
                for rel in rels_root.findall('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                    target = rel.get('Target')
                    if not target or rel.get('TargetMode') == 'External':
                        continue
                    used_support_parts.add(cls._resolve_package_target(chart_part, target))

            content_types_root = ET.fromstring(source_zip.read('[Content_Types].xml'))
            kept_parts = set(names)
            relationship_ns = '{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'
            chart_relationship_type = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart'
            for rel in list(document_rels_root.findall(relationship_ns)):
                if rel.get('Type') != chart_relationship_type:
                    continue
                target = rel.get('Target')
                if not target:
                    continue
                part_name = cls._resolve_package_target('word/document.xml', target)
                if part_name not in used_chart_parts:
                    document_rels_root.remove(rel)
            for name in list(names):
                if re.fullmatch(r'word/charts/chart\d+\.xml', name) and name not in used_chart_parts:
                    kept_parts.discard(name)
                    kept_parts.discard(f"word/charts/_rels/{PurePosixPath(name).name}.rels")
                elif re.fullmatch(r'word/charts/(style|colors)\d+\.xml', name) and name not in used_support_parts:
                    kept_parts.discard(name)
                elif re.fullmatch(r'word/drawings/drawing\d+\.xml', name) and name not in used_support_parts:
                    kept_parts.discard(name)
                elif re.fullmatch(r'word/drawings/_rels/drawing\d+\.xml\.rels', name):
                    drawing_part = name.replace('/_rels/', '/').replace('.rels', '')
                    if drawing_part not in kept_parts:
                        kept_parts.discard(name)

            for override in list(content_types_root):
                part_name = (override.get('PartName') or '').lstrip('/')
                if part_name and part_name not in kept_parts:
                    content_types_root.remove(override)

            output_buffer = BytesIO()
            with ZipFile(output_buffer, 'w') as target_zip:
                for item in source_zip.infolist():
                    if item.filename not in kept_parts and item.filename != '[Content_Types].xml':
                        continue
                    if item.filename == '[Content_Types].xml':
                        data = ET.tostring(content_types_root, encoding='utf-8', xml_declaration=True)
                    elif item.filename == 'word/_rels/document.xml.rels':
                        data = ET.tostring(document_rels_root, encoding='utf-8', xml_declaration=True)
                    else:
                        data = source_zip.read(item.filename)
                    target_zip.writestr(item, data)
        return output_buffer.getvalue()

    @classmethod
    def _populate_wisc4_excel_charts(cls, docx_bytes: bytes, context: dict) -> bytes:
        replacements: dict[str, bytes] = {}
        chart_targets = iter(cls._document_chart_targets(docx_bytes))

        def load_chart(name: str):
            with ZipFile(BytesIO(docx_bytes), 'r') as source_zip:
                return ET.fromstring(source_zip.read(name))

        def dump_chart(name: str, root):
            replacements[name] = ET.tostring(root, encoding='utf-8', xml_declaration=True)

        def next_chart_target() -> str | None:
            return next(chart_targets, None)

        wisc_test = cls._find_test(context, 'wisc4')
        if wisc_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, values = cls._wisc4_chart_payload(wisc_test)
            cls._update_chart_series(root, 0, categories, values)
            dump_chart(chart_name, root)

        bpa_test = cls._find_test(context, 'bpa2')
        if bpa_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, series_values = cls._bpa_excel_series(bpa_test)
            for idx, values in enumerate(series_values):
                cls._update_chart_series(root, idx, categories, values)
            dump_chart(chart_name, root)

        ravlt_test = cls._find_test(context, 'ravlt')
        if ravlt_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, series_values = cls._ravlt_chart_payload(ravlt_test, context)
            for idx, values in enumerate(series_values):
                cls._update_chart_series(root, idx, categories, values)
            dump_chart(chart_name, root)

        fdt_test = cls._find_test(context, 'fdt')
        if fdt_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, series_values = cls._fdt_chart_payload(fdt_test, automatic=True)
            for idx, values in enumerate(series_values):
                cls._update_chart_series(root, idx, categories, values)
            dump_chart(chart_name, root)

            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, series_values = cls._fdt_chart_payload(fdt_test, automatic=False)
            for idx, values in enumerate(series_values):
                cls._update_chart_series(root, idx, categories, values)
            dump_chart(chart_name, root)

        etdah_pais_test = cls._find_test(context, 'etdah_pais')
        if etdah_pais_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            cls._update_chart_series(root, 0, ['F1 - R.E.', 'F2 - H.I.', 'F3 - C.A.', 'F4 - At.', 'TOTAL'], cls._etdah_pais_chart_values(etdah_pais_test))
            dump_chart(chart_name, root)

        etdah_ad_test = cls._find_test(context, 'etdah_ad')
        if etdah_ad_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            cls._update_chart_series(root, 0, ['F1 - D', 'F2 - I', 'F3 - AE', 'F4 - AMAA', 'F5 - H'], cls._etdah_ad_chart_values(etdah_ad_test))
            dump_chart(chart_name, root)

        scared_tests = cls._find_tests(context, 'scared')
        scared_self = next((item for item in scared_tests if (((item.get('classified_payload') or {}).get('form_type') or (item.get('raw_payload') or {}).get('form')) != 'parent')), None)
        scared_parent = next((item for item in scared_tests if (((item.get('classified_payload') or {}).get('form_type') or (item.get('raw_payload') or {}).get('form')) == 'parent')), None)
        if scared_self:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            cls._update_chart_series(root, 0, ['Panic / S.S.', 'AG', 'AS', 'FS', 'EE', 'TOTAL'], cls._scared_self_chart_values(scared_self))
            dump_chart(chart_name, root)
        elif scared_tests:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            cls._update_chart_series(root, 0, ['Panic / S.S.', 'AG', 'AS', 'FS', 'EE', 'TOTAL'], [0, 0, 0, 0, 0, 0])
            dump_chart(chart_name, root)
        if scared_parent:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories = ['Pa/SS', 'AG', 'AS', 'FS', 'EE', 'TOTAL']
            for idx, values in enumerate(cls._scared_parent_chart_series(scared_parent)):
                cls._update_chart_series(root, idx, categories, values)
            dump_chart(chart_name, root)
        elif scared_tests:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories = ['Pa/SS', 'AG', 'AS', 'FS', 'EE', 'TOTAL']
            for idx in range(4):
                cls._update_chart_series(root, idx, categories, [0, 0, 0, 0, 0, 0])
            dump_chart(chart_name, root)

        srs2_test = cls._find_test(context, 'srs2')
        if srs2_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories = ['Perc.S', 'Cog.S', 'Com.S', 'Mot.S', 'PRR', 'CIS', 'TOTAL']
            cls._update_chart_series(root, 0, categories, cls._srs2_chart_values(srs2_test))
            dump_chart(chart_name, root)

        source_buffer = BytesIO(docx_bytes)
        output_buffer = BytesIO()
        with ZipFile(source_buffer, 'r') as source_zip, ZipFile(output_buffer, 'w') as target_zip:
            for item in source_zip.infolist():
                data = replacements.get(item.filename, source_zip.read(item.filename))
                if re.fullmatch(r'word/charts/chart\d+\.xml', item.filename):
                    data = cls._sanitize_chart_xml_bytes(data)
                if item.filename.startswith('word/charts/_rels/chart') and item.filename.endswith('.rels'):
                    data = cls._strip_external_chart_relationships(data)
                target_zip.writestr(item, data)
        return cls._prune_unused_chart_parts(output_buffer.getvalue())

    @classmethod
    def _wais3_chart_payload(cls, test: dict | None):
        payload = cls._wais3_payload(test)
        indices = payload.get("indices") or {}
        items = [
            ("ICV", indices.get("compreensao_verbal")),
            ("IOP", indices.get("organizacao_perceptual")),
            ("IMO", indices.get("memoria_operacional")),
            ("IVP", indices.get("velocidade_processamento")),
            ("QIV", indices.get("qi_verbal")),
            ("QIE", indices.get("qi_execucao")),
            ("QIT", indices.get("qi_total")),
        ]
        labels = []
        values = []
        for label, item in items:
            score = (item or {}).get("pontuacao_composta")
            if score is None:
                continue
            labels.append(label)
            values.append(float(score))
        return labels, values

    @classmethod
    def _populate_wais3_excel_charts(cls, docx_bytes: bytes, context: dict) -> bytes:
        replacements: dict[str, bytes] = {}
        chart_targets = iter(cls._document_chart_targets(docx_bytes))

        def load_chart(name: str):
            with ZipFile(BytesIO(docx_bytes), 'r') as source_zip:
                return ET.fromstring(source_zip.read(name))

        def dump_chart(name: str, root):
            replacements[name] = ET.tostring(root, encoding='utf-8', xml_declaration=True)

        def next_chart_target() -> str | None:
            return next(chart_targets, None)

        wais3_test = cls._find_test(context, 'wais3')
        if wais3_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, values = cls._wais3_chart_payload(wais3_test)
            cls._update_chart_series(root, 0, categories, values)
            dump_chart(chart_name, root)

        wasi_test = cls._find_test(context, 'wasi')
        if wasi_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, values = cls._wasi_chart_payload(wasi_test)
            cls._update_chart_series(root, 0, categories, values)
            dump_chart(chart_name, root)

        bpa_test = cls._find_test(context, 'bpa2')
        if bpa_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, series_values = cls._bpa_excel_series(bpa_test)
            for idx, values in enumerate(series_values):
                cls._update_chart_series(root, idx, categories, values)
            dump_chart(chart_name, root)

        ravlt_test = cls._find_test(context, 'ravlt')
        if ravlt_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, series_values = cls._ravlt_chart_payload(ravlt_test, context)
            for idx, values in enumerate(series_values):
                cls._update_chart_series(root, idx, categories, values)
            dump_chart(chart_name, root)

        fdt_test = cls._find_test(context, 'fdt')
        if fdt_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, series_values = cls._fdt_chart_payload(fdt_test, automatic=True)
            for idx, values in enumerate(series_values):
                cls._update_chart_series(root, idx, categories, values)
            dump_chart(chart_name, root)

            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, series_values = cls._fdt_chart_payload(fdt_test, automatic=False)
            for idx, values in enumerate(series_values):
                cls._update_chart_series(root, idx, categories, values)
            dump_chart(chart_name, root)

        etdah_ad_test = cls._find_test(context, 'etdah_ad')
        if etdah_ad_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            cls._update_chart_series(root, 0, ['F1 - D', 'F2 - I', 'F3 - AE', 'F4 - AMAA', 'F5 - H'], cls._etdah_ad_chart_values(etdah_ad_test))
            dump_chart(chart_name, root)

        scared_tests = cls._find_tests(context, 'scared')
        scared_self = next((item for item in scared_tests if (((item.get('classified_payload') or {}).get('form_type') or (item.get('raw_payload') or {}).get('form')) != 'parent')), None)
        if scared_self:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            cls._update_chart_series(root, 0, ['Panic / S.S.', 'AG', 'AS', 'FS', 'EE', 'TOTAL'], cls._scared_self_chart_values(scared_self))
            dump_chart(chart_name, root)
        elif scared_tests:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            cls._update_chart_series(root, 0, ['Panic / S.S.', 'AG', 'AS', 'FS', 'EE', 'TOTAL'], [0, 0, 0, 0, 0, 0])
            dump_chart(chart_name, root)

        srs2_test = cls._find_test(context, 'srs2')
        if srs2_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories = ['Perc.S', 'Cog.S', 'Com.S', 'Mot.S', 'PRR', 'CIS', 'TOTAL']
            cls._update_chart_series(root, 0, categories, cls._srs2_chart_values(srs2_test))
            dump_chart(chart_name, root)

        source_buffer = BytesIO(docx_bytes)
        output_buffer = BytesIO()
        with ZipFile(source_buffer, 'r') as source_zip, ZipFile(output_buffer, 'w') as target_zip:
            for item in source_zip.infolist():
                data = replacements.get(item.filename, source_zip.read(item.filename))
                if re.fullmatch(r'word/charts/chart\d+\.xml', item.filename):
                    data = cls._sanitize_chart_xml_bytes(data)
                if item.filename.startswith('word/charts/_rels/chart') and item.filename.endswith('.rels'):
                    data = cls._strip_external_chart_relationships(data)
                target_zip.writestr(item, data)
        return cls._prune_unused_chart_parts(output_buffer.getvalue())

    @classmethod
    def _populate_wasi_excel_charts(cls, docx_bytes: bytes, context: dict) -> bytes:
        replacements: dict[str, bytes] = {}
        chart_targets = list(cls._document_chart_targets(docx_bytes))
        chart_targets_iter = iter(chart_targets)

        def load_chart(name: str):
            with ZipFile(BytesIO(docx_bytes), 'r') as source_zip:
                return ET.fromstring(source_zip.read(name))

        def dump_chart(name: str, root):
            replacements[name] = ET.tostring(root, encoding='utf-8', xml_declaration=True)

        def next_chart_target() -> str | None:
            return next(chart_targets_iter, None)

        wasi_test = cls._find_test(context, 'wasi')
        if wasi_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, values = cls._wasi_chart_payload(wasi_test)
            cls._update_chart_series(root, 0, categories, values)
            dump_chart(chart_name, root)

        bpa_test = cls._find_test(context, 'bpa2')
        if bpa_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, series_values = cls._bpa_excel_series(bpa_test)
            for idx, values in enumerate(series_values):
                cls._update_chart_series(root, idx, categories, values)
            dump_chart(chart_name, root)

        ravlt_test = cls._find_test(context, 'ravlt')
        if ravlt_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, series_values = cls._ravlt_chart_payload(ravlt_test, context)
            for idx, values in enumerate(series_values):
                cls._update_chart_series(root, idx, categories, values)
            dump_chart(chart_name, root)

        fdt_test = cls._find_test(context, 'fdt')
        if fdt_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, series_values = cls._fdt_chart_payload(fdt_test, automatic=True)
            for idx, values in enumerate(series_values):
                cls._update_chart_series(root, idx, categories, values)
            dump_chart(chart_name, root)

            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories, series_values = cls._fdt_chart_payload(fdt_test, automatic=False)
            for idx, values in enumerate(series_values):
                cls._update_chart_series(root, idx, categories, values)
            dump_chart(chart_name, root)

        etdah_ad_test = cls._find_test(context, 'etdah_ad')
        if etdah_ad_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            cls._update_chart_series(
                root,
                0,
                ['F1 - D', 'F2 - I', 'F3 - AE', 'F4 - AMAA', 'F5 - H'],
                cls._etdah_ad_chart_values(etdah_ad_test),
            )
            dump_chart(chart_name, root)

        srs2_test = cls._find_test(context, 'srs2')
        if srs2_test:
            chart_name = next_chart_target()
            if chart_name is None:
                return docx_bytes
            root = load_chart(chart_name)
            categories = ['Perc.S', 'Cog.S', 'Com.S', 'Mot.S', 'PRR', 'CIS', 'TOTAL']
            cls._update_chart_series(root, 0, categories, cls._srs2_chart_values(srs2_test))
            dump_chart(chart_name, root)

        cls._zero_unused_wasi_charts(replacements, chart_targets, load_chart, dump_chart)

        source_buffer = BytesIO(docx_bytes)
        output_buffer = BytesIO()
        with ZipFile(source_buffer, 'r') as source_zip, ZipFile(output_buffer, 'w') as target_zip:
            for item in source_zip.infolist():
                data = replacements.get(item.filename, source_zip.read(item.filename))
                if re.fullmatch(r'word/charts/chart\d+\.xml', item.filename):
                    data = cls._sanitize_chart_xml_bytes(data)
                if item.filename.startswith('word/charts/_rels/chart') and item.filename.endswith('.rels'):
                    data = cls._strip_external_chart_relationships(data)
                target_zip.writestr(item, data)
        return output_buffer.getvalue()

    @classmethod
    def _zero_unused_wasi_charts(cls, replacements: dict, all_chart_targets: list, load_chart, dump_chart):
        used_targets = set(k for k in replacements.keys() if re.match(r'word/charts/chart\d+\.xml', k))
        for target in all_chart_targets:
            if target not in used_targets:
                try:
                    root = load_chart(target)
                    cls._zero_chart_series(root)
                    dump_chart(target, root)
                except Exception:
                    pass

    @classmethod
    def _build_fallback_document(cls, report: Report, context: dict):
        document = Document()
        cls._ensure_model_table_styles(document)
        cls._apply_base_styles(document)

        patient = context.get("patient") or {}
        author_name = cls.FIXED_AUTHOR
        interested_party = cls.FIXED_INTERESTED_PARTY
        purpose = cls.FIXED_PURPOSE

        cls._add_center_title(document, report.title or "Laudo Neuropsicológico")
        cls._add_center_text(document, "Documento gerado sem papel timbrado do template.")

        cls._append_heading(document, "1. IDENTIFICAÇÃO")
        cls._append_label_value(document, "Autora", author_name)
        cls._append_label_value(document, "Interessado", interested_party)
        cls._append_label_value(document, "Finalidade", purpose)
        cls._append_label_value(
            document, "Paciente", patient.get("full_name") or "Não informado"
        )
        cls._append_label_value(
            document,
            "Data de nascimento",
            cls._format_date_display(patient.get("birth_date")),
        )
        cls._append_label_value(document, "Idade", cls._age_text(context))

        raw_body_text = str(report.final_text or report.edited_text or report.generated_text or "")
        body_text = cls._sanitize_section_text_for_patient(raw_body_text, context)
        if body_text.strip():
            for raw_line in body_text.splitlines():
                line = raw_line.strip()
                if not line:
                    continue
                if line.startswith("### "):
                    cls._append_subheading(document, line[4:].strip())
                    continue
                if line.startswith("## "):
                    cls._append_heading(document, line[3:].strip())
                    continue
                cls._append_paragraph(document, line)
        else:
            appended_section = False
            for section in report.sections.all():
                raw_section_text = str(section.content_edited or section.content_generated or "")
                section_text = cls._sanitize_section_text_for_patient(raw_section_text, context)
                if not section_text:
                    continue
                appended_section = True
                cls._append_heading(document, section.title or section.key.replace("_", " ").title())
                for raw_line in section_text.splitlines():
                    line = raw_line.strip()
                    if not line:
                        continue
                    cls._append_paragraph(document, line)

            if not appended_section:
                cls._append_paragraph(
                    document,
                    "O laudo ainda não possui conteúdo textual consolidado para exportação.",
                )

        return document

    @classmethod
    def _template_has_editable_body(cls, document: Document) -> bool:
        if document.tables:
            return True
        return any(paragraph.text.strip() for paragraph in document.paragraphs)

    @classmethod
    def _is_adolescent_document(cls, document: Document, context: dict) -> bool:
        if any(
            paragraph.text.strip() == "Interessado Familiares"
            for paragraph in document.paragraphs
        ):
            return True
        birth_date = (context.get("patient") or {}).get("birth_date") or ""
        if birth_date:
            try:
                from datetime import date

                year, month, day = map(int, birth_date.split("-"))
                born = date(year, month, day)
                today = date.today()
                age = (
                    today.year
                    - born.year
                    - ((today.month, today.day) < (born.month, born.day))
                )
                return age < 18
            except ValueError:
                return False
        return False

    @classmethod
    def _is_adolescent_context(cls, context: dict, report) -> bool:
        return False

    @classmethod
    def _build_adolescent_document(
        cls, report, context: dict, sections: dict[str, str]
    ):
        wisc_test = cls._find_test(context, "wisc4")
        template_path = (
            cls.WISC4_TEMPLATE_PATH
            if wisc_test and cls.WISC4_TEMPLATE_PATH.exists()
            else cls._select_template_path(report, context)
        )
        document = (
            Document(str(template_path)) if template_path.exists() else Document()
        )
        template_chart_blocks = cls._extract_template_chart_blocks(document)
        cls._ensure_model_table_styles(document)
        cls._clear_document_body(document)
        cls._apply_base_styles(document)
        patient = context.get("patient") or {}
        evaluation = context.get("evaluation") or {}
        author_name = cls.FIXED_AUTHOR
        interested_party = cls.FIXED_INTERESTED_PARTY
        table_index = 1
        chart_index = 1
        primary_test = cls._primary_report_test_code(context)
        template_chart_map = {}
        if primary_test == "wais3" and cls.WAIS3_TEMPLATE_PATH.exists():
            wais3_template_blocks = cls._extract_chart_blocks_from_template_path(cls.WAIS3_TEMPLATE_PATH)
            template_chart_map = {
                "wais3": wais3_template_blocks[0] if len(wais3_template_blocks) > 0 else None,
                "bpa2": template_chart_blocks[1] if len(template_chart_blocks) > 1 else None,
                "ravlt": template_chart_blocks[2] if len(template_chart_blocks) > 2 else None,
                "fdt_auto": template_chart_blocks[3] if len(template_chart_blocks) > 3 else None,
                "fdt_control": template_chart_blocks[4] if len(template_chart_blocks) > 4 else None,
                "etdah_ad": template_chart_blocks[5] if len(template_chart_blocks) > 5 else None,
                "scared_pair": template_chart_blocks[7] if len(template_chart_blocks) > 7 else None,
                "srs2": template_chart_blocks[6] if len(template_chart_blocks) > 6 else None,
            }
        elif primary_test == "wasi" and cls.WASI_TEMPLATE_PATH.exists():
            wasi_template_blocks = cls._extract_chart_blocks_from_template_path(cls.WASI_TEMPLATE_PATH)
            template_chart_map = {
                "wasi": wasi_template_blocks[0] if len(wasi_template_blocks) > 0 else None,
                "bpa2": template_chart_blocks[1] if len(template_chart_blocks) > 1 else None,
                "ravlt": template_chart_blocks[2] if len(template_chart_blocks) > 2 else None,
            }
        else:
            template_chart_map = {
                "wisc4": template_chart_blocks[0] if len(template_chart_blocks) > 0 else None,
                "bpa2": template_chart_blocks[1] if len(template_chart_blocks) > 1 else None,
                "ravlt": template_chart_blocks[2] if len(template_chart_blocks) > 2 else None,
                "fdt_auto": template_chart_blocks[3] if len(template_chart_blocks) > 3 else None,
                "fdt_control": template_chart_blocks[4] if len(template_chart_blocks) > 4 else None,
                "etdah_pais": template_chart_blocks[5] if len(template_chart_blocks) > 5 else None,
                "etdah_ad": template_chart_blocks[6] if len(template_chart_blocks) > 6 else None,
                "scared_pair": template_chart_blocks[7] if len(template_chart_blocks) > 7 else None,
                "srs2": template_chart_blocks[9] if len(template_chart_blocks) > 9 else None,
                "wasi": template_chart_blocks[0] if template_chart_blocks else None,
            }

        def append_table_with_interpretation(
            rows,
            table_key: str,
            interpretation: str | None,
            caption: str,
        ):
            nonlocal table_index
            cls._append_table_with_interpretation(
                document,
                rows,
                table_key,
                interpretation,
                caption=cls._table_caption_text(table_index, caption),
            )
            if rows:
                table_index += 1

        def append_chart(
            caption: str,
            image_bytes: bytes | None,
            note: str | None = None,
            width=None,
            height=None,
            show_caption: bool = True,
            template_key: str | None = None,
        ):
            nonlocal chart_index
            if template_key and template_chart_map.get(template_key) is not None:
                cls._append_body_element_before_sectpr(document, template_chart_map[template_key])
                if show_caption:
                    p = document.add_paragraph()
                    r = p.add_run(cls._chart_caption_text(chart_index, caption))
                    r.font.name = cls.FONT_NAME
                    r.font.size = cls.BODY_SIZE
                    cls._format_caption_paragraph(p)
                if note:
                    note_paragraph = document.add_paragraph(note)
                    cls._format_chart_legend_paragraph(note_paragraph)
                chart_index += 1
                return
            if not image_bytes:
                return
            cls._append_chart(
                document,
                cls._chart_caption_text(chart_index, caption) if show_caption else None,
                image_bytes,
                width=width,
                height=height,
            )
            if note:
                note_paragraph = document.add_paragraph(note)
                cls._format_chart_legend_paragraph(note_paragraph)
            chart_index += 1

        def section_or_test_interpretation(
            section_key: str,
            fallback_section_key: str | None,
            test_payload: dict | None,
        ) -> str:
            return cls._resolve_interpretation_text(
                sections.get(section_key),
                sections.get(fallback_section_key) if fallback_section_key else None,
                test_payload,
                context,
            )

        purpose = cls.FIXED_PURPOSE

        cls._add_center_title(document, WASI_REPORT_SPEC["title"])
        cls._add_center_text(
            document,
            WASI_REPORT_SPEC["subtitle"],
        )

        cls._append_heading(document, "1. IDENTIFICAÇÃO")
        cls._append_subheading(document, "1.1. Identificação do Laudo")
        cls._append_label_value(document, "Autora", author_name)
        cls._append_label_value(document, "Interessado", interested_party)
        cls._append_label_value(document, "Finalidade", purpose)
        cls._append_subheading(document, "1.2. Identificação do Paciente")
        cls._append_label_value(
            document, "Nome", patient.get("full_name") or "Não informado"
        )
        cls._append_label_value(
            document, "Sexo", cls._format_sex_display(patient.get("sex"))
        )
        cls._append_label_value(
            document,
            "Data de nascimento",
            cls._format_date_display(patient.get("birth_date")),
        )
        cls._append_label_value(document, "Idade", cls._age_text(context))
        cls._append_label_value(document, "Filiação", cls._filiation_text(patient))
        cls._append_label_value(document, "Escolaridade", cls._schooling_text(patient))
        if patient.get("school_name"):
            cls._append_label_value(document, "Escola", patient.get("school_name"))

        cls._append_heading(document, "2. DESCRIÇÃO DA DEMANDA")
        cls._append_subheading(document, "2.1. Motivo do Encaminhamento")
        cls._append_paragraph(
            document,
            sections.get("descricao_demanda")
            or evaluation.get("referral_reason")
            or "Sem conteúdo disponível para esta seção.",
        )

        cls._append_heading(document, "3. PROCEDIMENTOS")
        for item in cls._adolescent_instruments(context):
            cls._append_procedure_bullet(document, item["name"], item["description"])

        cls._append_heading(document, "4. ANÁLISE")
        cls._append_subheading(document, "4.1. História Pessoal")
        cls._append_paragraph(
            document,
            sections.get("historia_pessoal")
            or "Sem conteúdo disponível para esta seção.",
        )

        patient_title = "da paciente" if (context.get("patient") or {}).get("sex") == "F" else "do paciente"
        cls._append_heading(document, "5. ANÁLISE QUALITATIVA")
        wais3_test = cls._find_test(context, "wais3")
        if wisc_test:
            cls._append_subheading(document, "Capacidade Cognitiva Global")
            cls._append_wisc_global_block(document, wisc_test, context)
            cls._append_wisc_indices_block(document, wisc_test)
            cls._append_subheading(document, f"Desempenho {patient_title} no WISC-IV")
            wisc_chart = cls._wisc_chart(wisc_test)
            append_chart(
                "WISC-IV - INDICES DE QI",
                wisc_chart,
                cls._wisc_chart_legend(chart_index),
                show_caption=False,
                template_key="wisc4",
            )
            cls._append_interpretation_block(
                document,
                cls._wisc_model_interpretation_text(wisc_test, context),
            )
            cls._append_subheading(document, "Subescalas WISC-IV")
            cls._append_subheading(document, "Função Executiva")
            append_table_with_interpretation(
                cls._wisc_rows(
                    cls._find_test(context, "wisc4"), context, "funcoes_executivas"
                ),
                "wisc",
                cls._wisc_section_text(sections, "funcoes_executivas", context),
                "Resultados da Função Executiva",
            )
            cls._append_subheading(document, "Linguagem")
            append_table_with_interpretation(
                cls._wisc_rows(cls._find_test(context, "wisc4"), context, "linguagem"),
                "wisc",
                cls._wisc_section_text(sections, "linguagem", context),
                "Resultados da Linguagem",
            )
            cls._append_subheading(document, "Gnosias e Praxias")
            append_table_with_interpretation(
                cls._wisc_rows(
                    cls._find_test(context, "wisc4"), context, "gnosias_praxias"
                ),
                "wisc",
                cls._wisc_section_text(sections, "gnosias_praxias", context),
                "Resultados de Gnosias e Praxias",
            )
            cls._append_subheading(document, "Memória e Aprendizagem")
            append_table_with_interpretation(
                cls._wisc_memory_rows(cls._find_test(context, "wisc4"), context),
                "wisc",
                cls._wisc_section_text(sections, "memoria_aprendizagem", context),
                "Resultados de Memória e Aprendizagem",
            )
        elif wais3_test:
            cls._append_subheading(document, f"5.1. Desempenho {patient_title} no WAIS-III")
            cls._append_paragraph(document, cls._wais3_intro_text(wais3_test, context))
            for lead, tail in cls._wais3_global_bullet_parts(wais3_test):
                p = document.add_paragraph()
                cls._append_wisc_global_bullet(p, lead, tail)
            append_chart(
                "WAIS III - INDICES DE QIS",
                cls._wais3_chart(wais3_test),
                show_caption=True,
                template_key="wais3",
            )

        next_section_number = 6

        def append_section_heading(title: str):
            nonlocal next_section_number
            if wisc_test:
                cls._append_subheading(document, title)
            else:
                cls._append_heading(document, f"{next_section_number}. {title}")
                next_section_number += 1

        bpa2_test = cls._find_test(context, "bpa2")
        if bpa2_test:
            append_section_heading("BPA-2 – BATERIA PSICOLÓGICA PARA AVALIAÇÃO DA ATENÇÃO")
            cls._append_paragraph(
                document,
                "A Bateria Psicológica para Avaliação da Atenção-2 (BPA-2) mensura a capacidade geral de atenção, avaliando individualmente atenção concentrada, dividida, alternada e geral.",
            )
            append_table_with_interpretation(
                cls._bpa_rows(bpa2_test, context),
                "bpa",
                None,
                "Atenção BPA-2 Resultados",
            )
            append_chart(
                "BPA-2 Resultados da Avaliação da Atenção",
                cls._bpa_chart_bytes(bpa2_test),
                template_key="bpa2",
            )
            cls._append_bpa_interpretation_block(document, bpa2_test, context)

        ravlt_test = cls._find_test(context, "ravlt")
        if ravlt_test:
            append_section_heading("RAVLT – REY AUDITORY VERBAL LEARNING TEST")
            ravlt_interpretation = sections.get("ravlt") or sections.get("memoria_aprendizagem")
            cls._append_ravlt_conceptual_paragraph(document)
            append_chart(
                "RAVLT Resultados",
                cls._ravlt_chart(ravlt_test),
                template_key="ravlt",
            )
            append_table_with_interpretation(
                cls._ravlt_rows(ravlt_test, context),
                "ravlt",
                ravlt_interpretation,
                cls._ravlt_table_caption(),
            )

        fdt_test = cls._find_test(context, "fdt")
        if fdt_test:
            append_section_heading("FDT – TESTE DOS CINCO DÍGITOS")
            cls._append_paragraph(document, cls._fdt_description_text())
            append_table_with_interpretation(
                cls._fdt_rows(fdt_test),
                "fdt",
                sections.get("fdt") or sections.get("funcoes_executivas"),
                "FDT Processos Automáticos e Controlados",
            )
            append_chart(
                "FDT Processos Automáticos",
                cls._fdt_chart(fdt_test, automatic=True),
                width=Cm(14),
                height=Cm(10),
                template_key="fdt_auto",
            )
            append_chart(
                "FDT Processos Controlados",
                cls._fdt_chart(fdt_test, automatic=False),
                width=Cm(14),
                height=Cm(10),
                template_key="fdt_control",
            )

        if cls._find_test(context, "etdah_pais"):
            append_section_heading("E-TDAH-PAIS")
            cls._append_paragraph(
                document,
                "A Escala E-TDAH-PAIS tem como objetivo identificar manifestações comportamentais e emocionais associadas ao Transtorno do Déficit de Atenção e Hiperatividade (TDAH) a partir da percepção dos pais ou responsáveis, avaliando domínios relacionados à regulação emocional, impulsividade, comportamento adaptativo e atenção. O instrumento fornece indicadores quantitativos e qualitativos do funcionamento atencional e comportamental, permitindo compreender o impacto desses aspectos no cotidiano da criança (Benczik, 2005).",
            )
            append_table_with_interpretation(
                cls._etdah_rows(cls._find_test(context, "etdah_pais")),
                "etdah_pais",
                None,
                "Resultados do E-TDAH-PAIS",
            )
            cls._append_etdah_pais_interpretation_block(document, cls._find_test(context, "etdah_pais"), context)
            append_chart(
                "E-TDAH-PAIS Percentis",
                cls._etdah_chart(cls._find_test(context, "etdah_pais")),
                template_key="etdah_pais",
            )

        if cls._find_test(context, "etdah_ad"):
            append_section_heading("E-TDAH-AD")
            cls._append_paragraph(
                document,
                "O E-TDAH-AD investiga sintomas relacionados à atenção, hiperatividade, impulsividade e aspectos emocionais a partir do autorrelato do adolescente.",
            )
            append_table_with_interpretation(
                cls._etdah_rows(cls._find_test(context, "etdah_ad")),
                "etdah_ad",
                None,
                "ETDAH-AD RESULTADO",
            )
            cls._append_etdah_ad_interpretation_block(
                document,
                cls._normalize_interpretation_text(
                    section_or_test_interpretation("etdah_ad", None, cls._find_test(context, "etdah_ad"))
                ),
            )
            append_chart(
                "E-TDAH-AD RESULTADO",
                cls._etdah_chart(cls._find_test(context, "etdah_ad")),
                template_key="etdah_ad",
            )

        scared_tests = cls._find_tests(context, "scared")
        if scared_tests:
            append_section_heading("SCARED")
            cls._append_paragraph(
                document,
                "O Screen for Child Anxiety Related Emotional Disorders – SCARED é um instrumento de rastreio destinado à identificação de sintomas ansiosos em crianças e adolescentes, avaliando manifestações relacionadas a pânico, ansiedade generalizada, ansiedade de separação, fobia social e evitação escolar (Birmaher et al., 1999).",
            )
            inserted_scared_pair = False
            for scared_test in scared_tests:
                form_label = cls._scared_form_label(scared_test)
                append_table_with_interpretation(
                    cls._scared_rows(scared_test),
                    cls._scared_table_key(scared_test),
                    cls._resolve_interpretation_text(None, None, scared_test, context),
                    f"SCARED - Resultados {form_label}",
                )
                if not inserted_scared_pair:
                    append_chart(
                        cls._scared_form_title(scared_test),
                        cls._scared_chart(scared_test),
                        template_key="scared_pair",
                    )
                    inserted_scared_pair = True

        if cls._find_test(context, "epq_j"):
            append_section_heading("EPQ-J")
            append_table_with_interpretation(
                cls._epq_rows(cls._find_test(context, "epq_j")),
                "epq",
                section_or_test_interpretation(
                    "epq_j", None, cls._find_test(context, "epq_j")
                ),
                "EPQ-J Resultados da personalidade",
            )
            append_chart(
                "EPQ-J - Percentis",
                cls._epq_chart(cls._find_test(context, "epq_j")),
            )

        if cls._find_test(context, "bai"):
            bai_test = cls._find_test(context, "bai")
            append_section_heading("BAI")
            cls._append_paragraph(document, cls._bai_description_text())
            append_table_with_interpretation(
                cls._bai_rows(bai_test),
                "scale_summary",
                section_or_test_interpretation("bai", None, bai_test),
                "BAI - Resultado da sintomatologia ansiosa",
            )

        if cls._find_test(context, "ebadep_a") or cls._find_test(context, "ebadep_ij") or cls._find_test(context, "ebaped_ij"):
            ebadep_test = cls._find_test(context, "ebadep_a") or cls._find_test(context, "ebadep_ij") or cls._find_test(context, "ebaped_ij")
            append_section_heading("EBADEP-A")
            cls._append_paragraph(document, cls._ebadep_description_text())
            append_table_with_interpretation(
                cls._ebadep_rows(ebadep_test),
                "scale_summary",
                cls._resolve_interpretation_text(
                    sections.get("aspectos_emocionais_comportamentais"),
                    None,
                    ebadep_test,
                ),
                "EBADEP-A - Resultado da sintomatologia",
            )

        if cls._find_test(context, "srs2"):
            srs2_test = cls._find_test(context, "srs2")
            append_section_heading("SRS-2 – ESCALA DE RESPONSIVIDADE SOCIAL")
            cls._append_paragraph(
                document,
                "A SRS-2 avalia aspectos da interação social e comportamentos associados ao espectro autista, exigindo sempre integração clínica cuidadosa.",
            )
            append_table_with_interpretation(
                cls._srs2_rows(srs2_test),
                "srs2",
                section_or_test_interpretation(
                    "srs2", None, srs2_test
                ),
                "SRS-2 Resultados",
            )
            append_chart(
                "SRS-2 Resultados",
                cls._srs2_chart(srs2_test),
                template_key="srs2",
            )

        if cls._find_test(context, "bfp"):
            bfp_test = cls._find_test(context, "bfp")
            append_section_heading("BFP – BATERIA FATORIAL DE PERSONALIDADE")
            cls._append_paragraph(document, cls._bfp_description_text())
            append_table_with_interpretation(
                cls._bfp_rows(bfp_test),
                "bfp",
                section_or_test_interpretation("bfp", "aspectos_emocionais_comportamentais", bfp_test),
                "BFP Resultados dos fatores",
            )
        if wisc_test:
            cls._append_subheading(document, "Conclusão")
        else:
            cls._append_heading(document, f"{next_section_number}. CONCLUSÃO")
        cls._append_paragraph(
            document,
            sections.get("conclusao") or "Sem conteúdo disponível para esta seção.",
        )
        if not wisc_test:
            next_section_number += 1
        if wisc_test:
            cls._append_subheading(document, "Sugestões de Conduta (Encaminhamentos)")
        else:
            cls._append_heading(document, f"{next_section_number}. SUGESTÕES DE CONDUTA (ENCAMINHAMENTOS)")
        for bullet in cls._split_bullets(sections.get("sugestoes_conduta") or ""):
            cls._append_bullet(document, bullet)

        if not wisc_test:
            next_section_number += 1
        if wisc_test:
            cls._append_subheading(document, "Referencia Bibliográfica")
        else:
            cls._append_heading(document, f"{next_section_number}. REFERÊNCIA BIBLIOGRÁFICA")
        for ref in cls._references_list(context):
            cls._append_paragraph(document, ref)
        return document

    @classmethod
    def _apply_base_styles(cls, document: Document):
        cls._apply_page_layout(document)
        for style_name in ("Normal",):
            style = document.styles[style_name]
            style.font.name = cls.FONT_NAME
            style.font.size = cls.BODY_SIZE
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = cls.BODY_SPACE_AFTER
            style.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
            style.paragraph_format.first_line_indent = cls.BODY_FIRST_LINE_INDENT
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    @classmethod
    def _apply_page_layout(cls, document: Document):
        for section in document.sections:
            sect_pr = section._sectPr
            has_header_footer_refs = any(
                child.tag in {qn("w:headerReference"), qn("w:footerReference")}
                for child in sect_pr
            )
            if not has_header_footer_refs:
                section.top_margin = Cm(3)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
                section.header_distance = Cm(1.27)
                section.footer_distance = Cm(1.27)
            pg_mar = sect_pr.find(qn("w:pgMar"))
            if pg_mar is None:
                pg_mar = OxmlElement("w:pgMar")
                sect_pr.append(pg_mar)
            pg_mar.set(qn("w:gutter"), "0")

    @classmethod
    def _normalize_model_header_footer(cls, document: Document):
        for section in document.sections:
            for paragraph in section.footer.paragraphs:
                for run in paragraph.runs:
                    if run.text == "pág. ":
                        run.text = "pág."

    @classmethod
    def _restore_template_header_footer(cls, docx_bytes: bytes, template_path: Path) -> bytes:
        with ZipFile(BytesIO(docx_bytes), "r") as output_zip, ZipFile(str(template_path), "r") as template_zip:
            template_names = set(template_zip.namelist())
            output_names = set(output_zip.namelist())
            template_document_xml = template_zip.read("word/document.xml")
            template_document_rels_xml = template_zip.read("word/_rels/document.xml.rels")
            output_document_xml = output_zip.read("word/document.xml")
            output_document_rels_xml = output_zip.read("word/_rels/document.xml.rels")

            template_document_root = ET.fromstring(template_document_xml)
            output_document_root = ET.fromstring(output_document_xml)
            template_rels_root = ET.fromstring(template_document_rels_xml)
            output_rels_root = ET.fromstring(output_document_rels_xml)

            relationship_ns = "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"
            word_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            office_rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/"
            office_doc_rel_attr = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
            header_footer_types = {
                f"{office_rel_ns}header",
                f"{office_rel_ns}footer",
            }

            for rel in list(output_rels_root.findall(relationship_ns)):
                if rel.get("Type") in header_footer_types:
                    output_rels_root.remove(rel)
            existing_rel_ids = {
                rel.get("Id")
                for rel in output_rels_root.findall(relationship_ns)
                if rel.get("Id")
            }
            rel_id_map: dict[str, str] = {}
            for rel in template_rels_root.findall(relationship_ns):
                if rel.get("Type") in header_footer_types:
                    cloned_rel = deepcopy(rel)
                    rel_id = cloned_rel.get("Id")
                    if rel_id in existing_rel_ids:
                        next_id = 1
                        while f"rId{next_id}" in existing_rel_ids:
                            next_id += 1
                        new_rel_id = f"rId{next_id}"
                        rel_id_map[rel_id] = new_rel_id
                        cloned_rel.set("Id", new_rel_id)
                        rel_id = new_rel_id
                    existing_rel_ids.add(rel_id)
                    output_rels_root.append(cloned_rel)

            output_body = output_document_root.find(f".//{word_ns}body")
            template_body = template_document_root.find(f".//{word_ns}body")
            output_sect_pr = output_body.find(f"{word_ns}sectPr") if output_body is not None else None
            template_sect_pr = template_body.find(f"{word_ns}sectPr") if template_body is not None else None
            if output_body is not None and template_sect_pr is not None:
                if output_sect_pr is not None:
                    output_body.remove(output_sect_pr)
                cloned_sect_pr = deepcopy(template_sect_pr)
                for child in cloned_sect_pr:
                    if child.tag not in {f"{word_ns}headerReference", f"{word_ns}footerReference"}:
                        continue
                    rel_id = child.get(office_doc_rel_attr)
                    if rel_id in rel_id_map:
                        child.set(office_doc_rel_attr, rel_id_map[rel_id])
                output_body.append(cloned_sect_pr)

            parts_to_copy: set[str] = {
                "word/document.xml",
                "word/_rels/document.xml.rels",
            }
            for rel in template_rels_root.findall(relationship_ns):
                if rel.get("Type") not in header_footer_types:
                    continue
                target = rel.get("Target")
                if not target:
                    continue
                part_name = cls._resolve_package_target("word/document.xml", target)
                parts_to_copy.add(part_name)
                rels_name = f"word/_rels/{PurePosixPath(part_name).name}.rels"
                if rels_name in template_names:
                    parts_to_copy.add(rels_name)
                    rels_root = ET.fromstring(template_zip.read(rels_name))
                    for sub_rel in rels_root.findall(relationship_ns):
                        sub_target = sub_rel.get("Target")
                        if not sub_target or sub_rel.get("TargetMode") == "External":
                            continue
                        parts_to_copy.add(cls._resolve_package_target(part_name, sub_target))

            content_types_root = ET.fromstring(output_zip.read("[Content_Types].xml"))
            template_content_types_root = ET.fromstring(template_zip.read("[Content_Types].xml"))
            existing_overrides = {
                override.get("PartName")
                for override in content_types_root
                if override.tag.endswith("Override")
            }
            for override in template_content_types_root:
                part_name = override.get("PartName")
                normalized = (part_name or "").lstrip("/")
                if normalized not in parts_to_copy or part_name in existing_overrides:
                    continue
                content_types_root.append(deepcopy(override))

            output_buffer = BytesIO()
            with ZipFile(output_buffer, "w") as target_zip:
                for item in output_zip.infolist():
                    if item.filename == "word/document.xml":
                        data = ET.tostring(output_document_root, encoding="utf-8", xml_declaration=True)
                    elif item.filename == "word/_rels/document.xml.rels":
                        data = ET.tostring(output_rels_root, encoding="utf-8", xml_declaration=True)
                    elif item.filename == "[Content_Types].xml":
                        data = ET.tostring(content_types_root, encoding="utf-8", xml_declaration=True)
                    elif item.filename in parts_to_copy and item.filename in template_names:
                        data = template_zip.read(item.filename)
                    else:
                        data = output_zip.read(item.filename)
                    target_zip.writestr(item, data)

                missing_parts = parts_to_copy - output_names
                for name in sorted(missing_parts):
                    target_zip.writestr(name, template_zip.read(name))

            return output_buffer.getvalue()

    @classmethod
    def _ensure_model_table_styles(cls, document: Document):
        try:
            document.styles["Table Grid"]
            return
        except KeyError:
            pass

        source_path = cls.TABLE_STYLE_SOURCE_PATH
        if not source_path.exists():
            return

        source_document = Document(str(source_path))
        source_styles = source_document.styles.element
        target_styles = document.styles.element

        existing_ids = {
            style.get(qn("w:styleId"))
            for style in target_styles.findall(qn("w:style"))
        }

        for style_id in ("Tabelanormal", "Tabelacomgrade", "Tabelacomgrade1"):
            if style_id in existing_ids:
                continue
            source_style = source_styles.find(f'.//w:style[@w:styleId="{style_id}"]', namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            if source_style is not None:
                target_styles.append(deepcopy(source_style))

    @classmethod
    def _apply_model_table_style(cls, table):
        for style_name in ("Table Grid", "Tabela com grade1"):
            try:
                table.style = style_name
                return
            except KeyError:
                continue

    @classmethod
    def _replace_simple_sections(cls, document: Document, report, sections: dict[str, str], context: dict):
        replacements = {
            "ANÁLISE": sections.get(
                "historia_pessoal", "Sem conteúdo disponível para esta seção."
            ),
            "Conclusão": sections.get(
                "conclusao", "Sem conteúdo disponível para esta seção."
            ),
            "Sugestões de Conduta (Encaminhamentos):": sections.get(
                "sugestoes_conduta", "Sem conteúdo disponível para esta seção."
            ),
            "Referências Bibliográficas": sections.get(
                "referencias_bibliograficas",
                "Sem conteúdo disponível para esta seção.",
            ),
            "Referencia Bibliográfica": sections.get(
                "referencias_bibliograficas",
                "Sem conteúdo disponível para esta seção.",
            ),
        }
        boundaries = {
            "IDENTIFICAÇÃO": "DESCRIÇÃO DA DEMANDA",
            "DESCRIÇÃO DA DEMANDA": "PROCEDIMENTOS",
            "PROCEDIMENTOS": "ANÁLISE",
            "ANÁLISE": "ANÁLISE QUALITATIVA" if cls._find_paragraph(document, "ANÁLISE QUALITATIVA") else "Conclusão",
            "Conclusão": "Sugestões de Conduta (Encaminhamentos):",
            "Sugestões de Conduta (Encaminhamentos):": "Considerações Finais",
            "Referências Bibliográficas": None,
            "Referencia Bibliográfica": None,
        }

        for heading, text in replacements.items():
            cls._replace_block_between_headings(
                document, heading, boundaries.get(heading), text
            )

        cls._rename_heading_if_present(document, "Conclusão", "14. CONCLUSÃO")
        cls._rename_heading_if_present(
            document,
            "Sugestões de Conduta (Encaminhamentos):",
            "15. SUGESTÕES DE CONDUTA (ENCAMINHAMENTOS):",
        )
        cls._rename_heading_if_present(document, "Considerações Finais", "16. CONSIDERAÇÕES FINAIS")
        cls._rename_heading_if_present(document, "Referências Bibliográficas", "17. REFERÊNCIAS BIBLIOGRÁFICAS")
        cls._rename_heading_if_present(document, "Referencia Bibliográfica", "17. REFERÊNCIAS BIBLIOGRÁFICAS")

        cls._replace_identification_block(document, report, context)
        cls._replace_wais3_demand_block(document, sections, context)
        cls._replace_wais3_procedures_block(document, report, sections, context)

    @classmethod
    def _sanitize_generated_document(cls, document: Document, report, context: dict):
        cls._remove_invalid_paragraph_patterns(document)
        cls._normalize_document_spacing(document)
        cls._remove_invalid_content_after_references(document)
        cls._remove_empty_paragraphs(document)
        cls._replace_foreign_patient_names(document, context)

    @classmethod
    def _paragraph_contains_chart(cls, paragraph) -> bool:
        return cls._element_contains_chart(paragraph._p)

    @classmethod
    def _replace_foreign_patient_names(cls, document: Document, context: dict):
        patient_name = ((context or {}).get("patient") or {}).get("full_name")
        if not patient_name:
            return
        patient_name = patient_name.strip()
        if not patient_name:
            return

        foreign_names = cls._foreign_patient_names_in_text(
            cls._document_text_before_references(document),
            patient_name,
        )
        if not foreign_names:
            return

        full_doc_text = "\n".join(
            paragraph.text or ""
            for paragraph in document.paragraphs
        )
        if not foreign_names:
            return

        for foreign_name in foreign_names:
            if foreign_name in full_doc_text:
                cls._replace_name_in_document(document, foreign_name, patient_name)

    @classmethod
    def _replace_name_in_document(cls, document: Document, old_name: str, new_name: str):
        for paragraph in document.paragraphs:
            if cls._paragraph_contains_chart(paragraph):
                continue
            full_text = paragraph.text or ""
            if old_name not in full_text:
                continue
            for run in paragraph.runs:
                if old_name in (run.text or ""):
                    run.text = (run.text or "").replace(old_name, new_name)

    @classmethod
    def _remove_invalid_paragraph_patterns(cls, document: Document):
        for paragraph in document.paragraphs:
            if cls._paragraph_contains_chart(paragraph):
                continue
            text = paragraph.text or ""
            updated = text
            for pattern in cls.INVALID_DOCX_PATTERNS:
                updated = updated.replace(pattern, "")
            if updated != text:
                paragraph.text = updated

    @classmethod
    def _normalize_document_spacing(cls, document: Document):
        for paragraph in document.paragraphs:
            if cls._paragraph_contains_chart(paragraph):
                continue
            normalized = PtBrTextService.normalize(paragraph.text or "")
            normalized = re.sub(r"[ \t]{2,}", " ", normalized)
            normalized = re.sub(r"\s+([,.;:])", r"\1", normalized)
            normalized = re.sub(r"\n{3,}", "\n\n", normalized)
            if normalized != (paragraph.text or ""):
                paragraph.text = normalized.strip()

    @classmethod
    def _remove_invalid_content_after_references(cls, document: Document):
        body = document._body._element
        found_references = False
        for element in list(body.iterchildren()):
            text = cls._body_element_text(element)
            upper_text = text.upper()
            if any(heading in upper_text for heading in cls.REFERENCE_SECTION_HEADINGS):
                found_references = True
                continue
            if not found_references:
                continue
            if element.tag.endswith("tbl"):
                body.remove(element)
                continue
            if cls._is_invalid_post_references_text(text):
                body.remove(element)

    @classmethod
    def _remove_empty_paragraphs(cls, document: Document):
        body = document._body._element
        for paragraph in list(document.paragraphs):
            if cls._paragraph_contains_chart(paragraph):
                continue
            if paragraph.text.strip():
                continue
            parent = paragraph._p.getparent()
            if parent is body:
                body.remove(paragraph._p)

    @classmethod
    def _foreign_patient_names_in_text(cls, text: str | None, patient_name: str | None) -> list[str]:
        patient_name = (patient_name or "").strip()
        if not patient_name:
            return []
        allowed_tokens = {token for token in patient_name.split() if token}
        technical_tokens = {
            "Raciocínio", "Matricial", "Execução", "Tabela", "Rey", "Auditory", "Verbal",
            "Learning", "Test", "Flexibilidade", "Cognitiva", "Big", "Five", "Interações",
            "Sociais", "Espectro", "Autista", "Pontuação", "Total", "Vocabulário",
            "Semelhanças", "Cubos", "Pânico", "Sintomas", "Somáticos", "Ansiedade",
            "Generalizada", "Separação", "Fobia", "Social", "Evitação", "Escolar",
            "Percepção", "Cognição", "Comunicação", "Motivação", "Realização", "Abertura",
        }
        ignored_names = {
            patient_name,
            cls.FIXED_AUTHOR.split("(", 1)[0].strip(),
            "Conselho Federal",
            "Microsoft Word",
            "Escala Wechsler",
            "Bateria Psicológica",
            "Teste dos",
            "Rey Auditory",
            "Escala Baptista",
            "Bateria Fatorial",
            "Social Responsiveness",
            "Screen for Child",
            "Rey Auditory Verbal Learning Test",
            "Raciocínio Matricial",
            "Flexibilidade Cognitiva",
            "Interações Sociais",
            "Espectro Autista",
            "Pontuação Total",
            "Big Five",
            "Execução Tabela",
        }
        candidates = re.findall(
            r"\b[A-ZÁÉÍÓÚÂÊÔÃÕÇ][a-záéíóúâêôãõç]+(?:\s+[A-ZÁÉÍÓÚÂÊÔÃÕÇ][a-záéíóúâêôãõç]+)+\b",
            text or "",
        )
        foreign_names = []
        for name in candidates:
            if name in ignored_names:
                continue
            if any(pattern.lower() in name.lower() for pattern in cls.SKIP_PATTERNS):
                continue
            words = name.split()
            if len(words) < 2 or len(words) > 5:
                continue
            if all(word in technical_tokens for word in words):
                continue
            first_name = words[0]
            if name == patient_name or first_name in allowed_tokens:
                continue
            if name not in foreign_names:
                foreign_names.append(name)
        return foreign_names

    @classmethod
    def _sanitize_section_text_for_patient(cls, text: str | None, context: dict) -> str:
        patient_name = ((context or {}).get("patient") or {}).get("full_name") or ""
        cleaned = str(text or "").strip()
        if not cleaned:
            return ""
        if cls._foreign_patient_names_in_text(cleaned, patient_name):
            return ""
        return cleaned

    @classmethod
    def _validate_patient_identity(cls, document: Document, report, context: dict):
        patient_name = ((context or {}).get("patient") or {}).get("full_name")
        if not patient_name and getattr(report, "patient", None) is not None:
            patient_name = getattr(report.patient, "full_name", "")
        patient_name = (patient_name or "").strip()
        if not patient_name:
            return

        foreign_names = cls._foreign_patient_names_in_text(
            cls._document_text_before_references(document),
            patient_name,
        )
        if foreign_names:
            raise ValueError(
                "Exportação bloqueada: o laudo contém nomes divergentes de pacientes: "
                + ", ".join(foreign_names)
            )

    @classmethod
    def _validate_unique_wasi_result(cls, document: Document):
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        patterns = (
            r"QIT\s*=\s*(\d+)",
            r"QI\s*Total\s*=\s*(\d+)",
            r"QI\s*TOTAL\s*(\d+)",
        )
        values = []
        for pattern in patterns:
            values.extend(re.findall(pattern, text, flags=re.IGNORECASE))
        unique_values = list(dict.fromkeys(values))
        if len(unique_values) > 1:
            raise ValueError(
                "Exportação bloqueada: resultados divergentes de QIT/QI Total encontrados: "
                + ", ".join(unique_values)
            )

    @classmethod
    def _document_text_before_references(cls, document: Document) -> str:
        lines = []
        for paragraph in document.paragraphs:
            text = (paragraph.text or "").strip()
            if any(heading in text.upper() for heading in cls.REFERENCE_SECTION_HEADINGS):
                break
            if text.startswith(("Autora:", "Filiação:")):
                continue
            lines.append(text)
        return "\n".join(lines)

    @staticmethod
    def _body_element_text(element) -> str:
        return "".join(node.text or "" for node in element.iter())

    @classmethod
    def _is_invalid_post_references_text(cls, text: str) -> bool:
        stripped = (text or "").strip()
        if not stripped:
            return True
        if any(pattern in stripped for pattern in cls.INVALID_DOCX_PATTERNS):
            return True
        if stripped.startswith(("{", "[")):
            return True
        if re.search(r'"[^\"]+"\s*:', stripped):
            return True
        return False

    @classmethod
    def _replace_identification_block(cls, document: Document, report, context: dict):
        start = cls._find_paragraph(document, "IDENTIFICAÇÃO")
        if start is None:
            return
        end = cls._find_paragraph(document, "DESCRIÇÃO DA DEMANDA")
        cls._remove_nodes_between(start, end)

        patient = (context or {}).get("patient") or {}
        author_name = cls.FIXED_AUTHOR
        interested_party = cls.FIXED_INTERESTED_PARTY
        purpose = cls.FIXED_PURPOSE

        anchor = cls._insert_paragraph_after(start, "1.1. Identificação do laudo:")
        cls._format_subtitle_paragraph(anchor)
        anchor = cls._insert_identification_label_value_after(anchor, "Autora", author_name)
        anchor = cls._insert_identification_label_value_after(anchor, "Interessado", interested_party)
        anchor = cls._insert_identification_label_value_after(anchor, "Finalidade", purpose)

        patient_heading = "1.2. Identificação da paciente:" if patient.get("sex") == "F" else "1.2. Identificação do paciente:"
        anchor = cls._insert_paragraph_after(anchor, patient_heading)
        cls._format_subtitle_paragraph(anchor)
        anchor = cls._insert_identification_label_value_after(
            anchor, "Nome", patient.get("full_name") or "Não informado"
        )
        anchor = cls._insert_identification_label_value_after(
            anchor, "Sexo", cls._format_sex_display(patient.get("sex"))
        )
        anchor = cls._insert_identification_label_value_after(
            anchor, "Data de nascimento", cls._format_date_display(patient.get("birth_date"))
        )
        anchor = cls._insert_identification_label_value_after(anchor, "Idade", cls._age_text(context))
        anchor = cls._insert_identification_label_value_after(anchor, "Filiação", cls._filiation_text(patient))
        cls._insert_identification_label_value_after(anchor, "Escolaridade", cls._schooling_text(patient))

    @classmethod
    def _insert_identification_label_value_after(cls, anchor, label: str, value: str):
        paragraph = cls._insert_paragraph_after(anchor, "")
        cls._append_identification_label_value(paragraph, label, value)
        return paragraph

    @classmethod
    def _insert_bullet_paragraph_after(cls, anchor, text: str):
        paragraph = cls._insert_paragraph_after(anchor, "")
        bullet_prefix = paragraph.add_run("• ")
        bullet_prefix.font.name = cls.FONT_NAME
        bullet_prefix.font.size = cls.BODY_SIZE
        label, separator, description = PtBrTextService.normalize(text).partition(":")
        label_run = paragraph.add_run(label)
        label_run.font.name = cls.FONT_NAME
        label_run.font.size = cls.BODY_SIZE
        label_run.bold = True
        if separator:
            separator_run = paragraph.add_run(f":{description}")
            separator_run.font.name = cls.FONT_NAME
            separator_run.font.size = cls.BODY_SIZE
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = cls.BODY_SPACE_AFTER
        paragraph.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Cm(0.63)
        return paragraph

    @staticmethod
    def _sanitize_procedures_section_text(text: str | None) -> str:
        normalized = PtBrTextService.normalize(text or "").strip()
        if not normalized or normalized == "Sem conteúdo disponível para esta seção.":
            return ""
        lines = [line.strip() for line in normalized.splitlines() if line.strip()]
        clean_lines = [
            line for line in lines
            if not line.startswith("•") and not line.startswith("-")
        ]
        if clean_lines:
            return clean_lines[0]
        return ""

    @classmethod
    def _append_identification_label_value(cls, paragraph, label: str, value: str):
        label = PtBrTextService.normalize(label)
        value = PtBrTextService.normalize(str(value or "Não informado"))
        label_run = paragraph.add_run(f"{label}: ")
        label_run.font.name = cls.FONT_NAME
        label_run.font.size = cls.BODY_SIZE
        label_run.bold = True
        value_run = paragraph.add_run(value)
        value_run.font.name = cls.FONT_NAME
        value_run.font.size = cls.BODY_SIZE
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = cls.IDENTIFICATION_LINE_SPACING
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)

    @classmethod
    def _replace_block_between_headings(
        cls, document: Document, start_heading: str, end_heading: str | None, text: str
    ):
        start = cls._find_paragraph(document, start_heading)
        if start is None:
            return
        end = cls._find_paragraph(document, end_heading) if end_heading else None
        cls._remove_nodes_between(start, end)
        anchor = start
        for line in (text or "Sem conteúdo disponível para esta seção.").split("\n"):
            anchor = cls._insert_paragraph_after(anchor, "")
            cls._append_body_text_with_bold_label(anchor, line)
            cls._format_body_paragraph(anchor)

    @classmethod
    def _replace_wais3_demand_block(cls, document: Document, sections: dict[str, str], context: dict):
        start = cls._find_paragraph(document, "DESCRIÇÃO DA DEMANDA")
        if start is None:
            return
        end = cls._find_paragraph(document, "PROCEDIMENTOS")
        cls._remove_nodes_between(start, end)
        anchor = cls._insert_paragraph_after(start, "Motivo do Encaminhamento")
        cls._format_subtitle_paragraph(anchor)
        anchor = cls._insert_paragraph_after(anchor, "")
        cls._append_body_text_with_bold_label(anchor, cls._wais3_demand_text(sections, context))
        cls._format_body_paragraph(anchor)

    @classmethod
    def _replace_wais3_procedures_block(cls, document: Document, report, sections: dict[str, str], context: dict):
        start = cls._find_paragraph(document, "PROCEDIMENTOS")
        if start is None:
            return
        end = cls._find_paragraph(document, "ANÁLISE")
        cls._remove_nodes_between(start, end)

        anchor = cls._insert_paragraph_after(
            start, cls._wais3_procedures_intro_text(report, sections)
        )
        cls._format_body_paragraph(anchor)

        for item in cls._procedure_items(context):
            anchor = cls._insert_bullet_paragraph_after(
                anchor,
                f"{item['name']}: {item['description']}",
            )

    @classmethod
    def _wais3_demand_text(cls, sections: dict[str, str], context: dict) -> str:
        section_text = PtBrTextService.normalize(sections.get("descricao_demanda") or "")
        if section_text and section_text != "Sem conteúdo disponível para esta seção.":
            return section_text
        return (
            "Levando em consideração a demanda apresentada, objetivou-se avaliar as funções cognitivas, aspectos comportamentais, emocionais e sociais do paciente, a fim de compreender seu funcionamento global e subsidiar orientações para o manejo adequado em contextos familiar e escolar."
        )

    @classmethod
    def _wais3_procedures_intro_text(cls, report, sections: dict[str, str]) -> str:
        section_text = cls._sanitize_procedures_section_text(sections.get("procedimentos"))
        if section_text:
            return section_text
        testing_sessions = getattr(getattr(report, "evaluation", None), "testing_sessions_count", None)
        if testing_sessions:
            return f"Para esta avaliação foram realizadas: uma sessão de anamnese, {testing_sessions} sessões de testagem com o paciente e uma sessão de devolutiva."
        return "Para esta avaliação foram realizadas entrevista clínica inicial, sessões de testagem com o paciente e devolutiva clínica, conforme a necessidade do caso."

    @classmethod
    def _rebuild_qualitative_section(
        cls, document: Document, sections: dict[str, str], context: dict
    ):
        has_analise_qualitativa = cls._find_paragraph(document, "ANÁLISE QUALITATIVA") is not None
        has_conclusao = cls._find_paragraph(document, "Conclusão") is not None
        
        if has_analise_qualitativa and has_conclusao:
            start = cls._find_paragraph(document, "ANÁLISE QUALITATIVA")
            end = cls._find_paragraph(document, "Conclusão")
            cls._remove_nodes_between(start, end)
        else:
            return
        
        tests = {
            item.get("instrument_code"): item
            for item in context.get("validated_tests") or []
        }
        table_index = 1
        chart_index = 1
        table_index = 1
        chart_index = 1
        anchor = start
        template_chart_blocks = cls._extract_template_chart_blocks(document)
        is_adolescent = cls._is_adolescent_document(document, context)
        template_chart_map = {}
        if tests.get('wisc4'):
            template_chart_map = {
                'wisc4': template_chart_blocks[0] if len(template_chart_blocks) > 0 else None,
                'bpa2': template_chart_blocks[1] if len(template_chart_blocks) > 1 else None,
                'ravlt': template_chart_blocks[2] if len(template_chart_blocks) > 2 else None,
                'fdt_auto': template_chart_blocks[3] if len(template_chart_blocks) > 3 else None,
                'fdt_control': template_chart_blocks[4] if len(template_chart_blocks) > 4 else None,
                'etdah_pais': template_chart_blocks[5] if len(template_chart_blocks) > 5 else None,
                'etdah_ad': template_chart_blocks[6] if len(template_chart_blocks) > 6 else None,
                'scared_pair': template_chart_blocks[7] if len(template_chart_blocks) > 7 else None,
                'epq': template_chart_blocks[8] if len(template_chart_blocks) > 8 else None,
                'srs2': template_chart_blocks[9] if len(template_chart_blocks) > 9 else None,
            }
        elif tests.get('wais3'):
            wais3_template_chart_blocks = cls._extract_chart_blocks_from_template_path(
                cls.WAIS3_TEMPLATE_PATH
            )
            wisc4_template_chart_blocks = cls._extract_chart_blocks_from_template_path(
                cls.WISC4_TEMPLATE_PATH
            )
            template_chart_map = {
                'wais3': (
                    template_chart_blocks[0]
                    if len(template_chart_blocks) > 0
                    else wais3_template_chart_blocks[0]
                    if len(wais3_template_chart_blocks) > 0
                    else None
                ),
                'bpa2': template_chart_blocks[1] if len(template_chart_blocks) > 1 else None,
                'ravlt': template_chart_blocks[2] if len(template_chart_blocks) > 2 else None,
                'fdt_auto': template_chart_blocks[3] if len(template_chart_blocks) > 3 else None,
                'fdt_control': template_chart_blocks[4] if len(template_chart_blocks) > 4 else None,
                'etdah_ad': template_chart_blocks[5] if len(template_chart_blocks) > 5 else None,
                'scared_pair': (
                    wisc4_template_chart_blocks[7]
                    if len(wisc4_template_chart_blocks) > 7
                    else None
                ),
                'srs2': template_chart_blocks[6] if len(template_chart_blocks) > 6 else None,
            }
        elif tests.get('wasi'):
            template_chart_map = {
                'wasi': template_chart_blocks[0] if len(template_chart_blocks) > 0 else None,
                'bpa2': template_chart_blocks[1] if len(template_chart_blocks) > 1 else None,
                'ravlt': template_chart_blocks[2] if len(template_chart_blocks) > 2 else None,
                'fdt_auto': template_chart_blocks[3] if len(template_chart_blocks) > 3 else None,
                'fdt_control': template_chart_blocks[4] if len(template_chart_blocks) > 4 else None,
                'etdah_ad': template_chart_blocks[5] if len(template_chart_blocks) > 5 else None,
                'srs2': template_chart_blocks[6] if len(template_chart_blocks) > 6 else None,
            }

        def add_title(text: str):
            nonlocal anchor
            anchor = cls._insert_paragraph_after(anchor, PtBrTextService.normalize(text))
            cls._format_subtitle_paragraph(anchor)

        def add_text(text: str):
            nonlocal anchor
            text = PtBrTextService.normalize(text)
            for line in (text or "").split("\n"):
                if not line.strip():
                    continue
                if line.lstrip().startswith("- "):
                    anchor = cls._insert_paragraph_after(anchor, "")
                    bullet_text = line.lstrip()[2:].strip()
                    p = anchor
                    lead, separator, tail = bullet_text.partition(" Avaliou ")
                    bullet_prefix = p.add_run("-\t")
                    bullet_prefix.font.name = cls.FONT_NAME
                    bullet_prefix.font.size = cls.BODY_SIZE
                    bold_run = p.add_run(lead if separator else bullet_text)
                    bold_run.font.name = cls.FONT_NAME
                    bold_run.font.size = cls.BODY_SIZE
                    bold_run.bold = True
                    if separator:
                        tail_run = p.add_run(f" Avaliou {tail}")
                        tail_run.font.name = cls.FONT_NAME
                        tail_run.font.size = cls.BODY_SIZE
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = cls.BODY_SPACE_AFTER
                    p.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
                    p.paragraph_format.left_indent = Cm(1.1)
                    p.paragraph_format.first_line_indent = Cm(-0.55)
                    p.paragraph_format.tab_stops.add_tab_stop(Cm(1.1))
                else:
                    anchor = cls._insert_paragraph_after(anchor, "")
                    if line.startswith("Capacidade Cognitiva Global:"):
                        cls._append_wisc_intro_paragraph(anchor, line)
                    else:
                        cls._append_body_text_with_bold_label(anchor, line)
                        cls._format_left_body_paragraph(anchor)

        def add_table(caption: str, rows: list[list[str]] | None, table_key: str):
            nonlocal anchor, table_index
            if not rows:
                return
            if table_key == "ravlt":
                table = cls._insert_ravlt_table_after(anchor, rows)
            else:
                table = cls._insert_table_after(anchor, rows, table_key)
                cls._format_table(table, table_key)
            anchor = cls._insert_paragraph_after_table(
                table, cls._table_caption_text(table_index, caption)
            )
            cls._format_caption_paragraph(anchor)
            table_index += 1

        def section_or_test_interpretation(
            section_key: str,
            fallback_section_key: str | None,
            test_payload: dict | None,
        ) -> str:
            return cls._resolve_interpretation_text(
                sections.get(section_key),
                sections.get(fallback_section_key) if fallback_section_key else None,
                test_payload,
                context,
            )

        def add_chart(
            caption: str,
            image_bytes: bytes | None,
            note: str | None = None,
            width=None,
            height=None,
            show_caption: bool = True,
            template_key: str | None = None,
        ):
            nonlocal anchor, chart_index
            if template_key and template_chart_map.get(template_key) is not None:
                anchor = cls._insert_paragraph_after(anchor, "")
                anchor.paragraph_format.first_line_indent = Pt(0)
                anchor.paragraph_format.left_indent = Pt(0)
                anchor.paragraph_format.right_indent = Pt(0)
                anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER
                anchor._p.addnext(deepcopy(template_chart_map[template_key]))
                anchor = Paragraph(anchor._p.getnext(), anchor._parent)
                if show_caption:
                    anchor = cls._insert_paragraph_after(
                        anchor, cls._chart_caption_text(chart_index, caption)
                    )
                    cls._format_caption_paragraph(anchor)
                if note:
                    anchor = cls._insert_paragraph_after(anchor, note)
                    cls._format_chart_legend_paragraph(anchor)
                chart_index += 1
                return
            if not image_bytes:
                return
            anchor = cls._insert_paragraph_after(anchor, "")
            anchor.paragraph_format.first_line_indent = Pt(0)
            anchor.paragraph_format.left_indent = Pt(0)
            anchor.paragraph_format.space_before = Pt(0)
            run = anchor.runs[0] if anchor.runs else anchor.add_run()
            picture_kwargs = {}
            if width is not None:
                picture_kwargs["width"] = width
            if height is not None:
                picture_kwargs["height"] = height
            if not picture_kwargs:
                picture_kwargs["width"] = cls.IMAGE_WIDTH
            run.add_picture(BytesIO(image_bytes), **picture_kwargs)
            anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if show_caption:
                anchor = cls._insert_paragraph_after(
                    anchor, cls._chart_caption_text(chart_index, caption)
                )
                cls._format_caption_paragraph(anchor)
            if note:
                anchor = cls._insert_paragraph_after(anchor, note)
                cls._format_chart_legend_paragraph(anchor)
            chart_index += 1

        patient_title = "da paciente" if (context.get("patient") or {}).get("sex") == "F" else "do paciente"
        if tests.get("wisc4"):
            add_title("Capacidade Cognitiva Global")
            add_text(cls._wisc_global_intro_text(tests.get("wisc4"), context))
            for lead, tail in cls._wisc_global_bullet_parts(tests.get("wisc4")):
                add_text(f"- {lead} {tail}")
            add_title(f"Desempenho {patient_title} no WISC-IV")
            add_chart(
                "WISC-IV - INDICES DE QI",
                cls._wisc_chart(tests.get("wisc4")),
                cls._wisc_chart_legend(chart_index),
                show_caption=False,
                template_key="wisc4",
            )
            anchor = cls._insert_interpretation_block_after(
                anchor,
                cls._wisc_model_interpretation_text(tests.get("wisc4"), context),
            )
        elif tests.get("wais3"):
            add_text(cls._wais3_intro_text(tests.get("wais3"), context))
            for lead, tail in cls._wais3_global_bullet_parts(tests.get("wais3")):
                add_text(f"- {lead} {tail}")
            add_title(f"Desempenho {patient_title} no WAIS III")
            add_chart(
                "WAIS III - INDICES DE QIS",
                cls._wais3_chart(tests.get("wais3")),
                show_caption=True,
                template_key="wais3",
            )
            anchor = cls._insert_interpretation_block_after(
                anchor,
                cls._normalize_interpretation_text(
                    section_or_test_interpretation(
                        "eficiencia_intelectual", None, tests.get("wais3")
                    )
                ),
            )
        elif tests.get("wasi"):
            add_text(cls._wasi_intro_text(tests.get("wasi"), context))
            for lead, tail in cls._wasi_global_bullet_parts(tests.get("wasi")):
                add_text(f"- {lead} - {tail}")
            add_title(f"Desempenho {patient_title} no WASI")
            add_chart(
                WASI_CHART_SPEC["title"],
                cls._wasi_chart_image(tests.get("wasi")),
                show_caption=True,
                template_key="wasi",
            )
            anchor = cls._insert_interpretation_block_after(
                anchor,
                cls._normalize_interpretation_text(
                    section_or_test_interpretation(
                        "eficiencia_intelectual", None, tests.get("wasi")
                    )
                ),
            )

        if is_adolescent and tests.get("wisc4"):
            add_title("Subescalas WISC-IV")

            def add_numbered_section(title: str):
                add_title(title)

            add_title("Função Executiva")
            add_table(
                "Resultado da Função executiva",
                cls._wisc_rows(tests.get("wisc4"), context, "funcoes_executivas"),
                "wisc",
            )
            add_text(cls._wisc_section_text(sections, "funcoes_executivas", context))

            add_title("Linguagem")
            add_table(
                "Resultados da Linguagem",
                cls._wisc_rows(tests.get("wisc4"), context, "linguagem"),
                "wisc",
            )
            add_text(cls._wisc_section_text(sections, "linguagem", context))

            add_title("Gnosias e Praxias")
            add_table(
                "Resultados da Gnosias e praxias",
                cls._wisc_rows(tests.get("wisc4"), context, "gnosias_praxias"),
                "wisc",
            )
            add_text(cls._wisc_section_text(sections, "gnosias_praxias", context))

            add_title("Memória e Aprendizagem")
            add_table(
                "Resultados de Memória e Aprendizagem",
                cls._wisc_memory_rows(tests.get("wisc4"), context),
                "wisc",
            )
            add_text(cls._wisc_section_text(sections, "memoria_aprendizagem", context))

            if tests.get("bpa2"):
                add_numbered_section("BPA-2 – BATERIA PSICOLÓGICA PARA AVALIAÇÃO DA ATENÇÃO")
                add_table("Atenção BPA-2 Resultados", cls._bpa_rows(tests.get("bpa2"), context), "bpa")
                anchor = cls._insert_bpa_interpretation_block_after(anchor, tests.get("bpa2"), context)
                add_chart(
                    "BPA-2 Resultados da Avaliação da Atenção",
                    cls._bpa_chart_bytes(tests.get("bpa2")),
                    template_key="bpa2",
                )

            if tests.get("ravlt"):
                add_numbered_section("RAVLT – REY AUDITORY VERBAL LEARNING TEST")
                ravlt_interpretation = cls._resolve_interpretation_text(
                    sections.get("ravlt"),
                    sections.get("memoria_aprendizagem"),
                    tests.get("ravlt"),
                )
                anchor = cls._insert_ravlt_conceptual_paragraph_after(anchor)
                add_table(
                    cls._ravlt_table_caption(), cls._ravlt_rows(tests.get("ravlt"), context), "ravlt"
                )
                if ravlt_interpretation:
                    anchor = cls._insert_interpretation_block_after(
                        anchor, cls._normalize_interpretation_text(ravlt_interpretation)
                    )
                add_chart(
                    "RAVLT Resultados",
                    cls._ravlt_chart(tests.get("ravlt")),
                    template_key="ravlt",
                )

            if tests.get("fdt"):
                add_numbered_section("FDT – TESTE DOS CINCO DÍGITOS")
                add_text(cls._fdt_description_text())
                add_table(
                    "FDT Processos Automáticos e Controlados",
                    cls._fdt_rows(tests.get("fdt")),
                    "fdt",
                )
                fdt_interpretation = cls._resolve_interpretation_text(
                    sections.get("fdt"),
                    sections.get("funcoes_executivas"),
                    tests.get("fdt"),
                )
                if fdt_interpretation:
                    anchor = cls._insert_interpretation_block_after(
                        anchor, cls._normalize_interpretation_text(fdt_interpretation)
                    )
                add_chart(
                    "FDT Processos Automáticos",
                    cls._fdt_chart(tests.get("fdt"), automatic=True),
                    width=Cm(14),
                    height=Cm(10),
                    template_key="fdt_auto",
                )
                add_chart(
                    "FDT Processos Controlados",
                    cls._fdt_chart(tests.get("fdt"), automatic=False),
                    width=Cm(14),
                    height=Cm(10),
                    template_key="fdt_control",
                )

            if tests.get("etdah_pais"):
                add_numbered_section("E-TDAH-PAIS")
                add_text(
                    "A Escala E-TDAH-PAIS tem como objetivo identificar manifestações comportamentais e emocionais associadas ao Transtorno do Déficit de Atenção e Hiperatividade (TDAH) a partir da percepção dos pais ou responsáveis, avaliando domínios relacionados à regulação emocional, impulsividade, comportamento adaptativo e atenção. O instrumento fornece indicadores quantitativos e qualitativos do funcionamento atencional e comportamental, permitindo compreender o impacto desses aspectos no cotidiano da criança (Benczik, 2005)."
                )
                add_table(
                    "Resultados do E-TDAH-PAIS",
                    cls._etdah_rows(tests.get("etdah_pais")),
                    "etdah_pais",
                )
                anchor = cls._insert_etdah_pais_interpretation_block_after(anchor, tests.get("etdah_pais"), context)
                add_chart(
                    "E-TDAH-PAIS Percentis",
                    cls._etdah_chart(tests.get("etdah_pais")),
                    template_key="etdah_pais",
                )

            if tests.get("etdah_ad"):
                add_numbered_section("E-TDAH-AD")
                add_text(cls._etdah_ad_description_text())
                add_table(
                    "ETDAH-AD RESULTADO",
                    cls._etdah_rows(tests.get("etdah_ad")),
                    "etdah_ad",
                )
                anchor = cls._insert_etdah_ad_interpretation_block_after(
                    anchor,
                    cls._normalize_interpretation_text(
                        section_or_test_interpretation("etdah_ad", None, tests.get("etdah_ad"))
                    ),
                )
                add_chart(
                    "E-TDAH-AD RESULTADO",
                    cls._etdah_chart(tests.get("etdah_ad")),
                    template_key="etdah_ad",
                )

            scared_tests = cls._find_tests(context, "scared")
            if scared_tests:
                add_numbered_section("SCARED")
                inserted_scared_pair = False
                for scared_test in scared_tests:
                    form_label = cls._scared_form_label(scared_test)
                    add_table(
                        f"SCARED Resultados {form_label}",
                        cls._scared_rows(scared_test),
                        cls._scared_table_key(scared_test),
                    )
                    anchor = cls._insert_interpretation_block_after(
                        anchor,
                        cls._normalize_interpretation_text(
                            cls._resolve_interpretation_text(None, None, scared_test, context)
                        ),
                    )
                    if not inserted_scared_pair:
                        add_chart(
                            cls._scared_form_title(scared_test),
                            cls._scared_chart(scared_test),
                            template_key="scared_pair",
                        )
                        inserted_scared_pair = True

            if tests.get("epq_j"):
                add_numbered_section("EPQ-J")
                add_text(
                    section_or_test_interpretation(
                        "epq_j", None, tests.get("epq_j")
                    )
                )
                add_table(
                    "EPQ-J Resultados da personalidade",
                    cls._epq_rows(tests.get("epq_j")),
                    "epq",
                )
                add_chart(
                    "EPQ-J Percentis",
                    cls._epq_chart(tests.get("epq_j")),
                    template_key="epq",
                )

            if tests.get("bai"):
                bai_test = tests.get("bai")
                add_numbered_section("BAI")
                add_text(cls._bai_description_text())
                add_table(
                    "BAI - Resultado da sintomatologia ansiosa",
                    cls._bai_rows(bai_test),
                    "bai_scores",
                )
                add_chart(
                    "BAI - Perfil do resultado",
                    cls._bai_chart(bai_test),
                )
                add_text(section_or_test_interpretation("bai", None, bai_test))

            if tests.get("ebadep_a") or tests.get("ebadep_ij") or tests.get("ebaped_ij"):
                ebadep_test = tests.get("ebadep_a") or tests.get("ebadep_ij") or tests.get("ebaped_ij")
                add_numbered_section("EBADEP-A")
                add_text(cls._ebadep_description_text())
                add_table(
                    "EBADEP-A - Resultado da sintomatologia",
                    cls._ebadep_rows(ebadep_test),
                    "scale_summary",
                )
                add_text(
                    cls._resolve_interpretation_text(
                        sections.get("aspectos_emocionais_comportamentais"),
                        None,
                        ebadep_test,
                        context,
                    )
                )

            if tests.get("bfp"):
                bfp_test = tests.get("bfp")
                add_numbered_section("BFP – BATERIA FATORIAL DE PERSONALIDADE")
                add_text(cls._bfp_description_text())
                bfp_tables = cls._bfp_rows(bfp_test)
                if bfp_tables:
                    for table_rows in bfp_tables:
                        add_table("Resultados gerais", table_rows, "bfp")
                add_chart(
                    "BFP – Bateria Fatorial de Personalidade",
                    cls._bfp_chart(bfp_test),
                    template_key="bfp",
                )
                anchor = cls._insert_interpretation_block_after(
                    anchor,
                    cls._normalize_interpretation_text(
                        section_or_test_interpretation("bfp", "aspectos_emocionais_comportamentais", bfp_test)
                    ),
                )

            if tests.get("srs2"):
                srs2_test = tests.get("srs2")
                add_numbered_section("SRS-2 – ESCALA DE RESPONSIVIDADE SOCIAL")
                add_table("SRS-2 Resultados", cls._srs2_rows(srs2_test), "srs2")
                anchor = cls._insert_interpretation_block_after(
                    anchor,
                    cls._normalize_interpretation_text(
                        section_or_test_interpretation("srs2", None, srs2_test)
                    ),
                )
                add_chart(
                    "SRS-2 Resultados",
                    cls._srs2_chart(srs2_test),
                    template_key="srs2",
                )
        elif tests.get("wais3"):
            add_title("Subescalas WAIS III")
            builder = cls._wais3_report_builder(context)
            for block in builder.subtest_blocks:
                add_title(block.title)
                add_table(
                    block.table_caption,
                    cls._wais3_domain_rows(tests.get("wais3"), block.interpretation_section),
                    block.table_key,
                )
                add_text(cls._wisc_section_text(sections, block.interpretation_section, context))

            if tests.get("bpa2"):
                add_title("BPA-2 Bateria Psicológica para Avaliação da Atenção")
                add_text(
                    "A Bateria Psicológica para Avaliação da Atenção – BPA-2 tem como objetivo mensurar a capacidade geral de atenção e avaliar individualmente atenção concentrada, atenção dividida e atenção alternada."
                )
                add_table("Atenção BPA-2 Resultados", cls._bpa_rows(tests.get("bpa2"), context), "bpa")
                anchor = cls._insert_bpa_interpretation_block_after(anchor, tests.get("bpa2"), context)
                add_chart(
                    "BPA-2 apresenta os resultados da avaliação da atenção",
                    cls._bpa_chart_bytes(tests.get("bpa2")),
                    template_key="bpa2",
                )

            if tests.get("ravlt"):
                add_title("RAVLT Rey Auditory Verbal Learning Test")
                anchor = cls._insert_ravlt_conceptual_paragraph_after(anchor)
                add_chart(
                    "RAVLT Resultados",
                    cls._ravlt_chart(tests.get("ravlt")),
                    template_key="ravlt",
                )
                add_table(
                    cls._ravlt_table_caption(),
                    cls._ravlt_rows(tests.get("ravlt"), context),
                    "ravlt",
                )
                ravlt_interpretation = sections.get("ravlt") or sections.get("memoria_aprendizagem")
                if ravlt_interpretation:
                    anchor = cls._insert_interpretation_block_after(
                        anchor, cls._normalize_interpretation_text(ravlt_interpretation)
                    )

            if tests.get("fdt"):
                add_title("FDT- TESTE DOS CINCO DÍGITOS")
                add_text(cls._fdt_description_text())
                add_table(
                    "FDT Processos Automáticos e Controlados",
                    cls._fdt_rows(tests.get("fdt")),
                    "fdt",
                )
                fdt_interpretation = sections.get("fdt") or sections.get("funcoes_executivas")
                if fdt_interpretation:
                    anchor = cls._insert_interpretation_block_after(
                        anchor, cls._normalize_interpretation_text(fdt_interpretation)
                    )
                add_chart(
                    "FDT Processos Automáticos",
                    cls._fdt_chart(tests.get("fdt"), automatic=True),
                    width=Cm(14),
                    height=Cm(10),
                    template_key="fdt_auto",
                )
                add_chart(
                    "FDT Processos Controlados",
                    cls._fdt_chart(tests.get("fdt"), automatic=False),
                    width=Cm(14),
                    height=Cm(10),
                    template_key="fdt_control",
                )

            if tests.get("etdah_ad"):
                add_title("ETDAH-AD")
                add_text(cls._etdah_ad_description_text())
                add_table(
                    "ETDAH-AD RESULTADO",
                    cls._etdah_rows(tests.get("etdah_ad")),
                    "etdah_ad",
                )
                anchor = cls._insert_etdah_ad_interpretation_block_after(
                    anchor,
                    cls._normalize_interpretation_text(
                        section_or_test_interpretation("etdah_ad", None, tests.get("etdah_ad"))
                    ),
                )
                add_chart(
                    "E-TDAH-AD RESULTADO",
                    cls._etdah_chart(tests.get("etdah_ad")),
                    template_key="etdah_ad",
                )

            if tests.get("iphexa"):
                add_title("iphexa Inventário de Personalidade Hexadimensional")
                anchor = cls._insert_interpretation_block_after(
                    anchor,
                    cls._normalize_interpretation_text(
                        cls._resolve_interpretation_text(
                            sections.get("aspectos_emocionais_comportamentais"),
                            None,
                            tests.get("iphexa"),
                        )
                    ),
                )

            if tests.get("bai"):
                bai_test = tests.get("bai")
                add_title("BAI")
                add_text(cls._bai_description_text())
                add_table(
                    "BAI - Resultado da sintomatologia ansiosa",
                    cls._bai_rows(bai_test),
                    "bai_scores",
                )
                add_chart(
                    "BAI - Perfil do resultado",
                    cls._bai_chart(bai_test),
                )
                anchor = cls._insert_interpretation_block_after(
                    anchor,
                    cls._normalize_interpretation_text(
                        section_or_test_interpretation("bai", None, bai_test)
                    ),
                )

            if tests.get("ebadep_a") or tests.get("ebadep_ij") or tests.get("ebaped_ij"):
                ebadep_test = tests.get("ebadep_a") or tests.get("ebadep_ij") or tests.get("ebaped_ij")
                add_title("EBADEP-A")
                add_text(cls._ebadep_description_text())
                add_table("EBADEP-A - Resultado da sintomatologia", cls._ebadep_rows(ebadep_test), "scale_summary")
                anchor = cls._insert_interpretation_block_after(
                    anchor,
                    cls._normalize_interpretation_text(
                        cls._resolve_interpretation_text(
                            sections.get("aspectos_emocionais_comportamentais"),
                            None,
                            ebadep_test,
                        )
                    ),
                )

            if tests.get("bfp"):
                bfp_test = tests.get("bfp")
                add_title("BFP – Bateria Fatorial de Personalidade")
                add_text(cls._bfp_description_text())
                bfp_tables = cls._bfp_rows(bfp_test)
                if bfp_tables:
                    for table_rows in bfp_tables:
                        add_table("Resultados gerais", table_rows, "bfp")
                add_chart(
                    "BFP – Bateria Fatorial de Personalidade",
                    cls._bfp_chart(bfp_test),
                    template_key="bfp",
                )
                anchor = cls._insert_interpretation_block_after(
                    anchor,
                    cls._normalize_interpretation_text(
                        section_or_test_interpretation("bfp", "aspectos_emocionais_comportamentais", bfp_test)
                    ),
                )

            scared_tests = cls._find_tests(context, "scared")
            if scared_tests:
                add_title("SCARED")
                inserted_scared_pair = False
                for scared_test in scared_tests:
                    form_label = cls._scared_form_label(scared_test)
                    add_table(
                        f"SCARED Resultados {form_label}",
                        cls._scared_rows(scared_test),
                        cls._scared_table_key(scared_test),
                    )
                    anchor = cls._insert_interpretation_block_after(
                        anchor,
                        cls._normalize_interpretation_text(
                            cls._resolve_interpretation_text(None, None, scared_test)
                        ),
                    )
                    if not inserted_scared_pair:
                        add_chart(
                            cls._scared_form_title(scared_test),
                            cls._scared_chart(scared_test),
                            template_key="scared_pair",
                        )
                        inserted_scared_pair = True

            if tests.get("srs2"):
                srs2_test = tests.get("srs2")
                add_title("SRS-2 Escala de Responsividade Social")
                add_text(cls._srs2_description_text())
                add_table("SRS-2 Resultados dos fatores", cls._srs2_rows(srs2_test), "srs2")
                anchor = cls._insert_interpretation_block_after(
                    anchor,
                    cls._normalize_interpretation_text(
                        section_or_test_interpretation("srs2", None, srs2_test)
                    ),
                )
                add_chart(
                    "SRS-2 Resultados da discrepância entre respondentes",
                    cls._srs2_chart(srs2_test),
                    template_key="srs2",
                )
        elif tests.get("wasi"):
            add_title("5.2. Subescalas WASI")

            verbal_rows = cls._wasi_subscale_rows(tests.get("wasi"), "verbal")
            if verbal_rows:
                add_title(WASI_TABLE_SPECS["verbal"]["section_title"])
                add_table(WASI_TABLE_SPECS["verbal"]["caption"], verbal_rows, "wisc")
                add_text(
                    cls._strip_markdown_heading_prefix(sections.get("linguagem"), "Linguagem")
                    or cls._wasi_verbal_interpretation_text(tests.get("wasi"), context)
                )

            execution_rows = cls._wasi_subscale_rows(tests.get("wasi"), "execucao")
            if execution_rows:
                add_title(WASI_TABLE_SPECS["execucao"]["section_title"])
                add_table(WASI_TABLE_SPECS["execucao"]["caption"], execution_rows, "wisc")
                add_text(
                    cls._strip_markdown_heading_prefix(sections.get("gnosias_praxias"), "Gnosias e Praxias")
                    or cls._wasi_execution_interpretation_text(tests.get("wasi"), context)
                )

            section_index = 6

            def add_numbered_section(title: str):
                nonlocal section_index
                add_title(f"{section_index}. {title}")
                section_index += 1

            if tests.get("bpa2"):
                add_numbered_section("BPA-2 BATERIA PSICOLÓGICA PARA AVALIAÇÃO DA ATENÇÃO")
                add_table("Atenção BPA-2 Resultados", cls._bpa_rows(tests.get("bpa2"), context), "bpa")
                anchor = cls._insert_bpa_interpretation_block_after(anchor, tests.get("bpa2"), context)
                add_chart(
                    "BPA-2 Resultados da Avaliação da Atenção",
                    cls._bpa_chart_bytes(tests.get("bpa2")),
                    template_key="bpa2",
                )

            if tests.get("ravlt"):
                add_numbered_section("RAVLT REY AUDITORY VERBAL LEARNING TEST")
                anchor = cls._insert_ravlt_conceptual_paragraph_after(anchor)
                add_table(
                    cls._ravlt_table_caption(), cls._ravlt_rows(tests.get("ravlt"), context), "ravlt"
                )
                ravlt_interpretation = cls._resolve_interpretation_text(
                    sections.get("ravlt"),
                    sections.get("memoria_aprendizagem"),
                    tests.get("ravlt"),
                )
                if ravlt_interpretation:
                    anchor = cls._insert_interpretation_block_after(
                        anchor, cls._normalize_interpretation_text(ravlt_interpretation)
                    )
                add_chart(
                    "RAVLT Resultados",
                    cls._ravlt_chart(tests.get("ravlt")),
                    template_key="ravlt",
                )

            if tests.get("fdt"):
                add_numbered_section("FDT- TESTE DOS CINCO DÍGITOS")
                add_text(cls._fdt_description_text())
                add_table(
                    "FDT Processos Automáticos e Controlados",
                    cls._fdt_rows(tests.get("fdt")),
                    "fdt",
                )
                fdt_interpretation = cls._resolve_interpretation_text(
                    sections.get("fdt"),
                    sections.get("funcoes_executivas"),
                    tests.get("fdt"),
                )
                if fdt_interpretation:
                    anchor = cls._insert_interpretation_block_after(
                        anchor, cls._normalize_interpretation_text(fdt_interpretation)
                    )
                add_chart(
                    "FDT Processos Automáticos",
                    cls._fdt_chart(tests.get("fdt"), automatic=True),
                    width=Cm(14),
                    height=Cm(10),
                    template_key="fdt_auto",
                )
                add_chart(
                    "FDT Processos Controlados",
                    cls._fdt_chart(tests.get("fdt"), automatic=False),
                    width=Cm(14),
                    height=Cm(10),
                    template_key="fdt_control",
                )

            if tests.get("etdah_ad"):
                add_numbered_section("ETDAH-AD")
                add_text(cls._etdah_ad_description_text())
                add_table(
                    "E-TDAH Resultados",
                    cls._etdah_rows(tests.get("etdah_ad")),
                    "etdah_ad",
                )
                anchor = cls._insert_etdah_ad_interpretation_block_after(
                    anchor,
                    cls._normalize_interpretation_text(
                        section_or_test_interpretation("etdah_ad", None, tests.get("etdah_ad"))
                    ),
                )
                add_chart(
                    "E-TDAH Resultados",
                    cls._etdah_chart(tests.get("etdah_ad")),
                    template_key="etdah_ad",
                )

            if tests.get("bfp"):
                add_numbered_section("BFP- BATERIA FATORIAL DE PERSONALIDADE")
                add_text(cls._bfp_description_text())
                for table_rows in cls._bfp_rows(tests.get("bfp")) or []:
                    add_table("BFP Resultados dos fatores", table_rows, "bfp")
                add_chart(
                    "BFP – Bateria Fatorial de Personalidade",
                    cls._bfp_chart(tests.get("bfp")),
                    template_key="bfp",
                )
                add_text(
                    section_or_test_interpretation(
                        "bfp",
                        "aspectos_emocionais_comportamentais",
                        tests.get("bfp"),
                    )
                )

            if tests.get("bai"):
                add_numbered_section("BAI INVENTÁRIO DE ANSIEDADE DE BECK")
                add_text(cls._bai_description_text())
                add_table(
                    "BAI - Resultado da sintomatologia ansiosa",
                    cls._bai_rows(tests.get("bai")),
                    "bai_scores",
                )
                add_chart(
                    "BAI - Perfil do resultado",
                    cls._bai_chart(tests.get("bai")),
                )
                add_text(section_or_test_interpretation("bai", None, tests.get("bai")))

            if tests.get("ebadep_a") or tests.get("ebadep_ij") or tests.get("ebaped_ij"):
                ebadep_test = tests.get("ebadep_a") or tests.get("ebadep_ij") or tests.get("ebaped_ij")
                add_numbered_section("EBADEP-A")
                add_text(cls._ebadep_description_text())
                add_table(
                    "Resultados da EBADEP",
                    cls._ebadep_rows(ebadep_test),
                    "scale_summary",
                )
                add_text(
                    cls._resolve_interpretation_text(
                        sections.get("aspectos_emocionais_comportamentais"),
                        None,
                        ebadep_test,
                    )
                )

            if tests.get("srs2"):
                add_numbered_section("SRS-2 ESCALA DE RESPONSIVIDADE SOCIAL")
                add_table("SRS-2 Resultados", cls._srs2_rows(tests.get("srs2")), "srs2")
                anchor = cls._insert_interpretation_block_after(
                    anchor,
                    cls._normalize_interpretation_text(
                        section_or_test_interpretation("srs2", None, tests.get("srs2"))
                    ),
                )
                add_chart(
                    "SRS-2 Resultados",
                    cls._srs2_chart(tests.get("srs2")),
                    template_key="srs2",
                )
        else:
            section_index = 1

            def add_numbered_section(title: str):
                nonlocal section_index
                add_title(f"5.{section_index}. {title}")
                section_index += 1

            if tests.get("bpa2"):
                add_numbered_section("Atenção")
                add_table("Atenção BPA-2 Resultados", cls._bpa_rows(tests.get("bpa2"), context), "bpa")
                add_chart(
                    "BPA-2 Resultados da Avaliação da Atenção",
                    cls._bpa_chart_bytes(tests.get("bpa2")),
                    template_key="bpa2",
                )
                anchor = cls._insert_bpa_interpretation_block_after(anchor, tests.get("bpa2"), context)

            if tests.get("ravlt"):
                add_numbered_section("Memória e Aprendizagem")
                anchor = cls._insert_ravlt_conceptual_paragraph_after(anchor)
                add_chart(
                    "RAVLT Resultados",
                    cls._ravlt_chart(tests.get("ravlt")),
                    template_key="ravlt",
                )
                add_table(
                    cls._ravlt_table_caption(), cls._ravlt_rows(tests.get("ravlt"), context), "ravlt"
                )

            if tests.get("fdt"):
                add_numbered_section("Funções Executivas")
                add_text(cls._fdt_description_text())
                add_table(
                    "FDT Processos Automáticos e Controlados",
                    cls._fdt_rows(tests.get("fdt")),
                    "fdt",
                )
                fdt_interpretation = sections.get("funcoes_executivas", "")
                if fdt_interpretation:
                    anchor = cls._insert_interpretation_block_after(
                        anchor, cls._normalize_interpretation_text(fdt_interpretation)
                    )
                add_chart(
                    "FDT Processos Automáticos",
                    cls._fdt_chart(tests.get("fdt"), automatic=True),
                    width=Cm(14),
                    height=Cm(10),
                    template_key="fdt_auto",
                )
                add_chart(
                    "FDT Processos Controlados",
                    cls._fdt_chart(tests.get("fdt"), automatic=False),
                    width=Cm(14),
                    height=Cm(10),
                    template_key="fdt_control",
                )

            has_emotional_tests = any(
                [
                    tests.get("etdah_ad"),
                    tests.get("etdah_pais"),
                    tests.get("bai"),
                    cls._find_tests(context, "scared"),
                    tests.get("epq_j"),
                    tests.get("bfp"),
                    tests.get("srs2"),
                    tests.get("ebadep_a") or tests.get("ebadep_ij") or tests.get("ebaped_ij"),
                ]
            )
            if has_emotional_tests:
                add_numbered_section("Aspectos Emocionais, Comportamentais e Escalas Complementares")
                add_text(sections.get("aspectos_emocionais_comportamentais", ""))
                if tests.get("etdah_ad") or tests.get("etdah_pais"):
                    add_table(
                        "E-TDAH Resultados",
                        cls._etdah_rows(tests.get("etdah_ad") or tests.get("etdah_pais")),
                        "etdah_ad" if tests.get("etdah_ad") else "etdah_pais",
                    )
                    add_chart(
                        "E-TDAH Resultados",
                        cls._etdah_chart(tests.get("etdah_ad") or tests.get("etdah_pais")),
                        template_key="etdah_ad" if tests.get("etdah_ad") else "etdah_pais",
                    )
                if tests.get("bai"):
                    add_title("BAI")
                    add_text(cls._bai_description_text())
                    add_table(
                        "BAI - Resultado da sintomatologia ansiosa",
                        cls._bai_rows(tests.get("bai")),
                        "scale_summary",
                    )
                    add_text(section_or_test_interpretation("bai", None, tests.get("bai")))
                for scared_test in cls._find_tests(context, "scared"):
                    form_label = cls._scared_form_label(scared_test)
                    add_table(
                        f"SCARED - Resultados {form_label}",
                        cls._scared_rows(scared_test),
                        cls._scared_table_key(scared_test)
                    )
                    anchor = cls._insert_interpretation_block_after(
                        anchor,
                        cls._normalize_interpretation_text(
                            cls._resolve_interpretation_text(None, None, scared_test)
                        ),
                    )
                    add_chart(
                        cls._scared_form_title(scared_test),
                        cls._scared_chart(scared_test),
                        template_key="scared_pair",
                    )
                if tests.get("epq_j"):
                    add_table(
                        "EPQ-J Resultados da personalidade",
                        cls._epq_rows(tests.get("epq_j")),
                        "epq",
                    )
                    add_chart(
                        "EPQ-J - Percentis",
                        cls._epq_chart(tests.get("epq_j")),
                        template_key="epq",
                    )
                if tests.get("bfp"):
                    add_title("BFP – Bateria Fatorial de Personalidade")
                    add_text(cls._bfp_description_text())
                    for table_rows in cls._bfp_rows(tests.get("bfp")) or []:
                        add_table("BFP Resultados dos fatores", table_rows, "bfp")
                    add_text(
                        section_or_test_interpretation(
                            "bfp",
                            "aspectos_emocionais_comportamentais",
                            tests.get("bfp"),
                        )
                    )
                if tests.get("srs2"):
                    add_table("SRS-2 Resultados", cls._srs2_rows(tests.get("srs2")), "srs2")
                    add_chart(
                        "SRS-2 Resultados",
                        cls._srs2_chart(tests.get("srs2")),
                        template_key="srs2",
                    )
                if tests.get("ebadep_a") or tests.get("ebadep_ij") or tests.get("ebaped_ij"):
                    ebadep_test = tests.get("ebadep_a") or tests.get("ebadep_ij") or tests.get("ebaped_ij")
                    add_title("EBADEP-A")
                    add_text(cls._ebadep_description_text())
                    add_table(
                        "Resultados da EBADEP",
                        cls._ebadep_rows(ebadep_test),
                        "scale_summary",
                    )
                    add_text(
                        cls._resolve_interpretation_text(
                            sections.get("aspectos_emocionais_comportamentais"),
                            None,
                            ebadep_test,
                        )
                    )

    @classmethod
    def _find_paragraph(cls, document: Document, text: str | None):
        if not text:
            return None
        for paragraph in document.paragraphs:
            if paragraph.text.strip() == text:
                return paragraph
        return None

    @classmethod
    def _rename_heading_if_present(cls, document: Document, old_text: str, new_text: str):
        paragraph = cls._find_paragraph(document, old_text)
        if paragraph is None:
            return
        paragraph.text = new_text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = cls.HEADING_SPACE_BEFORE
        paragraph.paragraph_format.space_after = cls.HEADING_SPACE_AFTER
        paragraph.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        paragraph.paragraph_format.first_line_indent = Pt(0)
        for run in paragraph.runs:
            run.font.name = cls.FONT_NAME
            run.font.size = cls.TITLE_SIZE
            run.bold = True

    @classmethod
    def _remove_nodes_between(cls, start_paragraph, end_paragraph):
        current = start_paragraph._p.getnext()
        end_element = end_paragraph._p if end_paragraph is not None else None
        while current is not None and current is not end_element:
            nxt = current.getnext()
            current.getparent().remove(current)
            current = nxt

    @classmethod
    def _insert_paragraph_after(cls, paragraph, text: str):
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_paragraph = Paragraph(new_p, paragraph._parent)
        if text is not None:
            new_paragraph.add_run(text)
        return new_paragraph

    @classmethod
    def _append_body_element_before_sectpr(cls, document: Document, element):
        body = document._body._element
        sect_pr = next((child for child in body.iterchildren() if child.tag.endswith('sectPr')), None)
        copied = deepcopy(element)
        if sect_pr is None:
            body.append(copied)
        else:
            sect_pr.addprevious(copied)
        return copied

    @classmethod
    def _insert_table_after(cls, paragraph, rows: list[list[str]], table_key: str | None = None):
        parent = paragraph._parent
        table_rows = cls._with_table_title(rows, table_key)
        table = parent.add_table(rows=0, cols=len(table_rows[0]), width=cls.DEFAULT_TABLE_WIDTH)
        tbl = table._tbl
        tbl.getparent().remove(tbl)
        paragraph._p.addnext(tbl)
        for row_values in table_rows:
            row = table.add_row().cells
            for index, value in enumerate(row_values):
                row[index].text = value
        return Table(tbl, parent)

    @classmethod
    def _insert_ravlt_table_after(cls, paragraph, rows: list[list[str]]):
        table = cls._insert_table_after(paragraph, rows, "ravlt")
        cls._format_ravlt_table(table)
        return table

    @classmethod
    def _with_table_title(cls, rows: list[list[str]] | None, table_key: str | None):
        if not rows or not table_key:
            return rows
        title = cls._table_title_text(table_key)
        if not title:
            return rows
        return [[title, *([""] * (len(rows[0]) - 1))], *rows]

    @classmethod
    def _table_title_text(cls, table_key: str) -> str | None:
        return {
            "bpa": "BPA-2",
            "fdt": "FDT - TESTE DOS CINCO DÍGITOS",
            "etdah": "E-TDAH-PAIS",
            "etdah_pais": "E-TDAH-PAIS",
            "etdah_ad": "E-TDAH-AD",
            "scared": "SCARED",
            "scared_parent": "SCARED",
            "scared_self": "SCARED - AUTORRELATO",
            "epq": "EPQ-J",
            "srs2": "SRS-2",
        }.get(table_key)

    @staticmethod
    def _is_scared_table_key(table_key: str | None) -> bool:
        return table_key in {"scared", "scared_parent", "scared_self"}

    @classmethod
    def _merge_title_row(cls, row):
        merged_text = row.cells[0].text.strip()
        merged = row.cells[0]
        for idx in range(1, len(row.cells)):
            merged = merged.merge(row.cells[idx])
        merged.text = merged_text

    @classmethod
    def _insert_paragraph_after_table(cls, table, text: str):
        new_p = OxmlElement("w:p")
        table._tbl.addnext(new_p)
        paragraph = Paragraph(new_p, table._parent)
        if text is not None:
            paragraph.add_run(text)
        return paragraph

    @classmethod
    def _clear_document_body(cls, document: Document):
        body = document._body._element
        for child in list(body):
            if child.tag.endswith("sectPr"):
                continue
            body.remove(child)

    @classmethod
    def _format_body_paragraph(cls, paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = cls.BODY_SPACE_AFTER
        paragraph.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        paragraph.paragraph_format.first_line_indent = cls.BODY_FIRST_LINE_INDENT
        for run in paragraph.runs:
            run.font.name = cls.FONT_NAME
            run.font.size = cls.BODY_SIZE

    @classmethod
    def _format_left_body_paragraph(cls, paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = cls.BODY_SPACE_AFTER
        paragraph.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        paragraph.paragraph_format.first_line_indent = cls.BODY_FIRST_LINE_INDENT
        for run in paragraph.runs:
            run.font.name = cls.FONT_NAME
            run.font.size = cls.BODY_SIZE

    @classmethod
    def _append_body_text_with_bold_label(cls, paragraph, text: str):
        stripped = text.lstrip()
        leading = text[: len(text) - len(stripped)]

        if leading:
            run = paragraph.add_run(leading)
            run.font.name = cls.FONT_NAME
            run.font.size = cls.BODY_SIZE

        for label in cls.INTERPRETATION_LABELS:
            if stripped.casefold().startswith(label.casefold()):
                bold_run = paragraph.add_run(stripped[: len(label)])
                bold_run.font.name = cls.FONT_NAME
                bold_run.font.size = cls.BODY_SIZE
                bold_run.bold = True

                remainder = stripped[len(label) :]
                if remainder:
                    run = paragraph.add_run(remainder)
                    run.font.name = cls.FONT_NAME
                    run.font.size = cls.BODY_SIZE
                return

        run = paragraph.add_run(stripped)
        run.font.name = cls.FONT_NAME
        run.font.size = cls.BODY_SIZE

    @classmethod
    def _append_wisc_intro_paragraph(cls, paragraph, text: str):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = cls.BODY_SPACE_AFTER
        paragraph.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        paragraph.paragraph_format.first_line_indent = cls.BODY_FIRST_LINE_INDENT

        qit_match = re.search(r"QI Total \(QIT [^)]+\)", text)
        class_match = re.search(r"ficando na classificação ([^,.]+)", text)
        age_phrase_match = re.search(
            r"quando comparado à média geral e com idade cognitiva estimada de [^.]+",
            text,
        )

        bold_segments = [
            "Capacidade Cognitiva Global:",
            "WISC IV",
        ]
        if qit_match:
            bold_segments.append(qit_match.group(0))
        if class_match:
            bold_segments.append(class_match.group(1).strip())
        if age_phrase_match:
            bold_segments.append(age_phrase_match.group(0))

        cursor = 0
        for segment in sorted(set(bold_segments), key=lambda item: text.find(item)):
            index = text.find(segment, cursor)
            if index < 0:
                continue
            if index > cursor:
                run = paragraph.add_run(text[cursor:index])
                run.font.name = cls.FONT_NAME
                run.font.size = cls.BODY_SIZE
            run = paragraph.add_run(segment)
            run.font.name = cls.FONT_NAME
            run.font.size = cls.BODY_SIZE
            run.bold = True
            cursor = index + len(segment)

        if cursor < len(text):
            run = paragraph.add_run(text[cursor:])
            run.font.name = cls.FONT_NAME
            run.font.size = cls.BODY_SIZE

    @staticmethod
    def _wisc_payload(test: dict | None) -> dict:
        return (
            (test or {}).get("structured_results")
            or (test or {}).get("classified_payload")
            or (test or {}).get("computed_payload")
            or {}
        )

    @classmethod
    def _wisc_global_intro_text(cls, test: dict | None, context: dict) -> str:
        payload = cls._wisc_payload(test)
        qit = payload.get("qit_data") or {}
        patient_name = cls._patient_reference_name(context or {})
        intro = (
            f"Capacidade Cognitiva Global: a avaliação neuropsicológica de {patient_name}, por meio da Escala Wechsler de Inteligência para Crianças – Quarta Edição (WISC-IV), possibilitou a análise do funcionamento intelectual global e dos principais domínios cognitivos. {patient_name} apresentou Quociente de Inteligência Total (QIT = {qit.get('escore_composto', 'não informado')}), classificado como {qit.get('classificacao', 'não informada')}"
        )
        idade_cognitiva = payload.get("idade_cognitiva")
        if idade_cognitiva:
            intro += (
                f", com idade cognitiva estimada de {idade_cognitiva}"
            )
        intro += "."
        return intro

    @classmethod
    def _wisc_model_interpretation_text(cls, test: dict | None, context: dict) -> str:
        interpretation = cls._resolve_interpretation_text(
            None,
            None,
            test,
            context,
        )
        normalized = cls._normalize_interpretation_text(interpretation)
        if normalized != "Interpretação e Observações Clínicas:":
            return normalized

        payload = cls._wisc_payload(test)
        patient_name = cls._patient_reference_name(context or {})
        generated = interpret_wisc4_profile(payload, patient_name)
        body = re.sub(r"^Interpretação e Observações Clínicas\s*", "", generated).strip()
        body = re.sub(r"\n\s*\n+", " ", body)
        return cls._normalize_interpretation_text(body)

    @classmethod
    def _wisc_global_bullet_parts(cls, test: dict | None) -> list[tuple[str, str]]:
        payload = cls._wisc_payload(test)
        indices = {item.get("indice"): item for item in payload.get("indices") or []}
        definitions = [
            ("Compreensão Verbal", "ICV", indices.get("icv"), "Avaliou o conhecimento verbal adquirido, o processo mental necessário para responder às questões formuladas, a capacidade de compreensão verbal e o raciocínio verbal."),
            ("Organização Perceptual", "IOP", indices.get("iop"), "Avaliou o raciocínio não verbal, a atenção para detalhes e a integração visomotora."),
            ("Memória Operacional", "IMO", indices.get("imt"), "Avaliou a atenção, a concentração e a memória operacional."),
            ("Velocidade de Processamento", "IVP", indices.get("ivp"), "Avaliou a capacidade de em realizar tarefas que demandam rapidez e precisão na análise de estímulos visuais."),
        ]
        rows: list[tuple[str, str]] = []
        for label, code, item, description in definitions:
            item = item or {}
            lead = (
                f"{label} ({code}) — {item.get('escore_composto', 'não informado')} — "
                f"{item.get('classificacao', 'não informada')}"
            )
            rows.append((lead, description))
        return rows

    @classmethod
    def _wais3_payload(cls, test: dict | None) -> dict:
        if not test or not isinstance(test, dict):
            return {}
        payload = (
            test.get("structured_results")
            or test.get("classified_payload")
            or test.get("computed_payload")
            or {}
        )
        return payload if isinstance(payload, dict) else {}

    @classmethod
    def _wais3_intro_text(cls, test: dict | None, context: dict) -> str:
        payload = cls._wais3_payload(test)
        indices = payload.get("indices") or {}
        qit = indices.get("qi_total") or {}
        patient_name = cls._patient_reference_name(context or {})
        if qit.get("pontuacao_composta") is None:
            return (
                f"Capacidade Cognitiva Global: {patient_name} realizou a Escala Wechsler de Inteligência para Adultos – Terceira Edição (WAIS-III), instrumento destinado à avaliação do funcionamento intelectual global em adultos. No momento, as tabelas normativas de conversão de pontos brutos em escores ponderados ainda não estão preenchidas de forma suficiente para cálculo automatizado dos índices compostos, percentis e intervalos de confiança."
            )
        return (
            f"Capacidade Cognitiva Global: {patient_name} obteve, a partir da Escala Wechsler de Inteligência para Adultos – Terceira Edição (WAIS-III), "
            f"Quociente de Inteligência Total (QIT = {qit.get('pontuacao_composta', 'não informado')}), permanecendo na classificação {qit.get('classificacao', 'não informada')} quando comparado à média geral da população normativa. "
            "Em relação aos índices fatoriais, apresentou os seguintes resultados:"
        )

    @classmethod
    def _wasi_intro_text(cls, test: dict | None, context: dict) -> str:
        if not test:
            return ""
        payload = cls._wasi_payload(test)
        composites = payload.get("composites") or {}
        qit = composites.get("qit_4") or composites.get("qit_2") or {}
        qit_value = qit.get("qi")
        qit_classification = qit.get("classification", "não informada")
        patient_name = cls._patient_reference_name(context or {})
        if qit_value is None:
            return (
                f"Capacidade Cognitiva Global: {patient_name} realizou a Escala Wechsler Abreviada de Inteligência (WASI), instrumento destinado à avaliação do funcionamento intelectual global. "
                "No momento, os dados necessários para o cálculo automatizado dos quocientes intelectuais ainda não estão disponíveis."
            )
        return (
            f"Capacidade Cognitiva Global: {patient_name} obteve, a partir da Escala Wechsler Abreviada de Inteligência (WASI), "
            f"QI Total (QIT = {qit_value}), ficando na classificação {qit_classification}, quando comparado à média geral. "
            "Em relação aos índices fatoriais (medidas mais apuradas da inteligência), apresentou os seguintes resultados:"
        )

    @classmethod
    def _wasi_global_bullet_parts(cls, test: dict | None) -> list[tuple[str, str]]:
        if not test:
            return []
        payload = cls._wasi_payload(test)
        composites = payload.get("composites") or {}

        qiv = composites.get("qi_verbal") or {}
        qie = composites.get("qi_execucao") or {}
        qit = composites.get("qit_4") or composites.get("qit_2") or {}

        definitions = [
            ("QIV (Coeficiente de Inteligência Verbal)", qiv, "Avalia os processos verbais e de conhecimento adquirido, tendo uma maior semelhança com o conceito de inteligência cristalizada. No WASI, ele é composto dos subtestes Semelhanças e Vocabulário."),
            ("QIE (Coeficiente de Inteligência de Execução)", qie, "Avalia a organização perceptual, capacidade de manipular estímulos visuais com rapidez e velocidade, e outros processos não verbais, assumindo maior proximidade com o conceito de inteligência fluida. No WASI, ele é composto dos subtestes Cubos e Raciocínio Matricial."),
            ("QIT (Coeficiente de Inteligência Total)", qit, "Avalia o nível geral do funcionamento intelectual."),
        ]

        rows: list[tuple[str, str]] = []
        for label, item, description in definitions:
            if not item:
                continue
            qi = item.get("qi")
            classification = item.get("classification", "")
            value_str = f"{qi} — {classification}" if qi and classification else str(qi) if qi else "-"
            rows.append((f"{label} — {value_str}", description))
        return rows

    @classmethod
    def _wais3_global_bullet_parts(cls, test: dict | None) -> list[tuple[str, str]]:
        payload = cls._wais3_payload(test)
        indices = payload.get("indices") or {}
        definitions = [
            ("Compreensão Verbal (ICV)", indices.get("compreensao_verbal"), "Avaliou o conhecimento verbal adquirido, o raciocínio verbal, a formação de conceitos e a compreensão verbal."),
            ("Organização Perceptual (IOP)", indices.get("organizacao_perceptual"), "Avaliou o raciocínio não verbal, a organização perceptual, a análise visuoespacial e a solução de problemas com estímulos visuais."),
            ("Memória Operacional (IMO)", indices.get("memoria_operacional"), "Avaliou atenção, concentração, memória operacional e manipulação mental de informações."),
            ("Velocidade de Processamento (IVP)", indices.get("velocidade_processamento"), "Avaliou rapidez, precisão, eficiência visuomotora e velocidade em tarefas simples e automatizadas."),
            ("Quociente Intelectual Verbal (QIV)", indices.get("qi_verbal"), "Avaliou recursos verbais globais, compreensão, expressão verbal e raciocínio mediado pela linguagem."),
            ("Quociente de Execução (QIE)", indices.get("qi_execucao"), "Avaliou raciocínio não verbal, organização visuoespacial, atenção a detalhes e solução prática de problemas."),
            ("Índice de Habilidade Geral (GAI)", indices.get("gai"), "Estimou habilidades intelectuais gerais com menor influência da memória operacional e da velocidade de processamento."),
        ]
        rows: list[tuple[str, str]] = []
        for label, item, description in definitions:
            item = item or {}
            score = item.get("pontuacao_composta")
            classification = item.get("classificacao")
            if score is None and item.get("subtestes_ausentes"):
                lead = f"{label} — não calculado"
            else:
                lead = f"{label} — {score if score is not None else 'não informado'} — {classification or 'não informada'}"
            rows.append((lead, description))
        return rows

    @classmethod
    def _wais3_indices_rows(cls, test: dict | None):
        payload = cls._wais3_payload(test)
        indices = payload.get("indices") or {}
        if not isinstance(indices, dict):
            indices = {}
        rows = [["Índice", "Soma Ponderada", "Pontuação Composta", "Percentil", "IC 95%", "Classificação"]]
        for key in (
            "qi_total",
            "qi_verbal",
            "qi_execucao",
            "compreensao_verbal",
            "organizacao_perceptual",
            "memoria_operacional",
            "velocidade_processamento",
        ):
            item = indices.get(key) or {}
            if not item or not isinstance(item, dict):
                continue
            rows.append([
                item.get("nome") or key,
                cls._num(item.get("soma_ponderada")),
                cls._num(item.get("pontuacao_composta")),
                cls._num(item.get("percentil")),
                item.get("ic_95") or item.get("ic_90") or "-",
                item.get("classificacao") or ("Não calculado" if item.get("subtestes_ausentes") else "-"),
            ])
        return rows if len(rows) > 1 else None

    @classmethod
    def _wais3_subtests_rows(cls, test: dict | None):
        payload = cls._wais3_payload(test)
        subtests = payload.get("subtestes") or {}
        if not isinstance(subtests, dict):
            subtests = {}
        rows = [["Subteste", "Pontos Brutos", "Escore Ponderado", "Classificação", "Observação"]]
        for key, item in subtests.items():
            if not isinstance(item, dict):
                continue
            rows.append([
                item.get("nome") or key,
                cls._num(item.get("pontos_brutos")),
                cls._num(item.get("escore_ponderado")),
                item.get("classificacao") or "-",
                item.get("warning") or "",
            ])
        return rows if len(rows) > 1 else None

    @classmethod
    def _wais3_domain_rows(cls, test: dict | None, section_key: str):
        table_rows = (test or {}).get("wais3_tables") or {}
        domain_rows = table_rows.get(section_key) or []
        rows = [[
            "Testes Utilizados",
            "Escore Máximo",
            "Escore Médio",
            "Escore Mínimo",
            "Escore Bruto",
            "Classificação",
        ]]
        for item in domain_rows:
            if item.get("note"):
                rows.append([
                    item.get("label") or "-",
                    item.get("note") or "-",
                    "",
                    "",
                    "",
                    "",
                ])
                continue
            rows.append(
                [
                    item.get("label") or "-",
                    item.get("maxScore") or "-",
                    item.get("avgScore") or "-",
                    item.get("minScore") or "-",
                    item.get("obtainedScore") or "-",
                    item.get("classification") or "-",
                ]
            )
        return rows if len(rows) > 1 else None

    @classmethod
    def _populate_wasi_tables(cls, document: Document, context: dict):
        validated = context.get('validated_tests') or []
        tests = {}
        for item in validated:
            if isinstance(item, dict):
                code = item.get('instrument_code')
                if code:
                    tests[code] = item

        for table in document.tables:
            if not table.rows:
                continue
            header_row = table.rows[0]
            cells_text = [c.text.strip() for c in header_row.cells]

            if 'Escore Máximo' in cells_text and len(table.columns) >= 6:
                subtest_names = [c.text.strip().lower() for c in (table.rows[1].cells if len(table.rows) > 1 else [])]

                if any('vocabul' in name for name in subtest_names if name):
                    wasi = tests.get('wasi')
                    if wasi:
                        payload = cls._wasi_payload(wasi)
                        subtests = payload.get('subtests') or {}
                        vc = subtests.get('vc') or {}
                        sm = subtests.get('sm') or {}
                        if len(table.rows) >= 2:
                            table.rows[1].cells[4].text = str(vc.get('t_score') or vc.get('raw_score') or '-')
                            table.rows[1].cells[5].text = vc.get('classification') or '-'
                        if len(table.rows) >= 3:
                            table.rows[2].cells[4].text = str(sm.get('t_score') or sm.get('raw_score') or '-')
                            table.rows[2].cells[5].text = sm.get('classification') or '-'

                elif any('cubo' in name for name in subtest_names if name):
                    wasi = tests.get('wasi')
                    if wasi:
                        payload = cls._wasi_payload(wasi)
                        subtests = payload.get('subtests') or {}
                        cb = subtests.get('cb') or {}
                        rm = subtests.get('rm') or {}
                        if len(table.rows) >= 2:
                            table.rows[1].cells[4].text = str(cb.get('t_score') or cb.get('raw_score') or '-')
                            table.rows[1].cells[5].text = cb.get('classification') or '-'
                        if len(table.rows) >= 3:
                            table.rows[2].cells[4].text = str(rm.get('t_score') or rm.get('raw_score') or '-')
                            table.rows[2].cells[5].text = rm.get('classification') or '-'

            elif 'ATENÇÃO BPA' in cells_text or cells_text[0] == 'ATENÇÃO BPA':
                bpa = tests.get('bpa2')
                if bpa:
                    sr = bpa.get('structured_results') or {}
                    subtestes = sr.get('subtestes') or []
                    subtest_map = {st.get('codigo'): st for st in subtestes}

                    code_map = [('ac', 1), ('ad', 2), ('aa', 3), ('ag', 4)]
                    for code, row_idx in code_map:
                        if row_idx < len(table.rows) and len(table.rows[row_idx].cells) >= 4:
                            st = subtest_map.get(code) or {}
                            table.rows[row_idx].cells[1].text = str(st.get('total', '-'))
                            table.rows[row_idx].cells[2].text = str(st.get('percentil', '-'))
                            table.rows[row_idx].cells[3].text = st.get('classificacao', '-')

            elif 'TESTE DOS CINCO DÍGITOS' in ' '.join(cells_text):
                fdt = tests.get('fdt')
                if fdt:
                    computed = fdt.get('computed_payload') or {}
                    metric_results = computed.get('metric_results') or []
                    stage_totals = computed.get('stage_totals') or {}

                    code_map = [
                        (2, 'leitura'),
                        (3, 'contagem'),
                        (4, 'escolha'),
                        (5, 'alternancia'),
                        (6, 'inibicao'),
                        (7, 'flexibilidade'),
                    ]
                    for row_idx, key in code_map:
                        st = next((m for m in metric_results if isinstance(m, dict) and m.get('codigo', '').lower() == key), None)
                        stage = stage_totals.get(key, {}) if isinstance(stage_totals, dict) else {}
                        if st and row_idx < len(table.rows) and len(table.rows[row_idx].cells) >= 6:
                            row = table.rows[row_idx]
                            row.cells[2].text = cls._fmt_fdt_num(stage.get('tempo') if isinstance(stage, dict) else None) or cls._fmt_fdt_num(st.get('valor'))
                            row.cells[3].text = str(stage.get('erros', '0') if isinstance(stage, dict) else '0')
                            row.cells[4].text = cls._fmt_fdt_num(st.get('percentil_num')) if st.get('percentil_num') is not None else '-'
                            row.cells[5].text = st.get('classificacao', '-')

            elif 'Desempenho' in cells_text and 'A1' in cells_text:
                ravlt = tests.get('ravlt')
                if ravlt:
                    classified = ravlt.get('classified_payload') or {}
                    computed = ravlt.get('computed_payload') or {}
                    resultados = classified.get('resultados') or computed.get('resultados') or []
                    result_map = {r.get('variavel'): r for r in resultados if isinstance(r, dict)}

                    chart_data = classified.get('chart') or computed.get('chart') or {}
                    series_list = chart_data.get('series', [])
                    obtenido_row = next((s for s in series_list if isinstance(s, dict) and s.get('key') == 'obtido'), None)
                    obtenido_vals = (obtenido_row.get('values', []) if isinstance(obtenido_row, dict) else []) if obtenido_row else []

                    labels = ['A1', 'A2', 'A3', 'A4', 'A5', 'B1', 'A6', 'A7', 'R', 'ALT', 'RET', 'I.P.', 'I.R.']
                    if len(table.rows) >= 4 and len(table.rows[3].cells) >= len(labels) + 1:
                        row = table.rows[3]
                        for col_idx, label in enumerate(labels, 1):
                            if label in result_map:
                                r = result_map[label]
                                val = r.get('bruto') if r.get('bruto') is not None else r.get('ponderado')
                                if val is not None:
                                    row.cells[col_idx].text = str(int(val)) if val == int(val) else f'{val:.2f}'.replace('.', ',')
                            elif col_idx - 1 < len(obtenido_vals):
                                val = obtenido_vals[col_idx - 1]
                                if val is not None:
                                    row.cells[col_idx].text = str(int(val)) if val == int(val) else f'{val:.2f}'.replace('.', ',')

    @classmethod
    def _fmt_fdt_num(cls, value):
        if value is None:
            return None
        try:
            return str(round(float(value), 2)).replace('.', ',')
        except (TypeError, ValueError):
            return str(value)

    @classmethod
    def _merge_nested_dicts(cls, base: dict, override: dict) -> dict:
        merged = dict(base or {})
        for key, value in (override or {}).items():
            current = merged.get(key)
            if isinstance(current, dict) and isinstance(value, dict):
                merged[key] = cls._merge_nested_dicts(current, value)
            else:
                merged[key] = value
        return merged

    @classmethod
    def _wasi_payload(cls, test: dict | None) -> dict:
        test = test or {}
        merged = {}
        for source in (
            test.get("structured_results") or {},
            test.get("computed_payload") or {},
        ):
            if isinstance(source, dict):
                merged = cls._merge_nested_dicts(merged, source)
        return merged

    @classmethod
    def _wasi_subscale_rows(cls, test: dict | None, scale: str):
        payload = cls._wasi_payload(test)
        subtests = payload.get("subtests") or {}
        if not isinstance(subtests, dict):
            subtests = {}

        scale_codes = {
            "verbal": ["vc", "sm"],
            "execucao": ["cb", "rm"],
        }
        rows = [[
            "Testes Utilizados",
            "Escore Máximo",
            "Escore Médio",
            "Escore Mínimo",
            "Escore Obtido",
            "Classificação",
        ]]
        for code in scale_codes.get(scale, []):
            item = subtests.get(code) or {}
            if not item:
                continue
            rows.append([
                item.get("name") or code.upper(),
                "80",
                "40 - 60",
                "20",
                cls._num(item.get("t_score") if item.get("t_score") is not None else item.get("raw_score")),
                item.get("classification") or "-",
            ])
        return rows if len(rows) > 1 else None

    @classmethod
    def _wasi_subscale_item(cls, test: dict | None, code: str) -> dict:
        payload = cls._wasi_payload(test)
        subtests = payload.get("subtests") or {}
        if not isinstance(subtests, dict):
            return {}
        return subtests.get(code) or {}

    @classmethod
    def _wasi_verbal_interpretation_text(cls, test: dict | None, context: dict) -> str:
        patient_name = cls._patient_reference_name(context or {})
        semelhancas = cls._wasi_subscale_item(test, "sm")
        vocabulario = cls._wasi_subscale_item(test, "vc")
        sem_class = semelhancas.get("classification", "não informada")
        voc_class = vocabulario.get("classification", "não informada")
        return (
            f"Interpretação: A avaliação da escala verbal de {patient_name} foi realizada por meio dos subtestes Vocabulário e Semelhanças da Escala Wechsler Abreviada de Inteligência – WASI, permitindo examinar aspectos relacionados ao raciocínio verbal, amplitude lexical, conhecimento semântico e capacidade de abstração conceitual.\n"
            f"No subteste Vocabulário, {patient_name} apresentou desempenho classificado na faixa {voc_class.lower()}, evidenciando repertório lexical adequado, capacidade funcional de nomeação, compreensão semântica e definição de palavras. Esse resultado sugere domínio preservado do conhecimento verbal adquirido, bem como organização do pensamento mediado pela linguagem, favorecendo a comunicação e a compreensão de conteúdos verbais estruturados.\n"
            f"No subteste Semelhanças, observou-se desempenho na faixa {sem_class.lower()}, indicando capacidade de abstração verbal, formação de conceitos e identificação de relações entre estímulos. Esse desempenho reflete funcionamento em processos de categorização, raciocínio lógico verbal e integração conceitual.\n"
            f"Em análise clínica, os resultados da Escala Verbal indicam funcionamento compatível com o perfil evidenciado nos subtestes avaliados, sugerindo recursos de compreensão verbal, raciocínio abstrato mediado pela linguagem e organização conceitual em nível coerente com o desempenho global observado no WASI. Esse perfil favorece a aprendizagem baseada em instruções verbais e a adaptação a contextos que demandam comunicação estruturada."
        )

    @classmethod
    def _wasi_execution_interpretation_text(cls, test: dict | None, context: dict) -> str:
        patient_name = cls._patient_reference_name(context or {})
        cubos = cls._wasi_subscale_item(test, "cb")
        matricial = cls._wasi_subscale_item(test, "rm")
        cubos_class = cubos.get("classification", "não informada")
        matricial_class = matricial.get("classification", "não informada")
        return (
            f"Interpretação: A avaliação da escala de execução de {patient_name} foi realizada por meio dos subtestes Cubos e Raciocínio Matricial da Escala Wechsler Abreviada de Inteligência – WASI, permitindo examinar habilidades de raciocínio visuoespacial, análise de padrões, integração perceptual e organização motora.\n"
            f"No subteste Cubos, {patient_name} apresentou desempenho classificado na faixa {cubos_class.lower()}, indicando organização visuoespacial, análise perceptiva e coordenação visomotora compatíveis com o perfil observado. Esse resultado sugere capacidade funcional de estruturar estímulos visuais, integrar partes em um todo e executar tarefas com demanda construtiva.\n"
            f"No subteste Raciocínio Matricial, o desempenho na faixa {matricial_class.lower()} evidencia raciocínio lógico não verbal, capacidade de identificar padrões, inferir regras e resolver problemas abstratos. Esse desempenho indica funcionamento em tarefas que exigem análise visual e raciocínio fluido.\n"
            f"Em análise clínica, os resultados da Escala de Execução indicam funcionamento compatível com o perfil evidenciado nos subtestes avaliados, sugerindo recursos de raciocínio não verbal, percepção de padrões, organização visuoespacial e resolução de problemas abstratos em nível coerente com o desempenho global observado no WASI."
        )

    @staticmethod
    def _strip_markdown_heading_prefix(text: str | None, heading: str) -> str:
        if not text:
            return ""
        normalized = str(text).strip()
        for variant in (f"**{heading}**", heading):
            if normalized.startswith(variant):
                normalized = normalized[len(variant):].lstrip("\n :.-")
        return normalized

    @classmethod
    def _wais3_chart(cls, test: dict | None):
        payload = cls._wais3_payload(test)
        indices = payload.get("indices") or {}
        items = [
            ("ICV", indices.get("compreensao_verbal")),
            ("IOP", indices.get("organizacao_perceptual")),
            ("IMO", indices.get("memoria_operacional")),
            ("IVP", indices.get("velocidade_processamento")),
            ("QIV", indices.get("qi_verbal")),
            ("QIE", indices.get("qi_execucao")),
            ("QIT", indices.get("qi_total")),
        ]
        labels = []
        values = []
        for label, item in items:
            score = (item or {}).get("pontuacao_composta")
            if score is None:
                continue
            labels.append(label)
            values.append(float(score))
        return cls._build_chart_png("bar", "Desempenho do paciente no WAIS III", labels, values, "Pontuação Composta")

    @classmethod
    def _append_wisc_global_bullet(cls, paragraph, lead: str, tail: str):
        bullet_prefix = paragraph.add_run("-\t")
        bullet_prefix.font.name = cls.FONT_NAME
        bullet_prefix.font.size = cls.BODY_SIZE
        bold_run = paragraph.add_run(lead)
        bold_run.font.name = cls.FONT_NAME
        bold_run.font.size = cls.BODY_SIZE
        bold_run.bold = True
        tail_run = paragraph.add_run(f" {tail}")
        tail_run.font.name = cls.FONT_NAME
        tail_run.font.size = cls.BODY_SIZE
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = cls.BODY_SPACE_AFTER
        paragraph.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        paragraph.paragraph_format.left_indent = Cm(1.1)
        paragraph.paragraph_format.first_line_indent = Cm(-0.55)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(1.1))

    @classmethod
    def _append_wisc_global_block(cls, document, test: dict | None, context: dict):
        intro = cls._wisc_global_intro_text(test, context)
        p = document.add_paragraph()
        cls._append_wisc_intro_paragraph(p, intro)

    @classmethod
    def _append_wisc_indices_block(cls, document, test: dict | None):
        for lead, tail in cls._wisc_global_bullet_parts(test):
            p = document.add_paragraph()
            cls._append_wisc_global_bullet(p, lead, tail)

    @classmethod
    def _format_subtitle_paragraph(cls, paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = cls.SUBHEADING_SPACE_BEFORE
        paragraph.paragraph_format.space_after = cls.SUBHEADING_SPACE_AFTER
        paragraph.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.first_line_indent = Pt(0)
        for run in paragraph.runs:
            run.font.name = cls.FONT_NAME
            run.font.size = cls.TITLE_SIZE
            run.bold = True

    @classmethod
    def _format_caption_paragraph(cls, paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            run.font.name = cls.FONT_NAME
            run.font.size = cls.CAPTION_SIZE
            run.italic = True

    @classmethod
    def _table_caption_text(cls, table_index: int, caption: str) -> str:
        return f"Tabela {table_index} {caption}"

    @classmethod
    def _chart_caption_text(cls, chart_index: int, caption: str) -> str:
        return f"Gráfico {chart_index} {caption}"

    @classmethod
    def _ravlt_table_caption(cls) -> str:
        return (
            'RAVLT Resultados: referente ao desempenho do paciente durante a prova RAVLT '
            '(memória auditiva e aprendizagem de palavras). Os intervalos "A" se referem '
            'à exposição de uma mesma lista de palavras e o desempenho da memória de curto '
            'prazo. O intervalo "B1" se refere a uma intrusão de uma nova lista de '
            'palavras. O intervalo "A6" se refere à memória após estímulo distrator. O '
            'intervalo "A7" se refere ao desempenho do examinando após um período de 20 '
            'minutos. O intervalo "R" se refere à memória de reconhecimento a partir de '
            'pistas. O intervalo "ALT" se refere ao aprendizado ao longo das tentativas. '
            'O intervalo "RET" se refere à velocidade de esquecimento. O intervalo "I. '
            'P." se refere à interferência de um aprendizado antigo em relação a um novo. '
            'O intervalo "I. R." se refere à interferência de aprendizado novo em relação '
            'a um antigo.'
        )

    @classmethod
    def _add_center_title(cls, document, text: str):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = cls.FONT_NAME
        r.font.size = cls.TITLE_SIZE
        r.bold = True
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = cls.HEADING_SPACE_AFTER
        p.paragraph_format.line_spacing = cls.IDENTIFICATION_LINE_SPACING

    @classmethod
    def _add_center_text(cls, document, text: str):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = cls.FONT_NAME
        r.font.size = cls.CAPTION_SIZE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = cls.IDENTIFICATION_LINE_SPACING

    @classmethod
    def _append_heading(cls, document, text: str):
        p = document.add_paragraph()
        if text.endswith("EFICIÊNCIA COGNITIVA") and ". " in text:
            number, heading_text = text.split(". ", 1)
            runs = [
                p.add_run(f"{number}."),
                p.add_run("\t\t"),
                p.add_run(heading_text),
            ]
        else:
            runs = [p.add_run(text)]
        for run in runs:
            run.font.name = cls.FONT_NAME
            run.font.size = cls.TITLE_SIZE
            run.bold = True
        p.paragraph_format.space_before = cls.HEADING_SPACE_BEFORE
        p.paragraph_format.space_after = cls.HEADING_SPACE_AFTER
        p.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        p.paragraph_format.first_line_indent = Pt(0)

    @classmethod
    def _append_subheading(cls, document, text: str):
        text = PtBrTextService.normalize(text)
        p = document.add_paragraph()
        r = p.add_run(text)
        r.font.name = cls.FONT_NAME
        r.font.size = cls.TITLE_SIZE
        r.bold = True
        p.paragraph_format.space_before = cls.SUBHEADING_SPACE_BEFORE
        p.paragraph_format.space_after = cls.SUBHEADING_SPACE_AFTER
        p.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        p.paragraph_format.first_line_indent = Pt(0)

    @classmethod
    def _append_paragraph(cls, document, text: str, center: bool = False):
        text = PtBrTextService.normalize(text)
        if not text:
            return
        p = document.add_paragraph()
        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
        )
        cls._append_body_text_with_bold_label(p, text)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = cls.BODY_SPACE_AFTER
        p.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        p.paragraph_format.first_line_indent = Pt(0) if center else cls.BODY_FIRST_LINE_INDENT

    @classmethod
    def _append_identification_text(cls, document, text: str):
        text = PtBrTextService.normalize(text)
        p = document.add_paragraph()
        r = p.add_run(text)
        r.font.name = cls.FONT_NAME
        r.font.size = cls.BODY_SIZE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = cls.IDENTIFICATION_LINE_SPACING
        p.paragraph_format.first_line_indent = Pt(0)

    @classmethod
    def _append_label_value(cls, document, label: str, value: str):
        label = PtBrTextService.normalize(label)
        value = PtBrTextService.normalize(str(value or "Não informado"))
        p = document.add_paragraph()
        a = p.add_run(f"{label}: ")
        a.font.name = cls.FONT_NAME
        a.font.size = cls.BODY_SIZE
        a.bold = True
        b = p.add_run(value)
        b.font.name = cls.FONT_NAME
        b.font.size = cls.BODY_SIZE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = cls.IDENTIFICATION_LINE_SPACING
        p.paragraph_format.first_line_indent = Pt(0)

    @classmethod
    def _append_bullet(cls, document, text: str):
        text = PtBrTextService.normalize(text)
        try:
            p = document.add_paragraph(style="List Bullet")
            r = p.add_run(text)
        except KeyError:
            p = document.add_paragraph()
            r = p.add_run(f"• {text}")
        r.font.name = cls.FONT_NAME
        r.font.size = cls.BODY_SIZE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = cls.BODY_SPACE_AFTER
        p.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.25)

    @classmethod
    def _append_procedure_bullet(cls, document, name: str, description: str):
        try:
            p = document.add_paragraph(style="List Bullet")
            cls._append_procedure_bullet_runs(p, name, description, add_bullet_prefix=False)
        except KeyError:
            p = document.add_paragraph()
            cls._append_procedure_bullet_runs(p, name, description, add_bullet_prefix=True)

    @classmethod
    def _append_procedure_bullet_runs(
        cls, paragraph, name: str, description: str, add_bullet_prefix: bool = True
    ):
        name = PtBrTextService.normalize(name)
        description = PtBrTextService.normalize(description)
        if add_bullet_prefix:
            bullet_run = paragraph.add_run("• ")
            bullet_run.font.name = cls.FONT_NAME
            bullet_run.font.size = cls.BODY_SIZE
        label_run = paragraph.add_run(f"{name}: ")
        label_run.font.name = cls.FONT_NAME
        label_run.font.size = cls.BODY_SIZE
        label_run.bold = True
        value_run = paragraph.add_run(description)
        value_run.font.name = cls.FONT_NAME
        value_run.font.size = cls.BODY_SIZE
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = cls.IDENTIFICATION_LINE_SPACING
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)

    @classmethod
    def _append_numbered_item(cls, document, text: str):
        text = PtBrTextService.normalize(text)
        p = document.add_paragraph()
        r = p.add_run(text)
        r.font.name = cls.FONT_NAME
        r.font.size = cls.BODY_SIZE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = cls.BODY_SPACE_AFTER
        p.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.25)

    @classmethod
    def _append_table_with_interpretation(
        cls,
        document,
        rows,
        table_key: str,
        interpretation: str | None,
        caption: str | None = None,
    ):
        if rows:
            table_rows = cls._with_table_title(rows, table_key)
            table = document.add_table(rows=0, cols=len(table_rows[0]))
            for row_values in table_rows:
                row = table.add_row().cells
                for idx, value in enumerate(row_values):
                    row[idx].text = str(value)
            if table_key == "ravlt":
                cls._format_ravlt_table(table)
            else:
                cls._format_table(table, table_key)
            if caption:
                caption_paragraph = document.add_paragraph()
                caption_run = caption_paragraph.add_run(caption)
                caption_run.font.name = cls.FONT_NAME
                caption_run.font.size = Pt(10)
                caption_run.italic = True
                cls._format_caption_paragraph(caption_paragraph)
        if interpretation:
            cls._append_interpretation_block(
                document, cls._normalize_interpretation_text(interpretation)
            )

    @classmethod
    def _append_interpretation_block(cls, document, text: str):
        cleaned = (text or "").strip()
        if not cleaned:
            return

        paragraphs = [
            item.strip()
            for item in re.split(r"\n\s*\n+", cleaned)
            if item.strip()
        ]

        for item in paragraphs:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cls._append_body_text_with_bold_label(p, item)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
            p.paragraph_format.first_line_indent = Pt(0)

    @classmethod
    def _append_ravlt_conceptual_paragraph(cls, document):
        text = cls._ravlt_conceptual_text()
        if not text:
            return
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        r.font.name = cls.FONT_NAME
        r.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.first_line_indent = Pt(0)

    @classmethod
    def _insert_ravlt_conceptual_paragraph_after(cls, anchor):
        anchor = cls._insert_paragraph_after(anchor, cls._ravlt_conceptual_text())
        anchor.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        anchor.paragraph_format.space_before = Pt(0)
        anchor.paragraph_format.space_after = Pt(0)
        anchor.paragraph_format.line_spacing = 1.15
        anchor.paragraph_format.first_line_indent = Pt(0)
        for run in anchor.runs:
            run.font.name = cls.FONT_NAME
            run.font.size = Pt(10)
        return anchor

    @classmethod
    def _insert_interpretation_block_after(cls, anchor, text: str):
        cleaned = (text or "").strip()
        if not cleaned:
            return anchor
        paragraphs = [
            item.strip()
            for item in re.split(r"\n\s*\n+", cleaned)
            if item.strip()
        ]
        for item in paragraphs:
            anchor = cls._insert_paragraph_after(anchor, "")
            cls._append_body_text_with_bold_label(anchor, item)
            anchor.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            anchor.paragraph_format.space_before = Pt(0)
            anchor.paragraph_format.space_after = Pt(0)
            anchor.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
            anchor.paragraph_format.first_line_indent = Pt(0)
        return anchor

    @classmethod
    def _bpa_intro_text(cls, patient_name: str) -> str:
        return (
            f"Interpretação e Observações Clínicas: A avaliação da atenção de {patient_name} foi realizada por meio da Bateria Psicológica para Avaliação da Atenção – BPA-2, instrumento que investiga diferentes aspectos do funcionamento atencional: atenção concentrada, dividida, alternada e atenção geral. Esses domínios estão relacionados ao funcionamento das redes atencionais descritas por Posner e Petersen, envolvendo os sistemas de alerta, orientação e controle executivo."
        )

    @classmethod
    def _bpa_section_title(cls, code: str) -> str:
        titles = {
            "ac": "Atenção Concentrada (AC):",
            "ad": "Atenção Dividida (AD):",
            "aa": "Atenção Alternada (AA):",
            "ag": "Atenção Geral (AG):",
        }
        return titles.get(code, f"{(code or '').upper()}:")

    @classmethod
    def _bpa_section_text(cls, item: dict, patient_name: str) -> str:
        code = (item.get("codigo") or "").lower()
        classificacao = item.get("classificacao") or "Não classificado"
        percentil = item.get("percentil") or 0
        opening = (SUBTEST_OPENINGS.get(code) or {}).get("opening") or ""
        detail_template = (SUBTEST_OPENINGS.get(code) or {}).get(classificacao)
        if detail_template:
            detail = detail_template.format(name=patient_name, percentil=cls._num(percentil))
        else:
            detail = (
                f"{patient_name} apresentou desempenho classificado como {str(classificacao).lower()} "
                f"(percentil {cls._num(percentil)})."
            )
        return PtBrTextService.normalize(f"{opening} {detail}".strip())

    @classmethod
    def _append_bpa_section_heading(cls, document, text: str):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(text)
        run.font.name = cls.FONT_NAME
        run.font.size = cls.BODY_SIZE
        run.bold = True

    @classmethod
    def _append_bpa_body_paragraph(cls, document, text: str):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = cls.BODY_SPACE_AFTER
        p.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        p.paragraph_format.first_line_indent = Cm(1.5)
        p.paragraph_format.left_indent = Pt(0)
        run = p.add_run(PtBrTextService.normalize(text))
        run.font.name = cls.FONT_NAME
        run.font.size = cls.BODY_SIZE

    @classmethod
    def _insert_bpa_section_heading_after(cls, anchor, text: str):
        anchor = cls._insert_paragraph_after(anchor, text)
        anchor.alignment = WD_ALIGN_PARAGRAPH.LEFT
        anchor.paragraph_format.space_before = Pt(0)
        anchor.paragraph_format.space_after = Pt(0)
        anchor.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        anchor.paragraph_format.first_line_indent = Pt(0)
        for run in anchor.runs:
            run.font.name = cls.FONT_NAME
            run.font.size = cls.BODY_SIZE
            run.bold = True
        return anchor

    @classmethod
    def _insert_bpa_body_paragraph_after(cls, anchor, text: str):
        anchor = cls._insert_paragraph_after(anchor, PtBrTextService.normalize(text))
        anchor.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        anchor.paragraph_format.space_before = Pt(0)
        anchor.paragraph_format.space_after = cls.BODY_SPACE_AFTER
        anchor.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        anchor.paragraph_format.first_line_indent = Cm(1.5)
        anchor.paragraph_format.left_indent = Pt(0)
        for run in anchor.runs:
            run.font.name = cls.FONT_NAME
            run.font.size = cls.BODY_SIZE
        return anchor

    @classmethod
    def _append_bpa_interpretation_block(cls, document, test: dict | None, context: dict | None = None):
        payload = (test or {}).get("classified_payload") or {}
        subtests = payload.get("subtestes") or []
        if not subtests:
            return
        patient_name = cls._patient_reference_name(context or {})
        cls._append_paragraph(document, cls._bpa_intro_text(patient_name))
        for code in ("ac", "ad", "aa", "ag"):
            item = next((subtest for subtest in subtests if (subtest.get("codigo") or "").lower() == code), None)
            if not item:
                continue
            cls._append_bpa_section_heading(document, cls._bpa_section_title(code))
            cls._append_bpa_body_paragraph(document, cls._bpa_section_text(item, patient_name))
        cls._append_bpa_section_heading(document, "Análise Clínica:")
        cls._append_bpa_body_paragraph(document, build_clinical_summary(subtests, patient_name))

    @classmethod
    def _insert_bpa_interpretation_block_after(cls, anchor, test: dict | None, context: dict | None = None):
        payload = (test or {}).get("classified_payload") or {}
        subtests = payload.get("subtestes") or []
        if not subtests:
            return anchor
        patient_name = cls._patient_reference_name(context or {})
        anchor = cls._insert_interpretation_block_after(anchor, cls._bpa_intro_text(patient_name))
        for code in ("ac", "ad", "aa", "ag"):
            item = next((subtest for subtest in subtests if (subtest.get("codigo") or "").lower() == code), None)
            if not item:
                continue
            anchor = cls._insert_bpa_section_heading_after(anchor, cls._bpa_section_title(code))
            anchor = cls._insert_bpa_body_paragraph_after(anchor, cls._bpa_section_text(item, patient_name))
        anchor = cls._insert_bpa_section_heading_after(anchor, "Análise Clínica:")
        anchor = cls._insert_bpa_body_paragraph_after(anchor, build_clinical_summary(subtests, patient_name))
        return anchor

    @classmethod
    def _etdah_pais_factor_body(cls, factor_key: str, result: dict) -> str:
        paragraph = PtBrTextService.normalize(build_etdah_pais_factor_paragraph(factor_key, result))
        label = ETDAH_PAIS_FACTOR_LABELS.get(factor_key)
        if not label:
            return paragraph
        prefix = f"{label}:"
        if paragraph.startswith(prefix):
            return paragraph[len(prefix):].strip()
        alt_prefix = f"No {label},"
        if paragraph.startswith(alt_prefix):
            return paragraph[len(alt_prefix):].strip()
        return paragraph

    @classmethod
    def _append_etdah_pais_interpretation_block(cls, document, test: dict | None, context: dict | None = None):
        payload = (test or {}).get("classified_payload") or {}
        results = payload.get("results") or {}
        if not results:
            return
        first_name = cls._patient_reference_name(context or {})
        cls._append_paragraph(
            document,
            f"Interpretação e Observações Clínicas: A avaliação comportamental de {first_name} por meio da Escala E-TDAH-PAIS permitiu investigar aspectos relacionados à regulação emocional, hiperatividade/impulsividade, comportamento adaptativo e atenção, a partir da percepção dos responsáveis, fornecendo subsídios para a compreensão do funcionamento comportamental no contexto familiar e cotidiano.",
        )
        for factor_key in ETDAH_PAIS_FACTOR_ORDER:
            result = results.get(factor_key)
            if not result:
                continue
            cls._append_bpa_section_heading(document, ETDAH_PAIS_FACTOR_LABELS.get(factor_key, factor_key))
            cls._append_bpa_body_paragraph(document, cls._etdah_pais_factor_body(factor_key, result))
        cls._append_bpa_section_heading(document, "Análise Clínica:")
        cls._append_bpa_body_paragraph(document, build_etdah_pais_analysis_text(first_name, results))

    @classmethod
    def _insert_etdah_pais_interpretation_block_after(cls, anchor, test: dict | None, context: dict | None = None):
        payload = (test or {}).get("classified_payload") or {}
        results = payload.get("results") or {}
        if not results:
            return anchor
        first_name = cls._patient_reference_name(context or {})
        anchor = cls._insert_interpretation_block_after(
            anchor,
            f"Interpretação e Observações Clínicas: A avaliação comportamental de {first_name} por meio da Escala E-TDAH-PAIS permitiu investigar aspectos relacionados à regulação emocional, hiperatividade/impulsividade, comportamento adaptativo e atenção, a partir da percepção dos responsáveis, fornecendo subsídios para a compreensão do funcionamento comportamental no contexto familiar e cotidiano.",
        )
        for factor_key in ETDAH_PAIS_FACTOR_ORDER:
            result = results.get(factor_key)
            if not result:
                continue
            anchor = cls._insert_bpa_section_heading_after(anchor, ETDAH_PAIS_FACTOR_LABELS.get(factor_key, factor_key))
            anchor = cls._insert_bpa_body_paragraph_after(anchor, cls._etdah_pais_factor_body(factor_key, result))
        anchor = cls._insert_bpa_section_heading_after(anchor, "Análise Clínica:")
        anchor = cls._insert_bpa_body_paragraph_after(anchor, build_etdah_pais_analysis_text(first_name, results))
        return anchor

    @classmethod
    def _normalize_interpretation_text(cls, interpretation: str) -> str:
        label = "Interpretação e Observações Clínicas:"
        text = PtBrTextService.normalize(
            cls._strip_embedded_caption_lines(
                cls._strip_markdown_emphasis(cls._strip_legacy_srs2_table(interpretation))
            )
        )
        if not text:
            return label
        if text.casefold().startswith(label.casefold()):
            return text
        return f"{label} {text}"

    @staticmethod
    def _strip_embedded_caption_lines(text: str | None) -> str:
        cleaned = str(text or "").strip()
        if not cleaned:
            return ""
        cleaned = re.sub(r"(?im)^\s*(?:Gr[áa]fico|Tabela)\s+\d+\s+.+$", "", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned.strip()

    @staticmethod
    def _strip_markdown_emphasis(text: str | None) -> str:
        cleaned = (text or "").strip()
        if not cleaned:
            return ""
        cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned, flags=re.DOTALL)
        cleaned = re.sub(r"__(.*?)__", r"\1", cleaned, flags=re.DOTALL)
        return cleaned

    @staticmethod
    def _parse_etdah_ad_factor_paragraph(text: str) -> tuple[str, str] | None:
        text_stripped = text.strip()
        
        match = re.match(r"^No\s+(Fator\s+\d+\s+[—-]\s+[^,]+),\s*(.+)$", text_stripped, flags=re.DOTALL)
        if match:
            heading = match.group(1).replace(" - ", " – ").replace(" — ", " – ").strip()
            body = match.group(2).strip()
            return heading, body
        
        match = re.match(r"^(Fator\s+\d+\s+[—-]\s+[^,]+),\s*(.+)$", text_stripped, flags=re.DOTALL)
        if match:
            heading = match.group(1).replace(" - ", " – ").replace(" — ", " – ").strip()
            body = match.group(2).strip()
            return heading, body
        
        return None

    @classmethod
    def _append_etdah_ad_interpretation_block(cls, document, interpretation: str):
        cleaned = (interpretation or "").strip()
        if not cleaned:
            return
        paragraphs = [item.strip() for item in re.split(r"\n\s*\n+", cleaned) if item.strip()]
        for item in paragraphs:
            factor_block = cls._parse_etdah_ad_factor_paragraph(item)
            if factor_block:
                heading, body = factor_block
                cls._append_bpa_section_heading(document, heading)
                cls._append_bpa_body_paragraph(document, body)
                continue
            if item.startswith("Interpretação e Observações Clínicas:"):
                cls._append_interpretation_block(document, item)
                continue
            cls._append_bpa_body_paragraph(document, item)

    @classmethod
    def _insert_etdah_ad_interpretation_block_after(cls, anchor, interpretation: str):
        cleaned = (interpretation or "").strip()
        if not cleaned:
            return anchor
        paragraphs = [item.strip() for item in re.split(r"\n\s*\n+", cleaned) if item.strip()]
        for item in paragraphs:
            factor_block = cls._parse_etdah_ad_factor_paragraph(item)
            if factor_block:
                heading, body = factor_block
                anchor = cls._insert_bpa_section_heading_after(anchor, heading)
                anchor = cls._insert_bpa_body_paragraph_after(anchor, body)
                continue
            if item.startswith("Interpretação e Observações Clínicas:"):
                anchor = cls._insert_interpretation_block_after(anchor, item)
                continue
            anchor = cls._insert_bpa_body_paragraph_after(anchor, item)
        return anchor

    @classmethod
    def _strip_legacy_srs2_table(cls, text: str | None) -> str:
        cleaned = (text or "").strip()
        if not cleaned:
            return ""

        cleaned = re.sub(
            r"={20,}\s*\n"
            r"Fator\s+Pts\s*Brts\s+T-Score\s+Percentil\s+Classifica(?:ç|c)ão\s*\n"
            r"={20,}\s*\n"
            r".*?"
            r"\n={20,}",
            "",
            cleaned,
            flags=re.IGNORECASE | re.DOTALL,
        ).strip()

        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        if cleaned in cls.EMPTY_INTERPRETATION_MESSAGES:
            return ""
        return cleaned

    @classmethod
    def _wasi_candidate_has_stale_qi(cls, text: str | None, test_payload: dict | None) -> bool:
        if not text or not test_payload:
            return False
        if test_payload.get("instrument_code") != "wasi":
            return False
        computed_composites = ((test_payload.get("computed_payload") or {}).get("composites") or {})
        if not computed_composites:
            return False

        text_str = str(text)
        qi_pattern = re.compile(
            r"\b(\d{2,3})\b",
            re.IGNORECASE,
        )
        stale_qi_count = 0
        valid_qi_count = 0

        composite_values = {
            (computed_composites.get(k) or {}).get("qi")
            for k in ["qi_verbal", "qi_execucao", "qit_4"]
            if (computed_composites.get(k) or {}).get("qi") is not None
        }

        for match in qi_pattern.finditer(text_str):
            qi_val = int(match.group(1))
            if 60 <= qi_val <= 160:
                if qi_val in composite_values:
                    valid_qi_count += 1
                else:
                    stale_qi_count += 1

        return stale_qi_count > 0 and stale_qi_count >= valid_qi_count

    @classmethod
    def _resolve_interpretation_text(
        cls,
        primary_section: str | None,
        fallback_section: str | None,
        test_payload: dict | None,
        context: dict | None = None,
    ) -> str:
        if test_payload:
            candidates = [
                (test_payload or {}).get("clinical_interpretation"),
                cls._fallback_test_interpretation(test_payload),
                (test_payload or {}).get("summary"),
                primary_section,
                fallback_section,
            ]
        else:
            candidates = [
                primary_section,
                fallback_section,
            ]

        for candidate in candidates:
            cleaned = cls._strip_legacy_srs2_table(candidate)
            if context is not None:
                cleaned = cls._sanitize_section_text_for_patient(cleaned, context)
            if not cleaned:
                continue
            if (
                test_payload
                and test_payload.get("instrument_code") == "wasi"
                and cls._wasi_candidate_has_stale_qi(cleaned, test_payload)
            ):
                continue
            return cleaned
        
        patient_name_for_interp = (((context or {}).get("patient") or {}).get("full_name") or "Paciente").split()[0] or "Paciente"

        if test_payload and test_payload.get("instrument_code") == "wais3":
            from apps.tests.wais3.interpreters import build_wais3_interpretation
            merged = {
                **(test_payload.get("computed_payload") or {}),
                **(test_payload.get("structured_results") or {}),
            }
            return build_wais3_interpretation(merged, patient_name_for_interp)

        if test_payload and test_payload.get("instrument_code") == "wasi":
            from apps.tests.wasi.interpreters import build_wasi_interpretation
            merged = {
                **(test_payload.get("computed_payload") or {}),
                **(test_payload.get("structured_results") or {}),
            }
            return build_wasi_interpretation(merged, patient_name=patient_name_for_interp)

        return ""

    @classmethod
    def _merge_text_blocks(cls, *blocks: str | None) -> str:
        cleaned_blocks = []
        seen = set()
        for block in blocks:
            cleaned = PtBrTextService.normalize(block)
            if not cleaned:
                continue
            key = cleaned.casefold()
            if key in seen:
                continue
            seen.add(key)
            cleaned_blocks.append(cleaned)
        return "\n\n".join(cleaned_blocks)

    @classmethod
    def _fallback_test_interpretation(cls, test_payload: dict | None) -> str:
        if not test_payload:
            return ""

        instrument_code = test_payload.get("instrument_code")

        if instrument_code == "srs2":
            merged_data = {
                **(test_payload.get("computed_payload") or {}),
                **(
                    test_payload.get("classified_payload")
                    or test_payload.get("structured_results")
                    or {}
                ),
            }
            return (interpret_srs2_results(merged_data) or "").strip()

        if instrument_code == "wasi":
            merged_data = {
                **(test_payload.get("computed_payload") or {}),
                **(test_payload.get("classified_payload") or {}),
                **(test_payload.get("structured_results") or {}),
            }
            patient_name = (((test_payload or {}).get("patient_context") or {}).get("full_name") or "Paciente").split()[0] or "Paciente"
            from apps.tests.wasi.interpreters import build_wasi_interpretation
            return (build_wasi_interpretation(merged_data, patient_name=patient_name) or "").strip()

        if instrument_code == "bai":
            from apps.tests.bai.interpreters import get_report_interpretation

            merged_data = {
                **(test_payload.get("computed_payload") or {}),
                **(test_payload.get("classified_payload") or {}),
            }
            return (get_report_interpretation(merged_data) or "").strip()

        if instrument_code == "ebadep_a":
            from apps.tests.ebadep_a.interpreters import get_report_interpretation

            classified_payload = (test_payload.get("classified_payload") or {}).get("result") or (
                test_payload.get("classified_payload") or {}
            )
            classificacao = (
                classified_payload.get("classificacao")
                or (test_payload.get("classified_payload") or {}).get("classificacao")
            )
            return (get_report_interpretation(classificacao or "", "Paciente") or "").strip()

        return ""

    @classmethod
    def _ravlt_conceptual_text(cls) -> str:
        return (
            "O Rey Auditory Verbal Learning Test (RAVLT) é um teste neuropsicológico amplamente utilizado para avaliar a memória verbal, a capacidade de aprendizado auditivo e a retenção de informações ao longo do tempo. Desenvolvido por Rey (1958), o RAVLT permite analisar diferentes aspectos da memória, como a curva de aprendizado, a interferência, o esquecimento e o reconhecimento verbal (Lezak et al., 2004). Ele é frequentemente utilizado na investigação de déficits cognitivos associados a doenças neurodegenerativas, lesões cerebrais traumáticas e transtornos psiquiátricos (Strauss, Sherman & Spreen, 2006). Os resultados do teste auxiliam no diagnóstico diferencial de condições como doença de Alzheimer e TDAH, além de fornecerem subsídios para o planejamento de intervenções cognitivas (Salthouse, 2010). Assim, o RAVLT é uma ferramenta essencial para a avaliação da memória e da aprendizagem verbal."
        )

    @classmethod
    def _fdt_description_text(cls) -> str:
        return (
            "O Teste dos Cinco Dígitos (FDT) permite avaliar tantos processos automáticos, de baixa demanda executiva, quanto processos controlados, que exigem regulação intencional, inibição de respostas e flexibilidade cognitiva. Esses sistemas estão relacionados ao funcionamento coordenado entre redes temporais e pré-frontais, fundamentais para a autorregulação e o controle executivo (Norman & Shallice, 1986; Miyake et al., 2000; Diamond, 2013)."
        )

    @classmethod
    def _etdah_ad_description_text(cls) -> str:
        return (
            "A E-TDAH-AD é um instrumento de autorrelato destinado a adolescentes e adultos, voltado à investigação de indicadores de desatenção, impulsividade, hiperatividade, regulação emocional, motivação e autorregulação no cotidiano. Seus resultados devem ser integrados aos achados cognitivos e comportamentais da avaliação neuropsicológica."
        )

    @classmethod
    def _bai_description_text(cls) -> str:
        return (
            "O Inventário de Ansiedade de Beck (BAI) é um instrumento de autorrelato utilizado para mensurar a intensidade de sintomas ansiosos, com ênfase em manifestações fisiológicas, cognitivas e subjetivas de ansiedade. Seu resultado contribui para o rastreio clínico e deve ser interpretado em conjunto com a entrevista, a observação comportamental e os demais instrumentos do protocolo avaliativo."
        )

    @classmethod
    def _ebadep_description_text(cls) -> str:
        return (
            "A EBADEP-A avalia a presença e a intensidade de sintomas depressivos em adultos, contemplando dimensões cognitivas, afetivas, somáticas e motivacionais. Seu resultado contribui para o rastreio de indicadores clínicos, devendo ser interpretado em conjunto com a anamnese e a observação clínica."
        )

    @classmethod
    def _bfp_description_text(cls) -> str:
        return (
            "A Bateria Fatorial de Personalidade (BFP) foi utilizada para investigar traços de personalidade com base no modelo dos Cinco Grandes Fatores, permitindo compreender tendências emocionais, interpessoais, motivacionais e comportamentais. Seus resultados devem ser integrados à anamnese, observação clínica e demais instrumentos aplicados."
        )

    @classmethod
    def _bai_rows(cls, test: dict | None):
        payload = {
            **((test or {}).get("computed_payload") or {}),
            **((test or {}).get("classified_payload") or {}),
        }
        if not payload:
            return None

        summary = ((test or {}).get("computed_payload") or {}).get("tables", {}).get("summary_table", [])
        summary_row = summary[0] if summary else {}
        classificacao = payload.get("classificacao") or {}
        descricao = (
            summary_row.get("description")
            or classificacao.get("interpretation")
            or payload.get("interpretacao_faixa")
            or "-"
        )
        escala = summary_row.get("scale") or "Escore Total"
        bruto = summary_row.get("raw_score")
        norma = summary_row.get("norm_value")
        if bruto is None:
            bruto = payload.get("escore_total") or payload.get("total_raw_score")
        if norma is None:
            norma = payload.get("t_score")

        return [
            ["Escala", "Pontuação bruta", "Valor da norma"],
            [escala, cls._num(bruto), cls._num(norma)],
            [descricao, "", ""],
        ]

    @classmethod
    def _bai_chart(cls, test: dict | None):
        payload = {
            **((test or {}).get("computed_payload") or {}),
            **((test or {}).get("classified_payload") or {}),
        }
        if not payload:
            return None

        total_raw_score = payload.get("total_raw_score") or payload.get("escore_total")
        t_score = payload.get("t_score")
        if total_raw_score is None or t_score is None:
            return None

        min_score = 20
        max_score = 80
        safe_t = max(min_score, min(max_score, float(t_score)))
        ticks = list(range(20, 81, 5))
        special_labels = {20: "min", 40: "-s", 50: "m", 60: "+s", 80: "max"}

        fig, ax = plt.subplots(figsize=(11.5, 3.2))
        ax.set_xlim(16, 108)
        ax.set_ylim(0, 1)
        ax.axis("off")

        start_x = 20
        end_x = 80
        cell_height = 0.28
        cell_y = 0.34

        ax.text(16, 0.92, "Inventário de Ansiedade · Padrão", fontsize=10, fontweight="bold")
        ax.text(16, 0.82, "Amostra Geral · Escore T (50+10z)", fontsize=10, fontweight="bold")

        for x in (16, 22):
            ax.add_patch(Rectangle((x, cell_y), 6, cell_height, facecolor="#dff5f7", edgecolor="#ffffff"))
        ax.text(19, cell_y + cell_height / 2, f"{int(total_raw_score)}", ha="center", va="center", fontsize=10)
        ax.text(25, cell_y + cell_height / 2, f"{safe_t:.0f}", ha="center", va="center", fontsize=10)

        segment_width = (end_x - start_x) / 12
        for idx in range(12):
            x0 = 28 + idx * segment_width
            ax.add_patch(
                Rectangle(
                    (x0, cell_y),
                    segment_width,
                    cell_height,
                    facecolor="#b8b8b8" if idx % 2 == 0 else "#a7a7a7",
                    edgecolor="#ffffff",
                    linewidth=0.8,
                )
            )

        ax.add_patch(Rectangle((88, cell_y), 18, cell_height, facecolor="#dff5f7", edgecolor="#ffffff"))
        ax.text(89, cell_y + cell_height / 2, "Escore Total", va="center", fontsize=10, fontweight="bold")

        for value in ticks:
            x = 28 + ((value - start_x) / (end_x - start_x)) * 60
            ax.plot([x, x], [0.64, 0.69], color="#555555", linewidth=0.8)
            ax.text(x, 0.73, f"{value}", ha="center", va="bottom", fontsize=8)
            if value in special_labels:
                ax.text(x, 0.78, special_labels[value], ha="center", va="bottom", fontsize=10)

        point_x = 28 + ((safe_t - start_x) / (end_x - start_x)) * 60
        ax.scatter([point_x], [cell_y + cell_height / 2], s=48, color="#e11d48", zorder=3)
        ax.text(17.2, 0.59, "Dados brutos", rotation=90, fontsize=8, va="center")
        ax.text(23.2, 0.59, "Normas", rotation=90, fontsize=8, va="center")

        output = BytesIO()
        fig.tight_layout()
        fig.savefig(output, format="png", dpi=200, bbox_inches="tight")
        plt.close(fig)
        return output.getvalue()

    @classmethod
    def _bfp_rows(cls, test: dict | None):
        computed = (test or {}).get("computed_payload") or {}
        factors = computed.get("factors") or ((test or {}).get("structured_results") or {}).get("factors") or {}
        facets = computed.get("facets") or {}
        factor_order = computed.get("factor_order") or ["NN", "EE", "SS", "RR", "AA"]
        
        factor_facets_map = {
            "NN": ["N1", "N2", "N3", "N4"],
            "EE": ["E1", "E2", "E3", "E4"],
            "SS": ["S1", "S2", "S3"],
            "RR": ["R1", "R2", "R3"],
            "AA": ["A1", "A2", "A3"],
        }
        
        factor_names = {
            "NN": "NEUROTICISMO",
            "EE": "EXTROVERSÃO",
            "SS": "SOCIALIZAÇÃO",
            "RR": "REALIZAÇÃO",
            "AA": "ABERTURA",
        }
        
        facet_names = {
            "N1": "Vulnerabilidade", "N2": "Instabilidade Emocional", "N3": "Passividade", "N4": "Depressão",
            "E1": "Comunicação", "E2": "Altivez", "E3": "Dinamismo", "E4": "Interações Sociais",
            "S1": "Amabilidade", "S2": "Pró-sociabilidade", "S3": "Confiança nas Pessoas",
            "R1": "Competência", "R2": "Ponderação", "R3": "Empenho",
            "A1": "Abertura a Ideias", "A2": "Liberalismo", "A3": "Busca por Novidades",
        }
        
        tables = []
        for factor_code in factor_order:
            factor = factors.get(factor_code)
            if not factor:
                continue
            
            rows = [["Faceta/Dimensão", "Escore Bruto", "Percentil", "Classificação"]]
            
            for facet_code in factor_facets_map.get(factor_code, []):
                facet = facets.get(facet_code)
                if facet:
                    rows.append([
                        facet_names.get(facet_code, facet_code),
                        cls._num(facet.get("raw_score")),
                        cls._num(facet.get("percentile")),
                        facet.get("classification") or "-",
                    ])
            
            rows.append([
                factor_names.get(factor_code, factor_code),
                cls._num(factor.get("raw_score")),
                cls._num(factor.get("percentile")),
                factor.get("classification") or "-",
            ])
            
            if len(rows) > 1:
                tables.append(rows)
        
        return tables if tables else None

    @classmethod
    def _bfp_chart(cls, test: dict | None):
        computed = (test or {}).get("computed_payload") or {}
        facets = computed.get("facets") or {}
        
        BFP_FACETAS_RADAR = [
            "Vulnerabilidade",
            "Instabilidade\nEmocional",
            "Passividade",
            "Depressão",
            "Comunicação",
            "Altivez",
            "Dinamismo",
            "Interações\nSociais",
            "Amabilidade",
            "Pró-sociabilidade",
            "Confiança nas\nPessoas",
            "Competência",
            "Ponderação",
            "Empenho",
            "Abertura a\nIdeias",
            "Liberalismo",
            "Busca por\nNovidades",
        ]
        
        facet_codes = ["N1", "N2", "N3", "N4", "E1", "E2", "E3", "E4", "S1", "S2", "S3", "R1", "R2", "R3", "A1", "A2", "A3"]
        
        percentis = []
        for code in facet_codes:
            facet = facets.get(code)
            if facet is None:
                percentis.append(50.0)
                continue
            if isinstance(facet, int):
                pct = facet
            elif isinstance(facet, dict):
                pct = facet.get("percentile")
                if pct is None:
                    pct = 50
            else:
                pct = 50
            try:
                percentis.append(float(pct))
            except (TypeError, ValueError):
                percentis.append(50.0)
        
        if len(percentis) != 17:
            return None
        
        norma = [50] * 17
        n = len(BFP_FACETAS_RADAR)
        
        angles = np.linspace(0, 2 * np.pi, n, endpoint=False).tolist()
        angles_closed = angles + angles[:1]
        
        values = list(percentis) + [percentis[0]]
        norm_values = norma + [norma[0]]
        
        fig = plt.figure(figsize=(10, 10), dpi=180)
        ax = plt.subplot(111, polar=True)
        
        ax.set_theta_offset(np.pi / 2)
        ax.set_theta_direction(-1)
        
        ax.plot(angles_closed, values, linewidth=2.5, marker="o", color="#2F6DB3", label="Resultado do avaliado")
        ax.fill(angles_closed, values, alpha=0.28, color="#2F6DB3")
        
        ax.plot(angles_closed, norm_values, linewidth=2.3, linestyle="-", color="#5E8E3E", label="Amostra normativa")
        ax.fill(angles_closed, norm_values, alpha=0.03, color="#5E8E3E")
        
        ax.set_ylim(0, 100)
        ax.set_yticks([20, 40, 60, 80, 100])
        ax.set_yticklabels(["20", "40", "60", "80", "100"], fontsize=9)
        
        ax.set_xticks(angles)
        ax.set_xticklabels(BFP_FACETAS_RADAR, fontsize=9)
        
        ax.grid(True, linewidth=0.8, alpha=0.35)
        ax.spines["polar"].set_alpha(0.25)
        
        plt.title("RADAR DE AVALIAÇÃO DAS FACETAS - BFP", fontsize=17, fontweight="bold", pad=35)
        
        ax.legend(loc="upper center", bbox_to_anchor=(0.5, 1.15), ncol=2, frameon=False, fontsize=10)
        
        output = BytesIO()
        fig.savefig(output, format="png", bbox_inches="tight", dpi=300)
        plt.close(fig)
        return output.getvalue()

    @classmethod
    def _srs2_description_text(cls) -> str:
        return (
            "A SRS-2 é uma escala destinada à investigação de dificuldades em comunicação social, cognição social, motivação social, percepção social e padrões restritos e repetitivos, contribuindo para o rastreio de traços associados ao Transtorno do Espectro Autista. Seus resultados devem ser compreendidos em articulação com a observação clínica e a história do desenvolvimento."
        )

    @classmethod
    def _append_chart(
        cls,
        document,
        caption: str | None,
        image_bytes: bytes | None,
        width=None,
        height=None,
    ):
        if not image_bytes:
            return
        img = document.add_paragraph()
        img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img.paragraph_format.first_line_indent = Pt(0)
        img.paragraph_format.left_indent = Pt(0)
        img.paragraph_format.space_before = Pt(0)
        picture_kwargs = {}
        if width is not None:
            picture_kwargs["width"] = width
        if height is not None:
            picture_kwargs["height"] = height
        if not picture_kwargs:
            picture_kwargs["width"] = cls.IMAGE_WIDTH
        img.add_run().add_picture(BytesIO(image_bytes), **picture_kwargs)
        if caption:
            p = document.add_paragraph()
            r = p.add_run(caption)
            r.font.name = cls.FONT_NAME
            r.font.size = cls.BODY_SIZE
            cls._format_caption_paragraph(p)

    @classmethod
    def _format_chart_legend_paragraph(cls, paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = cls.BODY_SPACE_AFTER
        paragraph.paragraph_format.line_spacing = cls.BODY_LINE_SPACING
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        for run in paragraph.runs:
            run.font.name = cls.FONT_NAME
            run.font.size = cls.CAPTION_SIZE
            run.bold = False
            run.italic = True

    @classmethod
    def _wisc_chart_legend(cls, chart_index: int) -> str:
        return (
            f"Gráfico {chart_index} WISC-IV - INDICES DE QI :"
            "Índice de Compreensão Verbal (ICV): composto por provas que avaliam as habilidades verbais por meio do raciocínio, compreensão e conceituação. "
            "Índice de Organização Perceptual (IOP): constituído por atividades que examinam o grau e a qualidade do contato não verbal do indivíduo com o ambiente, assim como a capacidade de integrar estímulos perceptuais e respostas motoras pertinentes, o nível de rapidez com o qual executa uma atividade e o modo como avalia informações visuoespaciais. "
            "Índice de Memória Operacional (IMO): formado por provas que analisam atenção, concentração e memória de trabalho. "
            "Índice de Velocidade de Processamento (IVP): constitui-se de atividades que avaliam agilidade mental e processamento grafomotor. "
            "Coeficiente de Inteligência Total (QIT): avalia o nível geral do funcionamento intelectual."
        )

    @classmethod
    def _find_test(cls, context: dict, code: str):
        for item in context.get("validated_tests") or []:
            if item.get("instrument_code") == code:
                return item
        return None

    @classmethod
    def _find_tests(cls, context: dict, code: str):
        tests = [
            item
            for item in context.get("validated_tests") or []
            if item.get("instrument_code") == code
        ]
        form_order = {"child": 0, "parent": 1}
        return sorted(
            tests,
            key=lambda item: form_order.get(
                ((item.get("classified_payload") or {}).get("form_type") or (item.get("raw_payload") or {}).get("form") or "child"),
                99,
            ),
        )

    @staticmethod
    def _scared_form_label(test: dict | None) -> str:
        form_type = ((test or {}).get("classified_payload") or {}).get("form_type") or ((test or {}).get("raw_payload") or {}).get("form")
        return "Mãe" if form_type == "parent" else "Autorrelato"

    @classmethod
    def _scared_form_title(cls, test: dict | None) -> str:
        return f"SCARED - {cls._scared_form_label(test)}"

    @classmethod
    def _scared_table_key(cls, test: dict | None) -> str:
        payload = (test or {}).get("classified_payload") or {}
        return "scared_parent" if payload.get("form_type") == "parent" else "scared_self"

    @classmethod
    def _format_date_display(cls, value: str | None) -> str:
        if not value:
            return "Não informado"
        try:
            year, month, day = value.split("-")
            return f"{day}/{month}/{year}"
        except ValueError:
            return value

    @staticmethod
    def _format_sex_display(value: str | None) -> str:
        return {"M": "Masculino", "F": "Feminino", "O": "Outro"}.get(
            value or "", value or "Não informado"
        )

    @classmethod
    def _age_text(cls, context: dict) -> str:
        birth_date = (context.get("patient") or {}).get("birth_date")
        if not birth_date:
            return "Não informada"
        try:
            from datetime import date

            year, month, day = map(int, birth_date.split("-"))
            born = date(year, month, day)
            today = date.today()
            years = (
                today.year
                - born.year
                - ((today.month, today.day) < (born.month, born.day))
            )
            months = (today.month - born.month) % 12
            return f"{years} anos e {months} meses"
        except ValueError:
            return "Não informada"

    @staticmethod
    def _filiation_text(patient: dict) -> str:
        names = [patient.get("father_name"), patient.get("mother_name")]
        names = [name for name in names if name]
        return (
            " e ".join(names)
            if names
            else patient.get("responsible_name") or "Não informada"
        )

    @staticmethod
    def _schooling_text(patient: dict) -> str:
        return patient.get("grade_year") or patient.get("schooling") or "Não informada"

    @classmethod
    def _adolescent_instruments(cls, context: dict):
        return cls._procedure_items(context)

    @classmethod
    def _procedure_items(cls, context: dict):
        validated_codes = []
        seen_codes = {"anamnese"}
        for test in context.get("validated_tests") or []:
            code = test.get("instrument_code")
            if not code or code in seen_codes:
                continue
            seen_codes.add(code)
            validated_codes.append(code)

        ordered_codes = [
            code
            for code in cls.PROCEDURES_ORDER
            if code == "anamnese" or code in validated_codes
        ]

        items = []
        for code in ordered_codes:
            catalog_entry = cls.INSTRUMENT_CATALOG.get(code)
            if not catalog_entry:
                continue
            items.append(
                {
                    "name": catalog_entry["nome"],
                    "description": catalog_entry["descricao"],
                }
            )

        remaining_codes = [
            code for code in validated_codes if code not in cls.PROCEDURES_ORDER
        ]
        for code in remaining_codes:
            catalog_entry = cls.INSTRUMENT_CATALOG.get(code)
            if not catalog_entry:
                continue
            items.append(
                {
                    "name": catalog_entry["nome"],
                    "description": catalog_entry["descricao"],
                }
            )

        return items

    @staticmethod
    def _patient_reference_name(context: dict) -> str:
        patient_name = ((context or {}).get("patient") or {}).get("full_name") or "Paciente"
        return patient_name.strip().split()[0] if patient_name.strip() else "Paciente"

    @staticmethod
    def _split_bullets(text: str):
        if not text:
            return []
        parts = [part.strip(" -") for part in text.split("\n") if part.strip()]
        return parts or [text]

    @classmethod
    def _team_paragraphs(cls, report):
        email = getattr(getattr(report, "author", None), "email", "") or ""
        return [
            "A Avaliação Neuropsicológica, quando bem fundamentada, é essencial para direcionar a reabilitação cognitiva e fornecer subsídios para outros profissionais em suas respectivas áreas de atuação.",
            "É importante que o documento seja lido na íntegra, e não apenas a conclusão, para que se compreenda plenamente o raciocínio clínico utilizado ao longo de todo o processo avaliativo.",
            f"Com a devida autorização por escrito dos responsáveis, coloco-me à disposição para esclarecimentos e discussões sobre o exame{f', podendo ser contatada pelo e-mail {email}' if email else ''}.",
        ]

    @classmethod
    def _important_document_items(cls):
        return [
            "1. Não deve ser utilizado para fins diferentes daqueles especificados no item de identificação do documento.",
            "2. Possui caráter sigiloso e extrajudicial, não cabendo à psicóloga a responsabilidade pelo seu uso indevido ou pela entrega do laudo sem autorização adequada.",
            "3. A análise isolada deste laudo não possui valor diagnóstico se não for considerada em conjunto com dados clínicos, epistemológicos, exames de neuroimagem e laboratoriais adicionais referentes ao paciente.",
            "4. Esta avaliação está em conformidade com as Resoluções CRP 09/2018 e 06/2019. Em conformidade com o Código de Ética Profissional, este exame deve ser tratado como confidencial.",
        ]

    @classmethod
    def _signature_lines(cls, report):
        author_name = getattr(
            getattr(report, "author", None), "display_name", "Profissional responsável"
        )
        created = getattr(report, "created_at", None)
        city = "Goiânia"
        date_text = created.strftime("%d/%m/%Y") if created else ""
        return [f"{city}, {date_text}", "", author_name]

    @staticmethod
    def _references_list(context: dict):
        return build_references(context.get("validated_tests") or [])

    @classmethod
    def _wisc_memory_rows(cls, test: dict | None, context: dict):
        payload = (test or {}).get("classified_payload") or {}
        ncp_table = cls._wisc_ncp_table(context)
        rows = [[
            "Testes Utilizados",
            "Escore Máximo",
            "Escore Médio",
            "Escore Mínimo",
            "Escore Bruto",
            "Classificação",
        ]]
        labels = {
            "SNL": "Seq. Núm. e Letras",
            "DG": "Dígitos",
        }
        for code in ("SNL", "DG"):
            item = next(
                (entry for entry in payload.get("subtestes") or [] if entry.get("codigo") == code),
                None,
            )
            if not item:
                continue
            score_max, score_mid, score_min = cls._wisc_reference_scores(
                ncp_table, code
            )
            rows.append(
                [
                    labels.get(code, item.get("subteste") or "-"),
                    score_max,
                    score_mid,
                    score_min,
                    cls._num(item.get("escore_bruto")),
                    item.get("classificacao") or "-",
                ]
            )
        return rows if len(rows) > 1 else None

    @classmethod
    def _ravlt_norm_band(cls, test: dict | None, context: dict | None = None) -> str:
        context = context or {}
        patient = context.get("patient") or {}
        birth_date = patient.get("birth_date")
        applied_on = (test or {}).get("applied_on")
        if birth_date and applied_on:
            try:
                age = _calcular_idade(str(birth_date), str(applied_on)[:10])
                if isinstance(age, tuple):
                    age = age[0]
                return get_ravlt_age_band(int(age))
            except Exception:
                pass

        payload = (test or {}).get("classified_payload") or {}
        band = payload.get("faixa_etaria")
        if band in RAVLT_NORMS:
            return band
        return "21-30"

    @classmethod
    def _ravlt_obtained_map(cls, test: dict | None) -> dict:
        classified = (test or {}).get("classified_payload") or {}
        computed = (test or {}).get("computed_payload") or {}
        # Build a normalized map from classified resultados for flexible key matching
        def _normalize_key(s: str | None) -> str | None:
            if not s:
                return None
            # lower-case and keep only ascii alphanumerics
            return re.sub(r"[^a-z0-9]", "", str(s).casefold())

        results_map = {}
        for item in classified.get("resultados") or []:
            if not isinstance(item, dict):
                continue
            var = item.get("variavel") or item.get("variavel_nome") or item.get("nome")
            key = _normalize_key(var) or None
            if key:
                results_map[key] = item

        def _get_bruto_from_results(names: list[str]) -> object:
            for name in names:
                key = _normalize_key(name)
                if key and key in results_map:
                    return results_map[key].get("bruto")
            return None

        # Try to read values from the classified resultados if available
        if results_map:
            return {
                "A1": _get_bruto_from_results(["A1", "a1"]),
                "A2": _get_bruto_from_results(["A2", "a2"]),
                "A3": _get_bruto_from_results(["A3", "a3"]),
                "A4": _get_bruto_from_results(["A4", "a4"]),
                "A5": _get_bruto_from_results(["A5", "a5"]),
                "B1": _get_bruto_from_results(["B1", "b1", "b"]),
                "A6": _get_bruto_from_results(["A6", "a6"]),
                "A7": _get_bruto_from_results(["A7", "a7"]),
                "R": _get_bruto_from_results(["R", "r", "reconhecimentolistaa", "reconhecimento lista a"]),
                "ALT": _get_bruto_from_results(["ALT", "alt", "aprendlongodastentativas", "aprend longo das tentativas"]),
                "RET": _get_bruto_from_results(["RET", "ret", "velocidadedeesquecimento", "velocidade de esquecimento"]),
                "I.P.": _get_bruto_from_results(["IP", "ip", "interferenciaproativa", "interferência proativa"]),
                "I.R.": _get_bruto_from_results(["IR", "ir", "interferenciaretroativa", "interferência retroativa"]),
            }

        # Fallback: try to read from computed_payload or classified top-level short keys
        return {
            "A1": computed.get("a1", {}).get("escore") or classified.get("a1"),
            "A2": computed.get("a2", {}).get("escore") or classified.get("a2"),
            "A3": computed.get("a3", {}).get("escore") or classified.get("a3"),
            "A4": computed.get("a4", {}).get("escore") or classified.get("a4"),
            "A5": computed.get("a5", {}).get("escore") or classified.get("a5"),
            "B1": computed.get("b", {}).get("escore") or classified.get("b1") or classified.get("b"),
            "A6": computed.get("a6", {}).get("escore") or classified.get("a6"),
            "A7": computed.get("a7", {}).get("escore") or classified.get("a7"),
            "R": computed.get("reconhecimento", {}).get("escore") or classified.get("r"),
            "ALT": computed.get("aprend_longo", {}).get("escore") or classified.get("alt"),
            "RET": computed.get("velocidade_esquecimento", {}).get("escore") or classified.get("ret"),
            "I.P.": computed.get("interferencia_proativa", {}).get("escore") or classified.get("ip"),
            "I.R.": computed.get("interferencia_retroativa", {}).get("escore") or classified.get("ir"),
        }

    @classmethod
    def _ravlt_rows(cls, test: dict | None, context: dict | None = None):
        # Columns in the same order as the reference model.
        columns = [
            "A1",
            "A2",
            "A3",
            "A4",
            "A5",
            "B1",
            "A6",
            "A7",
            "R",
            "ALT",
            "RET",
            "I.P.",
            "I.R.",
        ]
        norm_key_map = {
            "A1": "A1",
            "A2": "A2",
            "A3": "A3",
            "A4": "A4",
            "A5": "A5",
            "A6": "A6",
            "A7": "A7",
            "B1": "B1",
            "R": "Reconhecimento Lista A",
            "ALT": "Aprend. longo das Tentativas",
            "RET": "Velocidade de Esquecimento",
            "I.P.": "Interferência Proativa",
            "I.R.": "Interferência Retroativa",
        }
        band = cls._ravlt_norm_band(test, context)
        norms = RAVLT_NORMS.get(band, RAVLT_NORMS["21-30"])

        # Prefer explicit payload values if present: ravlt_esperado / ravlt_minimo / ravlt_obtido
        payload_source = (test or {})
        classified = (test or {}).get("classified_payload") or {}
        computed = (test or {}).get("computed_payload") or {}

        # Merge likely locations for a potential ravlt payload
        candidate_payloads = [payload_source, classified, computed]

        def _find_ravlt_map(key_name: str) -> dict | None:
            for src in candidate_payloads:
                val = src.get(key_name)
                if isinstance(val, dict) and val:
                    return val
            return None

        ravlt_expected_raw = _find_ravlt_map("ravlt_esperado")
        ravlt_minimum_raw = _find_ravlt_map("ravlt_minimo")
        ravlt_obtained_raw = _find_ravlt_map("ravlt_obtido")

        def _normalize_payload_map(m: dict | None) -> dict:
            if not m:
                return {}
            mapping = {}
            short_map = {
                "a1": "A1",
                "a2": "A2",
                "a3": "A3",
                "a4": "A4",
                "a5": "A5",
                "b1": "B1",
                "b": "B1",
                "a6": "A6",
                "a7": "A7",
                "r": "R",
                "alt": "ALT",
                "ret": "RET",
                "ip": "I.P.",
                "ir": "I.R.",
            }
            for k, v in (m or {}).items():
                if not k:
                    continue
                key_norm = str(k).casefold()
                if key_norm in short_map:
                    mapping[short_map[key_norm]] = v
                else:
                    # also accept keys already matching final labels
                    mapping_key = k if k in columns else None
                    if mapping_key:
                        mapping[mapping_key] = v
            return mapping

        expected_map = _normalize_payload_map(ravlt_expected_raw)
        minimum_map = _normalize_payload_map(ravlt_minimum_raw)
        obtained_map = _normalize_payload_map(ravlt_obtained_raw)

        # If obtained_map is empty, fallback to computed/classified extraction
        if not any(v is not None for v in obtained_map.values()):
            obtained_map = cls._ravlt_obtained_map(test)

        if not any(value is not None for value in obtained_map.values()):
            return None

        # Header: first cell is the label 'Desempenho', then the numeric column labels
        rows = [["Desempenho", *columns], ["Esperado"], ["Mínimo"], ["Obtido"]]
        for column in columns:
            # Prefer explicit payload expected/minimum, otherwise use norms
            if expected_map.get(column) is not None:
                rows[1].append(cls._num(expected_map.get(column)))
            else:
                norm = norms.get(norm_key_map[column])
                rows[1].append(cls._num(norm.p50 if norm else None))

            if minimum_map.get(column) is not None:
                rows[2].append(cls._num(minimum_map.get(column)))
            else:
                norm = norms.get(norm_key_map[column])
                rows[2].append(cls._num(norm.p25 if norm else None))

            rows[3].append(cls._num(obtained_map.get(column)))
        return rows

    @classmethod
    def _format_table(cls, table, table_key: str):
        table_spec = cls._table_layout_spec(table_key)
        table.style = table_spec.get("style")
        table.alignment = (
            WD_TABLE_ALIGNMENT.LEFT
            if table_spec.get("alignment") == "left"
            else WD_TABLE_ALIGNMENT.CENTER
        )
        table.autofit = False
        cls._set_table_layout_fixed(table)
        cls._set_table_cell_margins(
            table,
            top=cls.TABLE_CELL_MARGIN_TOP,
            start=cls.TABLE_CELL_MARGIN_START,
            bottom=cls.TABLE_CELL_MARGIN_BOTTOM,
            end=cls.TABLE_CELL_MARGIN_END,
        )
        cls._apply_table_borders_profile(table_spec, table)
        cls._apply_table_widths(table, table_key)
        title_text = cls._table_title_text(table_key)
        for row_index, row in enumerate(table.rows):
            if row_index == 0:
                cls._set_repeat_table_header(row)
            for cell_index, cell in enumerate(row.cells):
                cls._apply_table_cell_shading(table_spec, title_text, row, cell, row_index, cell_index)
                cls._apply_table_cell_no_wrap(table_spec, title_text, row, cell, row_index, cell_index)
                cls._apply_table_cell_borders(table_spec, cell)
                for paragraph in cell.paragraphs:
                    cls._apply_table_paragraph_style(
                        table_spec,
                        title_text,
                        row,
                        paragraph,
                        row_index,
                        cell_index,
                    )
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = cls.FONT_NAME
                        run.font.size = cls._table_font_size(table_key, row_index, cell_index, bool(title_text))
                        run.bold = row_index in cls._table_header_row_indices(bool(title_text))
                        if row_index == 0 and title_text:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        if row_index == 0 and title_text:
                            run.bold = False
                        if table_key in {"wisc", "bpa"} and row_index > 0 and cell == row.cells[0]:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        if table_key == "fdt":
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        if cls._is_etdah_table_key(table_key) or table_key == "srs2" or cls._is_scared_table_key(table_key):
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        if table_key == "srs2" and row_index == len(table.rows) - 1:
                            run.bold = True
                        if table_key == "srs2" and row_index == len(table.rows) - 1:
                            run.bold = True
                        # BFP: linha do fator (última linha) em negrito
                        if table_key == "bfp" and row_index == len(table.rows) - 1:
                            run.bold = True
                        if table_key == "bai_scores" and row_index == 2:
                            run.bold = False
                cell.vertical_alignment = (
                    WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    if table_spec.get("vertical_align_center")
                    else None
                )

            if table_key == "wisc" and row_index > 0 and row.cells[0].text == "Fala Espontânea":
                merged_text = row.cells[1].text.strip()
                merged = row.cells[1]
                for idx in range(2, len(row.cells)):
                    merged = merged.merge(row.cells[idx])
                merged.text = merged_text
                row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                for paragraph in row.cells[1].paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.5
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = cls.FONT_NAME
                        run.font.size = cls.TABLE_SIZE
                row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            if title_text and row_index == 0:
                cls._merge_title_row(row)
                for paragraph in row.cells[0].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.5
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = cls.FONT_NAME
                        run.font.size = cls.TABLE_HEADER_SIZE
                        run.bold = False
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            if table_key == "bai_scores" and row_index == 2:
                merged_text = row.cells[0].text.strip()
                merged = row.cells[0]
                for idx in range(1, len(row.cells)):
                    merged = merged.merge(row.cells[idx])
                merged.text = merged_text
                paragraph = row.cells[0].paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.5
                paragraph.paragraph_format.first_line_indent = Pt(0)
                for run in paragraph.runs:
                    run.font.name = cls.FONT_NAME
                    run.font.size = cls.TABLE_SIZE
                    run.bold = False
                row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    @classmethod
    def _format_ravlt_table(cls, table):
        table.style = None
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cls._set_table_layout_fixed(table)
        cls._clear_table_width(table)
        cls._set_table_cell_margins(
            table,
            top=cls.TABLE_CELL_MARGIN_TOP,
            start=cls.TABLE_CELL_MARGIN_START,
            bottom=cls.TABLE_CELL_MARGIN_BOTTOM,
            end=cls.TABLE_CELL_MARGIN_END,
        )
        cls._set_table_borders(table, color="D9D9D9", size=4)

        widths = cls._table_widths("ravlt") or []
        title_text = cls._table_title_text("ravlt")
        if widths:
            cls._set_table_grid_widths(table, widths)

        for row_index, row in enumerate(table.rows):
            if row_index == 0:
                cls._set_repeat_table_header(row)
            for cell_index, cell in enumerate(row.cells):
                if row_index == 0 and title_text:
                    cls._set_cell_shading(cell, cls.TABLE_TITLE_FILL)
                elif row_index in {0, 1} and title_text:
                    cls._set_cell_shading(cell, cls.RAVLT_HEADER_FILL)
                elif row_index == 0:
                    cls._set_cell_shading(cell, cls.RAVLT_HEADER_FILL)
                if widths and cell_index < len(widths):
                    cell.width = widths[cell_index]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cls._set_cell_no_wrap(cell, True)

                for paragraph in cell.paragraphs:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                        if row_index in cls._table_header_row_indices(bool(title_text))
                        else WD_ALIGN_PARAGRAPH.LEFT if cell_index == 0
                        else WD_ALIGN_PARAGRAPH.CENTER
                    )

                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.5
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.left_indent = Pt(0)
                    paragraph.paragraph_format.right_indent = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = cls.FONT_NAME
                        run.font.size = cls._table_font_size("ravlt", row_index, cell_index, bool(title_text))
                        run.bold = row_index in cls._table_header_row_indices(bool(title_text))
                        if row_index == 0 and title_text:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            if title_text and row_index == 0:
                cls._merge_title_row(row)
                for paragraph in row.cells[0].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.5
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = cls.FONT_NAME
                        run.font.size = cls.TABLE_HEADER_SIZE
                        run.bold = False
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    @classmethod
    def _apply_table_widths(cls, table, table_key: str):
        widths = cls._table_layout_spec(table_key).get("widths")
        if not widths:
            return
        cls._set_table_grid_widths(table, widths)
        for row in table.rows:
            for idx, cell in enumerate(row.cells[: len(widths)]):
                cell.width = widths[idx]

    @classmethod
    def _table_layout_spec(cls, table_key: str) -> dict:
        if cls._is_etdah_table_key(table_key):
            table_key = "etdah"
        return TABLE_LAYOUT_SPECS.get(table_key) or TABLE_LAYOUT_SPECS["default"]

    @classmethod
    def _apply_table_cell_shading(cls, table_spec: dict, title_text: str, row, cell, row_index: int, cell_index: int):
        profile = table_spec.get("shading_profile")
        if profile == "wisc":
            if row_index == 0:
                cls._set_cell_shading(cell, cls.WISC_HEADER_FILL)
            elif cell == row.cells[0]:
                cls._set_cell_shading(cell, cls.WISC_NAME_FILL)
            else:
                cls._set_cell_shading(cell, cls.WISC_VALUE_FILL)
            return
        if profile == "bpa":
            if row_index == 0 and title_text:
                cls._set_cell_shading(cell, cls.TABLE_TITLE_FILL)
            elif row_index in {0, 1} and title_text:
                cls._set_cell_shading(cell, cls.BPA_HEADER_FILL)
            elif row_index == 0:
                cls._set_cell_shading(cell, cls.BPA_HEADER_FILL)
            else:
                cls._set_cell_shading(cell, cls.BPA_BODY_FILL)
            return
        if profile in {"fdt", "banded_fdt"}:
            if row_index == 0 and title_text:
                cls._set_cell_shading(cell, cls.TABLE_TITLE_FILL)
            elif row_index in {0, 1} and title_text:
                cls._set_cell_shading(cell, cls.FDT_HEADER_FILL)
            elif row_index == 0:
                cls._set_cell_shading(cell, cls.FDT_HEADER_FILL)
            elif cell == row.cells[0]:
                cls._set_cell_shading(cell, cls.FDT_HEADER_FILL)
            else:
                cls._set_cell_shading(cell, cls.FDT_BODY_FILL)
            return
        if profile == "bai_scores":
            if row_index == 0:
                cls._set_cell_shading(cell, cls.HEADER_FILL)
            elif row_index == 1 and cell_index in {1, 2}:
                cls._set_cell_shading(cell, "E9FBFD")
            elif row_index == 2:
                cls._set_cell_shading(cell, "DFF5F7")
            return
        if profile == "epq":
            cls._set_cell_shading(cell, cls.HEADER_FILL)
            return
        if profile == "default" and row_index == 0:
            cls._set_cell_shading(cell, cls.HEADER_FILL)

    @classmethod
    def _apply_table_cell_no_wrap(cls, table_spec: dict, title_text: str, row, cell, row_index: int, cell_index: int):
        header_rows = set(table_spec.get("header_no_wrap_rows_with_title") if title_text else table_spec.get("header_no_wrap_rows", ()))
        if row_index in header_rows:
            cls._set_cell_no_wrap(cell, True)
        if table_spec.get("first_col_no_wrap") and cell_index == 0:
            cls._set_cell_no_wrap(cell, True)
        if table_spec.get("last_col_no_wrap") and cell_index == len(row.cells) - 1:
            cls._set_cell_no_wrap(cell, True)

    @classmethod
    def _apply_table_cell_borders(cls, table_spec: dict, cell):
        if table_spec.get("border_profile") != "bottom_only":
            return
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        for edge in ["top", "left", "right", "bottom"]:
            edge_elem = OxmlElement(f"w:{edge}")
            edge_elem.set(qn("w:val"), "nil")
            tc_borders.append(edge_elem)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:color"), "6F6F6F")
        tc_borders.append(bottom)
        tc_pr.append(tc_borders)

    @classmethod
    def _apply_table_borders_profile(cls, table_spec: dict, table):
        profile = table_spec.get("border_profile")
        if profile == "subtle_grid":
            cls._set_table_borders(table, color="D9D9D9", size=4)
        elif profile == "bottom_only":
            cls._set_table_borders(table, color="FFFFFF", size=0)

    @classmethod
    def _apply_table_paragraph_style(
        cls,
        table_spec: dict,
        title_text: str,
        row,
        paragraph,
        row_index: int,
        cell_index: int,
    ):
        profile = table_spec.get("paragraph_profile")
        if profile == "wisc":
            if row_index == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif cell_index in {0, len(row.cells) - 1}:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif profile == "banded_first_left":
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 and title_text
                else WD_ALIGN_PARAGRAPH.LEFT if cell_index == 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
        elif profile == "center_all":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif profile == "bai_scores":
            if row_index in {0, 1}:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if cell_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)

    @classmethod
    def _set_table_grid_widths(cls, table, widths):
        tbl_grid = table._tbl.tblGrid
        if tbl_grid is None:
            tbl_grid = OxmlElement("w:tblGrid")
            table._tbl.insert(1, tbl_grid)
        for child in list(tbl_grid):
            tbl_grid.remove(child)
        for width in widths:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(int(width.twips)))
            tbl_grid.append(grid_col)

    @classmethod
    def _table_widths(cls, table_key: str):
        return cls._table_layout_spec(table_key).get("widths")

    @classmethod
    def _set_table_layout_autofit(cls, table):
        tbl_pr = table._tbl.tblPr
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "autofit")

    @classmethod
    def _set_table_layout_fixed(cls, table):
        tbl_pr = table._tbl.tblPr
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")

    @classmethod
    def _set_table_width_pct(cls, table, width_pct: int):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "pct")
        tbl_w.set(qn("w:w"), str(width_pct))

    @classmethod
    def _clear_table_width(cls, table):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "auto")
        tbl_w.set(qn("w:w"), "0")

    @classmethod
    def _set_table_cell_margins(
        cls, table, top: int = 0, start: int = 0, bottom: int = 0, end: int = 0
    ):
        tbl_pr = table._tbl.tblPr
        tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
        if tbl_cell_mar is None:
            tbl_cell_mar = OxmlElement("w:tblCellMar")
            tbl_pr.append(tbl_cell_mar)
        for tag, value in {
            "w:top": top,
            "w:start": start,
            "w:bottom": bottom,
            "w:end": end,
        }.items():
            element = tbl_cell_mar.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tbl_cell_mar.append(element)
            element.set(qn("w:w"), str(value))
            element.set(qn("w:type"), "dxa")

    @classmethod
    def _set_table_borders(cls, table, color: str = "000000", size: int = 4):
        tbl_pr = table._tbl.tblPr
        tbl_borders = tbl_pr.find(qn("w:tblBorders"))
        if tbl_borders is None:
            tbl_borders = OxmlElement("w:tblBorders")
            tbl_pr.append(tbl_borders)
        for tag in ("w:top", "w:left", "w:bottom", "w:right", "w:insideH", "w:insideV"):
            element = tbl_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tbl_borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), str(size))
            element.set(qn("w:color"), color)
            element.set(qn("w:space"), "0")

    @classmethod
    def _set_cell_width_auto(cls, cell):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:type"), "auto")
        tc_w.set(qn("w:w"), "0")

    @classmethod
    def _set_cell_no_wrap(cls, cell, enabled: bool):
        tc_pr = cell._tc.get_or_add_tcPr()
        no_wrap = tc_pr.find(qn("w:noWrap"))
        if enabled and no_wrap is None:
            tc_pr.append(OxmlElement("w:noWrap"))
        if not enabled and no_wrap is not None:
            tc_pr.remove(no_wrap)

    @classmethod
    def _table_cell_alignment(
        cls, table_key: str, row_index: int, cell_index: int, is_wisc_observation_row: bool = False
    ):
        if row_index == 0:
            return WD_ALIGN_PARAGRAPH.CENTER
        if is_wisc_observation_row:
            return WD_ALIGN_PARAGRAPH.LEFT
        return WD_ALIGN_PARAGRAPH.LEFT if cell_index == 0 else WD_ALIGN_PARAGRAPH.CENTER

    @classmethod
    def _merge_wisc_observation_row(cls, row):
        merged_text = row.cells[1].text.strip()
        merged = row.cells[1]
        for idx in range(2, len(row.cells)):
            merged = merged.merge(row.cells[idx])
        merged.text = merged_text

    @classmethod
    def _table_header_row_indices(cls, has_title_row: bool) -> set[int]:
        return {1} if has_title_row else {0}

    @classmethod
    def _table_font_size(
        cls,
        table_key: str,
        row_index: int,
        cell_index: int,
        has_title_row: bool = False,
    ):
        return cls.TABLE_HEADER_SIZE if row_index in cls._table_header_row_indices(has_title_row) else cls.TABLE_SIZE

    @classmethod
    def _set_repeat_table_header(cls, row):
        tr_pr = row._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)

    @classmethod
    def _set_cell_shading(cls, cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    @staticmethod
    def _num(value):
        if value is None:
            return "-"
        if isinstance(value, float):
            return f"{value:.2f}".rstrip("0").rstrip(".").replace(".", ",")
        return str(value).replace(".", ",")

    @staticmethod
    def _to_float(value, default: float = 0.0) -> float:
        if value in (None, "", "-"):
            return default
        if isinstance(value, (int, float)):
            return float(value)
        text = str(value).strip()
        if not text:
            return default
        text = text.replace("%", "").replace(" ", "")
        if "," in text and "." in text:
            text = text.replace(".", "").replace(",", ".")
        else:
            text = text.replace(",", ".")
        try:
            return float(text)
        except ValueError:
            return default

    @staticmethod
    def _parse_iso_date(value):
        if not value:
            return None
        from datetime import date

        if isinstance(value, date):
            return value
        try:
            return date.fromisoformat(str(value)[:10])
        except ValueError:
            return None

    @classmethod
    def _wisc_ncp_table(cls, context: dict):
        birth_date = cls._parse_iso_date((context.get("patient") or {}).get("birth_date"))
        reference_date = cls._parse_iso_date((context.get("evaluation") or {}).get("start_date"))
        for test in context.get("validated_tests") or []:
            if test.get("instrument_code") == "wisc4" and test.get("applied_on"):
                reference_date = cls._parse_iso_date(test.get("applied_on"))
                break
        if not birth_date or not reference_date:
            return None
        try:
            years, months = _calcular_idade(birth_date, reference_date)
            return _carregar_tabela_ncp(years, months)
        except Exception:
            return None

    @staticmethod
    def _normalize_wisc_cell(cell: str) -> tuple[int | None, int | None]:
        raw = (cell or "").strip().replace(":", "-")
        if not raw or raw == "-":
            return (None, None)
        if "-" in raw:
            start, end = raw.split("-", 1)
            try:
                return (int(start), int(end))
            except ValueError:
                return (None, None)
        if raw.isdigit():
            if len(raw) == 2 and int(raw[1]) == int(raw[0]) + 1:
                return (int(raw[0]), int(raw[1]))
            if len(raw) == 4:
                left, right = raw[:2], raw[2:]
                if left.isdigit() and right.isdigit():
                    return (int(left), int(right))
            number = int(raw)
            return (number, number)
        return (None, None)

    @classmethod
    def _wisc_reference_scores(cls, table: list[dict] | None, code: str) -> tuple[str, str, str]:
        if not table:
            return ("-", "-", "-")

        rows_by_pp = {}
        for row in table:
            try:
                rows_by_pp[int(row.get("PP") or 0)] = row
            except ValueError:
                continue

        max_start, max_end = cls._normalize_wisc_cell((rows_by_pp.get(19) or {}).get(code, ""))
        if max_start is None:
            max_text = "-"
        else:
            max_text = str(max_start) if max_start == max_end else f"{max_start}-{max_end}"

        mid_start, _ = cls._normalize_wisc_cell((rows_by_pp.get(8) or {}).get(code, ""))
        _, mid_end = cls._normalize_wisc_cell((rows_by_pp.get(12) or {}).get(code, ""))
        if mid_start is None or mid_end is None:
            mid_text = "-"
        else:
            mid_text = str(mid_start) if mid_start == mid_end else f"{mid_start}-{mid_end}"

        min_start, _ = cls._normalize_wisc_cell((rows_by_pp.get(5) or {}).get(code, ""))
        min_text = str(min_start) if min_start is not None else "-"
        return (max_text, mid_text, min_text)

    @staticmethod
    def _extract_percentile_value(value):
        if isinstance(value, (int, float)):
            return float(value)
        if isinstance(value, str):
            numeric = "".join(ch for ch in value if ch.isdigit() or ch in ",.")
            if numeric:
                try:
                    return float(numeric.replace(",", "."))
                except ValueError:
                    return 0.0
        return 0.0

    @classmethod
    def _wisc_rows(cls, test: dict | None, context: dict, domain: str):
        payload = (test or {}).get("classified_payload") or {}
        ncp_table = cls._wisc_ncp_table(context)
        subtests = {item.get("codigo"): item for item in payload.get("subtestes") or []}
        domain_codes = {
            "funcoes_executivas": ["SM", "CN", "CO", "RM"],
            "linguagem": ["SM", "VC", "CO"],
            "gnosias_praxias": ["RM", "CB"],
        }
        rows = [[
            "Testes Utilizados",
            "Escore Máximo",
            "Escore Médio",
            "Escore Mínimo",
            "Escore Bruto",
            "Classificação",
        ]]
        for code in domain_codes.get(domain, []):
            item = subtests.get(code)
            if not item:
                continue
            score_max, score_mid, score_min = cls._wisc_reference_scores(
                ncp_table, code
            )
            rows.append(
                [
                    item.get("subteste") or "-",
                    score_max,
                    score_mid,
                    score_min,
                    cls._num(item.get("escore_bruto")),
                    item.get("classificacao") or "-",
                ]
            )
        if domain == "linguagem":
            rows.append(
                [
                    "Fala Espontânea",
                    "Dentro do esperado para a sua idade",
                    "",
                    "",
                    "",
                    "",
                ]
            )
        return rows if len(rows) > 1 else None

    @classmethod
    def _bpa_rows(cls, test: dict | None, context: dict | None = None):
        payload = (test or {}).get("classified_payload") or {}
        labels = {
            "ac": "Atenção Concentrada - AC",
            "ad": "Atenção Dividida - AD",
            "aa": "Atenção Alternada - AA",
            "ag": "Atenção Geral - AG",
        }
        rows = [["ATENÇÃO BPA", "Pontos", "Percentil", "Classificação"]]
        for item in payload.get("subtestes") or []:
            rows.append(
                [
                    labels.get(
                        item.get("codigo"),
                        item.get("subteste") or item.get("codigo") or "-",
                    ),
                    cls._num(item.get("total") or item.get("brutos")),
                    cls._num(item.get("percentil")),
                    item.get("classificacao") or "-",
                ]
            )
        return rows if len(rows) > 1 else None

    @staticmethod
    def _bpa_chart_bytes(test: dict | None):
        chart_data = (test or {}).get("bpa_chart_data") or {}
        return gerar_grafico_bpa_bytes(chart_data)

    @classmethod
    def _fdt_rows(cls, test: dict | None):
        payload = (test or {}).get("classified_payload") or {}
        errors = payload.get("erros") or {}
        rows = [
            [
                "Processo",
                "Tempo Médio",
                "Tempo Obtido",
                "Erros",
                "Desempenho",
                "Classificação",
            ]
        ]
        for item in payload.get("metric_results") or []:
            err = errors.get((item.get("codigo") or "").lower(), {})
            error_count = err.get("qtde_erros") if err else None
            table_result = cls._fdt_table_result(item.get("percentil_num"))
            rows.append(
                [
                    item.get("nome") or item.get("codigo") or "-",
                    cls._num(item.get("media")),
                    cls._num(item.get("valor")),
                    "" if error_count is None else cls._num(error_count),
                    cls._num(table_result["desempenho"]),
                    table_result["classificacao"],
                ]
            )
        return rows if len(rows) > 1 else None

    @classmethod
    def _fdt_table_result(cls, percentile: float | int | None) -> dict[str, str | int]:
        try:
            percentile_value = float(percentile)
        except (TypeError, ValueError):
            return {"desempenho": "-", "classificacao": "-"}

        if percentile_value <= 24:
            return {
                "desempenho": 5,
                "classificacao": "Indicativo de Déficit",
            }
        if percentile_value <= 30:
            return {
                "desempenho": 25,
                "classificacao": "Indicativo de Dificuldade Discreta",
            }
        if percentile_value > 65:
            return {
                "desempenho": 75,
                "classificacao": "Sem Indicativo de Déficit",
            }
        return {
            "desempenho": 50,
            "classificacao": "Sem Indicativo de Déficit",
        }

    @classmethod
    def _etdah_rows(cls, test: dict | None):
        payload = (test or {}).get("classified_payload") or {}
        results = payload.get("results") or {}
        rows = [["Escala", "Pontos Brutos", "Média", "Percentil", "Classificação"]]

        if any(key in results for key in ("D", "I", "AE", "AAMA", "H")):
            ordered_items = [results.get(key) for key in ("D", "I", "AE", "AAMA", "H")]
        elif any(
            key in results
            for key in ("fator_1", "fator_2", "fator_3", "fator_4", "escore_geral")
        ):
            ordered_items = [
                results.get(key)
                for key in ("fator_1", "fator_2", "fator_3", "fator_4", "escore_geral")
            ]
        else:
            ordered_items = list(results.values())

        for item in ordered_items:
            if not item:
                continue
            rows.append(
                [
                    item.get("name") or "-",
                    cls._num(item.get("raw_score")),
                    cls._num(item.get("mean")),
                    cls._num(item.get("percentile_text") or item.get("percentil")),
                    item.get("classification") or item.get("classificacao") or "-",
                ]
            )
        return rows if len(rows) > 1 else None

    @classmethod
    def _ebadep_rows(cls, test: dict | None):
        payload = (test or {}).get("classified_payload") or {}
        result = payload.get("result") or payload
        if not result:
            return None
        return [
            ["Indicador", "Valor"],
            [
                "Escore Total",
                cls._num(result.get("escore_total") or result.get("pontuacao_total")),
            ],
            [
                "Percentil",
                cls._num(
                    result.get("percentil")
                    or ((result.get("normas") or {}).get("percentil"))
                ),
            ],
            ["Classificação", result.get("classificacao") or "-"],
        ]

    @classmethod
    def _scared_rows(cls, test: dict | None):
        payload = (test or {}).get("classified_payload") or {}
        form_type = payload.get("form_type") or "child"
        factor_labels = {
            "panico_somatico": "Pânico/Sintomas Somáticos",
            "ansiedade_generalizada": "Ansiedade Generalizada",
            "ansiedade_separacao": "Ansiedade Separação",
            "fobia_social": "Fobia Social",
            "evitacao_escolar": "Evitação Escolar",
            "total": "Total",
        }
        if form_type == "parent":
            rows = [["Fator", "Pontos Bruto", "Nota de Corte", "Percentil", "Classificação"]]
        else:
            rows = [["Escala", "Pontos Brutos", "Média", "Percentil", "Classificação"]]
        for item in payload.get("analise_geral") or []:
            factor = (item.get("fator") or "").lower()
            rows.append(
                [
                    factor_labels.get(factor, str(item.get("fator", "")).replace("_", " ").title()),
                    cls._num(item.get("escore_bruto")),
                    cls._num(item.get("nota_corte") if form_type == "parent" else item.get("media")),
                    cls._num(item.get("percentil") or item.get("percentual")),
                    item.get("classificacao") or "-",
                ]
            )
        return rows if len(rows) > 1 else None

    @classmethod
    def _epq_rows(cls, test: dict | None):
        payload = (test or {}).get("classified_payload") or {}
        factors = payload.get("fatores") or {}
        rows = [
            ["EPQ-J - Inventário de Personal", "EPQ-J - Inventário de Personal", "EPQ-J - Inventário de Personal", "EPQ-J - Inventário de Personal", "EPQ-J - Inventário de Personal"],
            ["", "P", "E", "N", "S"],
        ]
        brute_values = []
        percentil_values = []
        classificacao_values = []
        for key in ("P", "E", "N", "S"):
            item = factors.get(key, {})
            brute_values.append(cls._num(item.get("escore")))
            percentil_values.append(cls._num(item.get("percentil")))
            classificacao_values.append((item.get("classificacao") or "-").upper())
        rows.append(["RESULTADO BRUTO"] + brute_values)
        rows.append(["PERCENTIL"] + percentil_values)
        rows.append(["CLASSIFICAÇÃO"] + classificacao_values)
        return rows if len(rows) > 2 else None

    @classmethod
    def _srs2_rows(cls, test: dict | None):
        payload = (test or {}).get("classified_payload") or {}
        rows = [["Escala", "Pontos Bruto", "T-Score", "Percentil", "Classificação"]]
        for item in payload.get("resultados") or []:
            rows.append(
                [
                    item.get("nome") or "-",
                    cls._num(item.get("bruto")),
                    cls._num(item.get("tscore")),
                    cls._num(item.get("percentil")),
                    item.get("classificacao") or "-",
                ]
            )
        return rows if len(rows) > 1 else None

    @classmethod
    def _build_chart_png(
        cls,
        kind: str,
        title: str,
        labels: list[str],
        values: list[float],
        ylabel: str,
        extra: dict | None = None,
    ) -> bytes | None:
        if not labels or not values:
            return None
        extra = extra or {}
        regular_font = cls._chart_font()
        title_font = regular_font.copy() if regular_font else None
        if title_font:
            title_font.set_size(16)
        label_font = regular_font.copy() if regular_font else None
        if label_font:
            label_font.set_size(10)
        fig, ax = plt.subplots(figsize=(8, 4.5), dpi=180)
        fig.patch.set_facecolor("white")
        cls._apply_figure_border(fig)
        ax.set_facecolor("white")
        title_kwargs = {"color": "#3D5F2A", "pad": 10, "fontweight": "normal"}
        if title_font:
            title_kwargs["fontproperties"] = title_font
        else:
            title_kwargs["fontsize"] = 16
        ax.set_title(title, **title_kwargs)
        ax.set_ylabel(ylabel, fontproperties=label_font or None, fontsize=None if label_font else 9)
        ax.grid(axis="y", color="#BFBFBF", linewidth=1)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(True)
        ax.spines["left"].set_color("#BFBFBF")
        ax.spines["right"].set_color("#BFBFBF")
        ax.spines["left"].set_linewidth(0.8)
        ax.spines["right"].set_linewidth(0.8)
        ax.spines["bottom"].set_color("#94A3B8")

        if kind == "bar":
            colors = extra.get("colors") or "#2F6DB3"
            bars = ax.bar(
                labels,
                values,
                color=colors,
                edgecolor=extra.get("edgecolor", "#1E3A5F"),
                linewidth=0.6,
                yerr=extra.get("errors"),
                ecolor="#1E3A5F",
                capsize=3 if extra.get("errors") else 0,
            )
            for bar, value in zip(bars, values):
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    bar.get_height(),
                    cls._num(value),
                    ha="center",
                    va="bottom",
                    fontsize=8,
                )
        else:
            x_positions = np.arange(1, len(labels) + 1)
            ax.plot(x_positions, values, marker="o", color="#5E8E3E", linewidth=3)
            if extra.get("expected"):
                ax.plot(
                    x_positions,
                    extra["expected"],
                    linestyle="--",
                    color="#8FBC6B",
                    linewidth=1.5,
                    label="Esperado",
                )
            if extra.get("minimum"):
                ax.plot(
                    x_positions,
                    extra["minimum"],
                    linestyle=":",
                    color="#D6A85C",
                    linewidth=1.5,
                    label="Mínimo",
                )
            if extra.get("expected") or extra.get("minimum"):
                ax.legend(fontsize=8)
            ax.set_xticks(x_positions)
            ax.set_xticklabels(labels)

        ymax = extra.get("y_max")
        if ymax is None:
            ymax = max(values + extra.get("expected", []) + extra.get("minimum", []) + [1]) * 1.2
        ymin = extra.get("y_min", 0)
        ax.set_ylim(ymin, ymax)
        if extra.get("y_ticks") is not None:
            ax.set_yticks(extra["y_ticks"])
        ax.tick_params(axis="x", labelrotation=extra.get("x_rotation", 0), labelsize=8, length=0)
        ax.tick_params(axis="y", labelsize=8, length=0)
        if label_font:
            for tick in ax.get_xticklabels() + ax.get_yticklabels():
                tick.set_fontproperties(label_font)
        plt.tight_layout()

        output = BytesIO()
        fig.savefig(output, format="png", bbox_inches="tight", facecolor="white")
        plt.close(fig)
        return output.getvalue()

    @classmethod
    def _build_wisc_chart_png(
        cls, labels: list[str], values: list[float], errors: list[float] | None = None
    ) -> bytes | None:
        if not labels or not values:
            return None

        errors = errors or [0.0] * len(values)
        regular_font = cls._chart_font()
        title_font = regular_font.copy() if regular_font else None
        if title_font:
            title_font.set_size(23)
        x_label_font = regular_font.copy() if regular_font else None
        if x_label_font:
            x_label_font.set_size(13)
        y_label_font = regular_font.copy() if regular_font else None
        if y_label_font:
            y_label_font.set_size(16)
        y_tick_font = regular_font.copy() if regular_font else None
        if y_tick_font:
            y_tick_font.set_size(12)
        x = list(range(len(labels)))

        fig, ax = plt.subplots(figsize=(10.2, 5.1), dpi=150)
        fig.patch.set_facecolor("#FFFFFF")
        cls._apply_figure_border(fig)
        ax.set_facecolor("white")
        fig.subplots_adjust(left=0.09, right=0.965, bottom=0.15, top=0.87)
        bands = [
            (40, 84, "#FFF86B"),
            (84, 116, "#9DB9DE"),
            (116, 160, "#B8EB70"),
        ]
        for y0, y1, color in bands:
            ax.axhspan(y0, y1, facecolor=color, alpha=1.0, zorder=0)

        ax.set_ylim(40, 160)
        ax.set_xlim(-0.45, len(labels) - 0.55)
        ax.yaxis.set_major_locator(plt.MultipleLocator(10))
        ax.yaxis.set_minor_locator(plt.MultipleLocator(5))
        ax.grid(axis="y", which="major", color="#D9D9D9", linewidth=0.9, zorder=1)
        ax.grid(axis="y", which="minor", color="#EDEDED", linewidth=0.5, zorder=1)
        ax.grid(axis="x", color="#E6E6E6", linewidth=0.8, zorder=1)
        ax.set_axisbelow(True)

        color_by_label = {
            "ICV": "#4F81BD",
            "IOP": "#C0504D",
            "IMO": "#76933C",
            "IVP": "#8064A2",
            "QI Total": "#F79646",
            "GAI": "#4A76AF",
            "CPI": "#7F2A24",
        }
        colors = [color_by_label.get(label, "#5B9BD5") for label in labels]

        ax.bar(
            x,
            values,
            width=0.72,
            color=colors,
            edgecolor="none",
            yerr=errors,
            ecolor="black",
            capsize=3,
            zorder=3,
        )

        for xi, value in zip(x, values):
            ax.text(
                xi,
                46,
                cls._num(value),
                ha="center",
                va="center",
                fontsize=12,
                color="#FFFFFF",
                fontproperties=regular_font,
                zorder=4,
            )

        title_kwargs = {
            "color": "#70AD47",
            "pad": 2,
            "fontweight": "normal",
        }
        if title_font:
            title_kwargs["fontproperties"] = title_font
        else:
            title_kwargs["fontsize"] = 23
        ax.set_title("WISC-IV INDICES QIs", **title_kwargs)
        ax.set_ylabel(
            "Pontos Compostos",
            color="#000000",
            fontproperties=y_label_font or regular_font,
        )
        ax.set_xticks(x, labels)
        for label in ax.get_xticklabels():
            label.set_fontsize(13)
            if x_label_font:
                label.set_fontproperties(x_label_font)
        ax.tick_params(axis="y", labelsize=12, colors="#444444")
        ax.tick_params(axis="x", pad=1, colors="#444444")
        if y_tick_font:
            for label in ax.get_yticklabels():
                label.set_fontproperties(y_tick_font)

        ax.spines["left"].set_color("#BFBFBF")
        ax.spines["left"].set_linewidth(0.8)
        ax.spines["right"].set_color("#BFBFBF")
        ax.spines["right"].set_linewidth(0.8)
        ax.spines["bottom"].set_color("#666666")
        ax.spines["bottom"].set_linewidth(0.6)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(True)

        output = BytesIO()
        fig.savefig(output, format="png", facecolor="white")
        plt.close(fig)
        return output.getvalue()

    @classmethod
    def _wisc_chart(cls, test: dict | None):
        payload = cls._wisc_payload(test)
        labels, values, errors = [], [], []
        index_labels = {
            "icv": "ICV",
            "iop": "IOP",
            "imt": "IMO",
            "ivp": "IVP",
        }
        for code in ("icv", "iop", "imt", "ivp"):
            item = next(
                (entry for entry in payload.get("indices") or [] if entry.get("indice") == code),
                None,
            )
            if not item:
                continue
            labels.append(index_labels[code])
            values.append(float(item.get("escore_composto") or 0))
            interval = item.get("intervalo_confianca") or (0, 0)
            if isinstance(interval, (list, tuple)) and len(interval) == 2:
                errors.append(abs(float(interval[1]) - float(interval[0])) / 2)
            else:
                errors.append(0.0)
        qit = payload.get("qit_data") or {}
        if qit.get("escore_composto"):
            labels.append("QI Total")
            values.append(float(qit.get("escore_composto") or 0))
            interval = qit.get("intervalo_confianca") or (0, 0)
            if isinstance(interval, (list, tuple)) and len(interval) == 2:
                errors.append(abs(float(interval[1]) - float(interval[0])) / 2)
            else:
                errors.append(0.0)
        for key, label in (("gai_data", "GAI"), ("cpi_data", "CPI")):
            item = payload.get(key) or {}
            if not item.get("escore_composto"):
                continue
            labels.append(label)
            values.append(float(item.get("escore_composto") or 0))
            interval = item.get("intervalo_confianca") or (0, 0)
            if isinstance(interval, (list, tuple)) and len(interval) == 2:
                errors.append(abs(float(interval[1]) - float(interval[0])) / 2)
            else:
                errors.append(0.0)
        return cls._build_wisc_chart_png(labels, values, errors)

    @classmethod
    def _wasi_chart(cls, test: dict | None):
        if not test:
            return None
        payload = cls._wasi_payload(test)
        composites = payload.get("composites") or {}
        labels, values, errors = [], [], []
        index_labels = {
            "qi_execucao": "QIE",
            "qi_verbal": "QIV",
            "qit_4": "QI TOTAL",
        }
        for code, label in index_labels.items():
            composite = composites.get(code) or {}
            qi = composite.get("qi")
            if qi is not None:
                labels.append(label)
                values.append(float(qi))
                interval = composite.get("confidence_interval") or (0, 0)
                if isinstance(interval, (list, tuple)) and len(interval) == 2:
                    errors.append(abs(float(interval[1]) - float(interval[0])) / 2)
                else:
                    errors.append(0.0)
        return cls._build_wisc_chart_png(labels, values, errors)

    @classmethod
    def _wasi_chart_image(cls, test: dict | None):
        if not test:
            return None
        payload = cls._wasi_payload(test)
        composites = payload.get("composites") or {}

        qie = composites.get("qi_execucao", {}).get("qi")
        qiv = composites.get("qi_verbal", {}).get("qi")
        qit = composites.get("qit_4", {}).get("qi")

        if qiv is None or qie is None or qit is None:
            return None

        return gerar_grafico_wasi_bytes(
            qi_verbal=int(qiv),
            qi_execucao=int(qie),
            qi_total=int(qit),
        )

    @classmethod
    def _ravlt_chart(cls, test: dict | None):
        payload = (test or {}).get("classified_payload") or {}
        chart = payload.get("chart") or {}
        labels = chart.get("labels") or []
        series = chart.get("series") or []
        if not labels or not series:
            return None
        y_axis = chart.get("y_axis") or {}

        regular_font = cls._chart_font()

        fig, ax = plt.subplots(figsize=(10.8, 4.6), dpi=120)
        background = "#ffffff"
        fig.patch.set_facecolor(background)
        cls._apply_figure_border(fig)
        ax.set_facecolor(background)

        for item in series:
            values = [float(value or 0) for value in item.get("values") or []]
            if len(values) != len(labels):
                continue
            ax.plot(
                labels,
                values,
                linewidth=3,
                color=item.get("color") or "#70AD47",
                solid_capstyle="round",
                label=item.get("label") or "Serie",
            )

        title_kwargs = {"color": "#5B8A3C", "pad": 18}
        title_font = regular_font.copy()
        title_font.set_size(24)
        title_kwargs["fontproperties"] = title_font
        ax.set_title(chart.get("title") or "RAVLT", **title_kwargs)

        ax.set_ylim(float(y_axis.get("min") or 0), float(y_axis.get("max") or 21))
        ax.set_yticks(y_axis.get("ticks") or [0, 5, 10, 15, 20])
        ax.tick_params(axis="y", labelsize=12, colors="#555555", length=0, pad=10)
        ax.tick_params(axis="x", labelsize=13, colors="#555555", length=0, pad=10)

        y_tick_font = regular_font.copy()
        y_tick_font.set_size(12)
        for label in ax.get_yticklabels():
            label.set_fontproperties(y_tick_font)
        x_tick_font = regular_font.copy()
        x_tick_font.set_size(13)
        for label in ax.get_xticklabels():
            label.set_fontproperties(x_tick_font)

        ax.grid(axis="y", color="#c9c9c9", linewidth=1)
        ax.grid(axis="x", visible=False)

        ax.spines["top"].set_visible(False)
        ax.spines["bottom"].set_visible(False)
        ax.spines["left"].set_visible(True)
        ax.spines["right"].set_visible(True)
        ax.spines["left"].set_color("#BFBFBF")
        ax.spines["right"].set_color("#BFBFBF")
        ax.spines["left"].set_linewidth(0.8)
        ax.spines["right"].set_linewidth(0.8)

        legend_kwargs = {
            "loc": "lower center",
            "bbox_to_anchor": (0.5, -0.35),
            "ncol": 3,
            "frameon": False,
            "handlelength": 2.0,
            "handletextpad": 0.4,
            "columnspacing": 1.2,
        }
        legend_font = regular_font.copy()
        legend_font.set_size(13)
        legend_kwargs["prop"] = legend_font
        legend = ax.legend(**legend_kwargs)
        for text in legend.get_texts():
            text.set_color("#666666")

        plt.subplots_adjust(left=0.06, right=0.98, top=0.83, bottom=0.28)
        output = BytesIO()
        fig.savefig(output, format="png", bbox_inches="tight", facecolor=background)
        plt.close(fig)
        return output.getvalue()

    @classmethod
    def _fdt_chart(cls, test: dict | None, automatic: bool):
        payload = (test or {}).get("classified_payload") or {}
        charts = payload.get("charts") or {}
        chart_key = "automaticos" if automatic else "controlados"
        chart = charts.get(chart_key) or {}

        categories = chart.get("categories") or []
        series = chart.get("series") or []
        if not categories or not series:
            return None

        regular_font = cls._chart_font()

        width_cm = 14
        height_cm = 10
        fig, ax = plt.subplots(
            figsize=(width_cm / 2.54, height_cm / 2.54),
            dpi=300,
        )
        background = "#F2F2F2"
        grid_color = "#C9C9C9"
        title_color = "#548235"
        fig.patch.set_facecolor(background)
        ax.set_facecolor(background)

        y = np.arange(len(categories))
        bar_height = 0.22 if automatic else 0.16
        offsets = [
            (index - ((len(series) - 1) / 2)) * bar_height
            for index in range(len(series))
        ]

        for offset, item in zip(offsets, series):
            values = [float(value or 0) for value in item.get("values") or []]
            if len(values) != len(categories):
                continue
            ax.barh(
                y + offset,
                values,
                height=bar_height,
                color=item.get("color") or "#4472C4",
                edgecolor="none",
                label=item.get("label") or "FDT",
            )

        title_kwargs = {"color": title_color, "pad": 10, "fontweight": "normal"}
        title_font = regular_font.copy()
        title_font.set_size(12)
        title_kwargs["fontproperties"] = title_font
        ax.set_title(chart.get("title") or "FDT", **title_kwargs)

        ax.set_xlim(0, 82)
        ax.set_xticks(np.arange(0, 81, 10))
        ax.set_yticks(y)
        ax.set_yticklabels(categories)
        ax.invert_yaxis()

        for label in ax.get_yticklabels():
            tick_font = regular_font.copy()
            tick_font.set_size(8)
            label.set_fontproperties(tick_font)
            label.set_color("#595959")
        for label in ax.get_xticklabels():
            tick_font = regular_font.copy()
            tick_font.set_size(7)
            label.set_fontproperties(tick_font)
            label.set_color("#595959")

        ax.tick_params(axis="x", colors="#595959", length=0)
        ax.tick_params(axis="y", colors="#595959", length=0, pad=5)
        ax.xaxis.grid(True, color=grid_color, linewidth=0.6)
        ax.yaxis.grid(False)
        ax.set_axisbelow(True)

        for spine in ["top", "right", "left", "bottom"]:
            ax.spines[spine].set_visible(False)

        legend_kwargs = {
            "loc": "lower center",
            "bbox_to_anchor": (0.5, -0.18),
            "ncol": 2 if automatic else 4,
            "frameon": False,
            "handlelength": 0.6,
            "handletextpad": 0.3,
            "columnspacing": 1.0 if automatic else 0.7,
        }
        legend_font = regular_font.copy()
        legend_font.set_size(7 if automatic else 6.8)
        legend_kwargs["prop"] = legend_font
        legend = ax.legend(**legend_kwargs)
        for text in legend.get_texts():
            text.set_color("#595959")

        plt.subplots_adjust(
            left=0.38,
            right=0.96,
            top=0.84,
            bottom=0.22,
        )
        output = BytesIO()
        fig.savefig(
            output,
            format="png",
            dpi=300,
            facecolor=fig.get_facecolor(),
        )
        plt.close(fig)
        return output.getvalue()

    @classmethod
    def _etdah_chart(cls, test: dict | None):
        payload = (test or {}).get("classified_payload") or {}
        results = payload.get("results") or {}
        if any(key in results for key in ("D", "I", "AE", "AAMA", "H")):
            labels = []
            values = []
            label_map = {
                "D": "F1 - D",
                "I": "F2 - I",
                "AE": "F3 - AE",
                "AAMA": "F4 - AMAA",
                "H": "F5 - H",
            }
            for key in ("D", "I", "AE", "AAMA", "H"):
                item = results.get(key)
                if not item:
                    continue
                labels.append(label_map[key])
                values.append(
                    cls._extract_percentile_value(
                        item.get("percentile_text") or item.get("percentil")
                    )
                )
            return cls._build_etdah_ad_chart_png(labels, values)

        labels = []
        values = []
        label_map = {
            "fator_1": "F1 - RE",
            "fator_2": "F2 - H/I",
            "fator_3": "F3 - CA",
            "fator_4": "F4 - A",
            "escore_geral": "EG",
        }
        for key in ("fator_1", "fator_2", "fator_3", "fator_4", "escore_geral"):
            item = results.get(key)
            if not item:
                continue
            labels.append(label_map[key])
            values.append(
                cls._extract_percentile_value(
                    item.get("percentile_text") or item.get("percentil")
                )
            )
        return cls._build_etdah_pais_chart_png(labels, values)

    @classmethod
    def _build_etdah_ad_chart_png(
        cls, labels: list[str], values: list[float]
    ) -> bytes | None:
        if not labels or not values:
            return None

        regular_font = cls._chart_font()
        title_font = regular_font.copy() if regular_font else None
        if title_font:
            title_font.set_size(22)
        axis_font = regular_font.copy() if regular_font else None
        if axis_font:
            axis_font.set_size(12)
        note_font = regular_font.copy() if regular_font else None
        if note_font:
            note_font.set_size(9)

        x = np.arange(len(labels))
        fig, ax = plt.subplots(figsize=(10, 4.5), dpi=150)
        fig.patch.set_facecolor("#FFFFFF")
        cls._apply_figure_border(fig)
        ax.set_facecolor("#FFFFFF")

        ax.axhspan(0, 25, color="#EAF4E2", alpha=0.8)
        ax.axhspan(25, 45, color="#F6F1D5", alpha=0.8)
        ax.axhspan(45, 65, color="#FCE8C8", alpha=0.8)
        ax.axhspan(65, 85, color="#F5D1C8", alpha=0.8)
        ax.axhspan(85, 100, color="#E8B6B0", alpha=0.8)

        ax.plot(
            x,
            values,
            color="#375623",
            linewidth=3,
            marker="o",
            markersize=8,
            markerfacecolor="#FFFFFF",
            markeredgecolor="#375623",
            markeredgewidth=2.2,
        )

        for xi, yi in zip(x, values):
            text_kwargs = {
                "ha": "center",
                "va": "bottom",
                "fontsize": 11,
                "color": "#375623",
                "fontweight": "bold",
            }
            if axis_font:
                text_kwargs["fontproperties"] = axis_font
            ax.text(xi, yi + 3, cls._num(yi), **text_kwargs)

        title_kwargs = {"color": "#2F4F1F", "pad": 18}
        if title_font:
            title_kwargs["fontproperties"] = title_font
        else:
            title_kwargs["fontsize"] = 22
        ax.set_title("E-TDAH-AD", **title_kwargs)

        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        ax.set_xlabel(
            "Fator",
            fontsize=12,
            color="#2F4F1F",
            labelpad=14,
            fontproperties=axis_font or None,
        )
        ax.set_ylim(0, 105)
        ax.set_ylabel(
            "Percentil",
            fontsize=13,
            color="#2F4F1F",
            labelpad=12,
            fontproperties=axis_font or None,
        )
        ax.grid(axis="y", color="#D9D9D9", linewidth=0.8)
        ax.grid(axis="x", visible=False)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(True)
        ax.spines["left"].set_color("#BFBFBF")
        ax.spines["right"].set_color("#BFBFBF")
        ax.spines["bottom"].set_color("#BFBFBF")
        ax.tick_params(axis="y", colors="#404040")
        ax.tick_params(axis="x", length=0, colors="#2F4F1F")
        if axis_font:
            for tick in ax.get_xticklabels() + ax.get_yticklabels():
                tick.set_fontproperties(axis_font)

        note_kwargs = {
            "transform": ax.transAxes,
            "ha": "center",
            "va": "center",
            "fontsize": 9,
            "color": "#404040",
        }
        if note_font:
            note_kwargs["fontproperties"] = note_font
        ax.text(
            0.5,
            -0.23,
            "F1-D: Desatenção   |   F2-I: Impulsividade   |   F3-AE: Aspectos Emocionais   |   F4-AMAA: Autorregulação   |   F5-H: Hiperatividade",
            **note_kwargs,
        )

        plt.tight_layout()
        output = BytesIO()
        fig.savefig(
            output,
            format="png",
            dpi=300,
            facecolor=fig.get_facecolor(),
            bbox_inches="tight",
        )
        plt.close(fig)
        return output.getvalue()

    @classmethod
    def _build_etdah_pais_chart_png(
        cls, labels: list[str], values: list[float]
    ) -> bytes | None:
        if not labels or not values:
            return None

        regular_font = cls._chart_font()
        title_font = regular_font.copy() if regular_font else None
        if title_font:
            title_font.set_size(22)
        axis_font = regular_font.copy() if regular_font else None
        if axis_font:
            axis_font.set_size(12)
        note_font = regular_font.copy() if regular_font else None
        if note_font:
            note_font.set_size(9)

        x = np.arange(len(labels))
        fig, ax = plt.subplots(figsize=(10, 4.5), dpi=150)
        fig.patch.set_facecolor("#FFFFFF")
        cls._apply_figure_border(fig)
        ax.set_facecolor("#FFFFFF")

        ax.axhspan(0, 25, color="#EAF4E2", alpha=0.8)
        ax.axhspan(25, 45, color="#F6F1D5", alpha=0.8)
        ax.axhspan(45, 65, color="#FCE8C8", alpha=0.8)
        ax.axhspan(65, 85, color="#F5D1C8", alpha=0.8)
        ax.axhspan(85, 100, color="#E8B6B0", alpha=0.8)

        ax.plot(
            x,
            values,
            color="#375623",
            linewidth=3,
            marker="o",
            markersize=8,
            markerfacecolor="#FFFFFF",
            markeredgecolor="#375623",
            markeredgewidth=2.2,
        )

        for xi, yi in zip(x, values):
            text_kwargs = {
                "ha": "center",
                "va": "bottom",
                "fontsize": 11,
                "color": "#375623",
                "fontweight": "bold",
            }
            if axis_font:
                text_kwargs["fontproperties"] = axis_font
            ax.text(xi, yi + 3, cls._num(yi), **text_kwargs)

        title_kwargs = {"color": "#2F4F1F", "pad": 18}
        if title_font:
            title_kwargs["fontproperties"] = title_font
        else:
            title_kwargs["fontsize"] = 22
        ax.set_title("E-TDAH-PAIS", **title_kwargs)

        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        ax.set_xlabel(
            "Fator",
            color="#2F4F1F",
            labelpad=14,
            fontproperties=axis_font or None,
        )
        ax.set_ylim(0, 105)
        ax.set_ylabel(
            "Percentil",
            color="#2F4F1F",
            labelpad=12,
            fontproperties=axis_font or None,
        )

        ax.grid(axis="y", color="#D9D9D9", linewidth=0.8)
        ax.grid(axis="x", visible=False)

        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(True)
        ax.spines["left"].set_color("#BFBFBF")
        ax.spines["right"].set_color("#BFBFBF")
        ax.spines["bottom"].set_color("#BFBFBF")

        ax.tick_params(axis="y", colors="#404040")
        ax.tick_params(axis="x", length=0, colors="#2F4F1F")
        if axis_font:
            for tick in ax.get_xticklabels() + ax.get_yticklabels():
                tick.set_fontproperties(axis_font)

        note_kwargs = {
            "transform": ax.transAxes,
            "ha": "center",
            "va": "center",
            "fontsize": 9,
            "color": "#404040",
        }
        if note_font:
            note_kwargs["fontproperties"] = note_font
        ax.text(
            0.5,
            -0.23,
            "F1-RE: Regulação Emocional   |   F2-H/I: Hiperatividade/Impulsividade   |   F3-CA: Comportamento Adaptativo   |   F4-A: Atenção   |   EG: Escore Geral",
            **note_kwargs,
        )

        plt.tight_layout()
        output = BytesIO()
        fig.savefig(
            output,
            format="png",
            dpi=300,
            facecolor=fig.get_facecolor(),
            bbox_inches="tight",
        )
        plt.close(fig)
        return output.getvalue()

    @classmethod
    def _scared_chart(cls, test: dict | None):
        payload = (test or {}).get("classified_payload") or {}
        labels = []
        values = []
        for item in payload.get("analise_geral") or []:
            labels.append(str(item.get("fator", "")).replace("_", " ").title()[:14])
            values.append(float(item.get("percentil") or item.get("percentual") or 0))
        return cls._build_chart_png(
            "bar", "Perfil no SCARED", labels, values, "Percentil"
        )

    @classmethod
    def _epq_chart(cls, test: dict | None):
        payload = (test or {}).get("classified_payload") or {}
        factors = payload.get("fatores") or {}
        labels = []
        values = []
        for key in ("P", "E", "N", "S"):
            item = factors.get(key)
            if not item:
                continue
            labels.append(key)
            values.append(float(item.get("percentil") or 0))
        return cls._build_epq_chart_png(labels, values)

    @classmethod
    def _build_epq_chart_png(cls, labels: list[str], values: list[float]) -> bytes | None:
        if not labels or not values:
            return None

        regular_font = cls._chart_font()
        title_font = regular_font.copy() if regular_font else None
        if title_font:
            title_font.set_size(18)
        label_font = regular_font.copy() if regular_font else None
        if label_font:
            label_font.set_size(11)

        x = np.arange(1, len(labels) + 1)

        fig, ax = plt.subplots(figsize=(7.5, 4), dpi=200)
        fig.patch.set_facecolor("#FFFFFF")
        cls._apply_figure_border(fig)
        ax.set_facecolor("#FFFFFF")

        ax.plot(x, values, color="#5E8E3E", linewidth=3)
        ax.scatter(x, values, color="#5E8E3E", s=40)

        for xi, yi in zip(x, values):
            text_kwargs = {
                "ha": "center",
                "fontsize": 11,
                "color": "#404040",
            }
            if label_font:
                text_kwargs["fontproperties"] = label_font
            ax.text(xi, yi + 3, cls._num(yi), **text_kwargs)

        title_kwargs = {
            "color": "#3D5F2A",
            "pad": 10,
        }
        if title_font:
            title_kwargs["fontproperties"] = title_font
        else:
            title_kwargs["fontsize"] = 18
        ax.set_title("PERCENTIL - EPQ-J", **title_kwargs)

        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        if label_font:
            for tick in ax.get_xticklabels() + ax.get_yticklabels():
                tick.set_fontproperties(label_font)
        else:
            ax.tick_params(axis="x", labelsize=11)
            ax.tick_params(axis="y", labelsize=11)

        ax.set_ylim(0, 90)
        ax.set_yticks(np.arange(0, 91, 10))
        ax.grid(True, axis="y", color="#BFBFBF", linewidth=1)
        ax.grid(False, axis="x")

        ax.spines["top"].set_visible(False)
        ax.spines["bottom"].set_visible(False)
        ax.spines["left"].set_visible(True)
        ax.spines["right"].set_visible(True)
        ax.spines["left"].set_color("#BFBFBF")
        ax.spines["right"].set_color("#BFBFBF")
        ax.spines["left"].set_linewidth(0.8)
        ax.spines["right"].set_linewidth(0.8)

        ax.tick_params(axis="x", length=0)
        ax.tick_params(axis="y", length=0)
        plt.tight_layout()

        output = BytesIO()
        fig.savefig(output, format="png", facecolor=fig.get_facecolor())
        plt.close(fig)
        return output.getvalue()

    @classmethod
    def _srs2_chart(cls, test: dict | None):
        payload = (test or {}).get("classified_payload") or {}
        labels = []
        values = []
        for item in payload.get("resultados") or []:
            labels.append((item.get("nome") or "Escala")[:12])
            values.append(float(item.get("tscore") or 0))
        return cls._build_chart_png(
            "bar", "Perfil social no SRS-2", labels, values, "T-Score"
        )
